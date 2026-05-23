#!/usr/bin/env python3
"""
Generate 04_Chuong_4.docx from 04_Chuong_4.md.

Images are in sample_images/chapter4/.
Mobile screenshots (portrait) scale to 8cm width.
Admin/landscape screenshots scale up to 15cm width.
Headings use proper Word Heading 1/2/3 styles for TOC and integration compatibility.
"""

import struct
from pathlib import Path

from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

import make_chapter1_docx as base

HERE = Path(__file__).parent
MD_FILE = HERE / "04_Chuong_4.md"
OUT_FILE = HERE / "04_Chuong_4.docx"
IMG_DIR = HERE / "sample_images" / "chapter4"

MAX_WIDTH_CM        = 15.0
MAX_HEIGHT_CM       = 18.0
MOBILE_MAX_WIDTH_CM = 8.0


def _png_size(path: Path):
    try:
        with path.open("rb") as f:
            sig = f.read(8)
            if sig != b"\x89PNG\r\n\x1a\n":
                return None
            f.read(4)
            chunk_type = f.read(4)
            if chunk_type != b"IHDR":
                return None
            width, height = struct.unpack(">II", f.read(8))
            return width, height
    except Exception:
        return None


def _setup_doc_styles(doc):
    """Apply Times New Roman 1.3× spacing to Heading styles for standalone viewing."""
    BLACK = RGBColor(0, 0, 0)

    specs = [
        ("Heading 1", 16, True,  False, WD_ALIGN_PARAGRAPH.CENTER,  0, 18),
        ("Heading 2", 14, True,  False, WD_ALIGN_PARAGRAPH.LEFT,   10,  4),
        ("Heading 3", 14, True,  True,  WD_ALIGN_PARAGRAPH.LEFT,    8,  4),
    ]
    for name, size, bold, italic, align, before, after in specs:
        try:
            s = doc.styles[name]
        except KeyError:
            continue
        f = s.font
        f.name = "Times New Roman"
        f.size = Pt(size)
        f.bold = bold
        f.italic = italic
        f.color.rgb = BLACK
        pf = s.paragraph_format
        pf.alignment = align
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.3


def _h1_para(doc, text):
    return doc.add_heading(text, level=1)


def _h2_para(doc, text):
    return doc.add_heading(text, level=2)


def _h3_para(doc, text):
    return doc.add_heading(text, level=3)


def chapter4_image_para(doc, filename):
    img_path = IMG_DIR / filename
    if not img_path.exists():
        base.normal_para(doc, f"[Hình: {filename} — không tìm thấy file]")
        return

    size = _png_size(img_path)
    px_w, px_h = size if size else (1, 1)
    aspect = px_w / px_h if px_h else 1.0
    is_portrait = aspect < 1.0

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base._para_fmt(p, before=12, after=0, indent=0)

    if is_portrait:
        width_cm = MOBILE_MAX_WIDTH_CM
        height_cm = width_cm / aspect
        if height_cm > MAX_HEIGHT_CM:
            height_cm = MAX_HEIGHT_CM
            width_cm = height_cm * aspect
    else:
        width_cm = MAX_WIDTH_CM
        height_cm = width_cm / aspect
        if height_cm > MAX_HEIGHT_CM:
            height_cm = MAX_HEIGHT_CM
            width_cm = height_cm * aspect

    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm), height=Cm(height_cm))


def main():
    base.MD_FILE = MD_FILE
    base.OUT_FILE = OUT_FILE
    base.IMG_DIR = IMG_DIR
    base.image_para = chapter4_image_para
    # Override heading functions → use proper Word heading styles
    base.h1_para = _h1_para
    base.h2_para = _h2_para
    base.h3_para = _h3_para

    lines = MD_FILE.read_text(encoding="utf-8").splitlines()
    doc = base.build_doc(lines)
    _setup_doc_styles(doc)
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    main()
