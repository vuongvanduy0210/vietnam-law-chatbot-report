#!/usr/bin/env python3
"""
Generate 01_Chuong_1.docx from 01_Chuong_1.md.
Format: Times New Roman 14pt, 1.3x line spacing, margins 3/1.5/2/2 cm.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
MD_FILE = HERE / "01_Chuong_1.md"
OUT_FILE = HERE / "01_Chuong_1.docx"
IMG_DIR = HERE / "sample_images"

TNR = "Times New Roman"
MONO = "Courier New"


def _set_run_font(run, name=TNR, size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    rf.set(qn("w:cs"), name)
    rPr.insert(0, rf)


def _para_fmt(para, line=1.3, before=0, after=0, indent=None):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.first_line_indent = Cm(indent)


def add_run(para, text, bold=False, italic=False, size=14, mono=False):
    run = para.add_run(text)
    if mono:
        run.font.name = MONO
        run.font.size = Pt(11)
        run.font.bold = False
        run.font.italic = False
    else:
        _set_run_font(run, TNR, size, bold, italic)
    return run


def add_inline(para, text, base_bold=False, base_italic=False, size=14):
    """Add runs to para, parsing **bold**, *italic*, and `code` inline markers."""
    for seg in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**") and len(seg) > 4:
            add_run(para, seg[2:-2], bold=True, italic=base_italic, size=size)
        elif seg.startswith("*") and seg.endswith("*") and len(seg) > 2 and not seg.startswith("**"):
            add_run(para, seg[1:-1], bold=base_bold, italic=True, size=size)
        elif seg.startswith("`") and seg.endswith("`") and len(seg) > 2:
            add_run(para, seg[1:-1], mono=True)
        else:
            add_run(para, seg, bold=base_bold, italic=base_italic, size=size)


def h1_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_fmt(p, before=12, after=12, indent=0)
    add_run(p, text, bold=True, size=16)


def h2_para(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, before=10, after=4, indent=0)
    add_run(p, text, bold=True, size=14)


def h3_para(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, before=8, after=4, indent=0)
    add_run(p, text, bold=True, italic=True, size=14)


def h4_para(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, before=6, after=4, indent=0)
    add_run(p, text, bold=False, italic=True, size=14)


def figure_caption_para(doc, text):
    """Tên hình vẽ: centered, italic, Before=6pt, After=12pt."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_fmt(p, before=6, after=12, indent=0)
    add_run(p, text, bold=False, italic=True, size=13)


def table_caption_para(doc, text):
    """Tên bảng: left-aligned, bold, Before=6pt, After=4pt."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_fmt(p, before=6, after=4, indent=0)
    add_run(p, text, bold=True, italic=False, size=13)


def image_para(doc, filename):
    """Embed an image centered, width=14cm."""
    img_path = IMG_DIR / filename
    if not img_path.exists():
        normal_para(doc, f"[Hình: {filename} — không tìm thấy file]")
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_fmt(p, before=12, after=0, indent=0)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(14))


def normal_para(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, indent=1.25)
    add_inline(p, text)


def list_para(doc, text, bullet="•"):
    p = doc.add_paragraph()
    _para_fmt(p, indent=0)
    p.paragraph_format.left_indent = Cm(1.25)
    add_run(p, bullet + " ", size=14)
    add_inline(p, text)


def code_para(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, line=1.1, indent=0)
    p.paragraph_format.left_indent = Cm(1.0)
    add_run(p, text, mono=True)


def _shade_cell(cell, fill="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def flush_table(doc, buf):
    """Build a Word table from accumulated markdown table lines."""
    rows = [l for l in buf if not re.match(r"^\|[-:| ]+\|$", l)]
    if not rows:
        return
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    ncols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(parsed):
        is_hdr = ri == 0
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            cell_text = row[ci] if ci < len(row) else ""
            if is_hdr:
                _shade_cell(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            add_inline(para, cell_text, base_bold=is_hdr, size=12)


def build_doc(lines):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Remove the auto-created empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    in_code = False
    code_buf = []
    tbl_buf = []
    i = 0

    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        # Code block toggle
        if s.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    code_para(doc, cl)
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # Table accumulation
        if s.startswith("|"):
            tbl_buf.append(s)
            i += 1
            continue
        elif tbl_buf:
            flush_table(doc, tbl_buf)
            tbl_buf = []

        # Blank line
        if not s:
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", s)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            [h1_para, h2_para, h3_para, h4_para][level - 1](doc, text)
            i += 1
            continue

        # Bullet list (- or *)
        if re.match(r"^[-*]\s", s):
            list_para(doc, s[2:].strip())
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", s)
        if m:
            list_para(doc, m.group(2), bullet=f"{m.group(1)}.")
            i += 1
            continue

        # Image embed: [IMG:filename]
        m_img = re.match(r"^\[IMG:([^\]]+)\]$", s)
        if m_img:
            image_para(doc, m_img.group(1))
            i += 1
            continue

        # Figure caption: *Hình X.X. ...*
        m_fig = re.match(r"^\*(Hình\s+.+)\*$", s)
        if m_fig:
            figure_caption_para(doc, m_fig.group(1))
            i += 1
            continue

        # Table caption: *Bảng X.X. ...*
        m_tbl = re.match(r"^\*(Bảng\s+.+)\*$", s)
        if m_tbl:
            table_caption_para(doc, m_tbl.group(1))
            i += 1
            continue

        # Normal paragraph
        normal_para(doc, s)
        i += 1

    if tbl_buf:
        flush_table(doc, tbl_buf)

    return doc


def main():
    lines = MD_FILE.read_text(encoding="utf-8").splitlines()
    doc = build_doc(lines)
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    main()
