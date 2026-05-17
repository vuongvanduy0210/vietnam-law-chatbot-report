from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


TARGET = Path(__file__).resolve().parents[2] / "Vương Văn Duy_Báo cáo.docx"


def set_code(run):
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), "Courier New")


def main() -> None:
    doc = Document(TARGET)
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "ession_id":
                    paragraph = cell.paragraphs[0]
                    for run in paragraph.runs:
                        run.text = ""
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.text = "Session_id"
                    set_code(run)
                    changed += 1
    if changed:
        doc.save(TARGET)
    print(changed)


if __name__ == "__main__":
    main()
