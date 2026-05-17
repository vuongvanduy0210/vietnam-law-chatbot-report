from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX = Path(__file__).resolve().parents[1] / "02_Chuong_2_ban_bo_sung_phan_tich_rag.docx"

NEW_SECTION = [
    "2.8.3. Quản lý lịch sử hội thoại và câu hỏi gợi ý",
    "Lịch sử hội thoại được sử dụng để duy trì mạch trao đổi giữa người dùng và hệ thống. Trong thực tế, người dùng thường không đặt toàn bộ thông tin trong một câu hỏi duy nhất mà có xu hướng hỏi nối tiếp, ví dụ sau khi nhận câu trả lời về một mức phạt, họ có thể hỏi tiếp về trường hợp tái phạm, giấy tờ cần chuẩn bị hoặc thủ tục khiếu nại. Nếu không truyền một phần lịch sử sang lượt xử lý mới, hệ thống sẽ khó hiểu được các đại từ như \"trường hợp đó\", \"mức này\" hoặc \"văn bản trên\" đang nói đến nội dung nào.",
    "Tuy nhiên, lịch sử hội thoại chỉ được xem là ngữ cảnh giao tiếp, không phải căn cứ pháp lý. Đây là điểm rất quan trọng trong hệ thống tư vấn pháp luật. Các câu trả lời trước đó có thể giúp xác định chủ đề đang trao đổi, nhưng mỗi lượt trả lời mới vẫn phải truy hồi lại nguồn pháp luật phù hợp trong lượt hiện tại. Cách làm này tránh tình trạng hệ thống dựa vào nội dung đã trả lời trước đó rồi tiếp tục suy luận mà không kiểm chứng lại với văn bản nguồn.",
    "Để cân bằng giữa khả năng hiểu ngữ cảnh và chi phí xử lý, hệ thống chỉ lấy một số tin nhắn gần nhất của hội thoại thay vì đưa toàn bộ lịch sử vào pipeline. Các câu trả lời của trợ lý thường dài vì chứa phân tích và trích dẫn, nên phần này được rút gọn trước khi gửi sang RAG Service. Nhờ vậy, truy vấn hiện tại vẫn giữ được bối cảnh cần thiết nhưng không làm phình to đầu vào, không tăng độ trễ quá mức và không làm giảm sự tập trung của bước phân tích câu hỏi.",
    "Trình tự lưu dữ liệu trong luồng chat cũng được thiết kế để bảo đảm trải nghiệm ổn định. Khi người dùng gửi câu hỏi, hệ thống tạo hoặc xác định hội thoại hiện tại, sau đó lưu ngay tin nhắn người dùng và phát sự kiện ready để giao diện có thể hiển thị bong bóng chat tức thời. Trong lúc RAG Service xử lý, các sự kiện progress tiếp tục cập nhật tiến trình. Chỉ khi câu trả lời đã được tạo và lưu thành công, hệ thống mới phát sự kiện done kèm tin nhắn của trợ lý. Nhờ đó, dữ liệu hiển thị trên giao diện luôn khớp với dữ liệu đã lưu trong cơ sở dữ liệu.",
    "Với hội thoại mới, tiêu đề hội thoại được suy ra từ chủ đề chính của câu hỏi nếu bước phân tích truy vấn xác định được; nếu không, hệ thống dùng một phần nội dung câu hỏi đầu tiên làm tiêu đề tạm. Cơ chế này giúp danh sách hội thoại có tên dễ nhận biết mà không cần người dùng đặt tên thủ công. Về phía người dùng, họ có thể quay lại một hội thoại cũ, tiếp tục hỏi trong cùng bối cảnh và vẫn xem được các nguồn tham chiếu đã được lưu cùng từng câu trả lời.",
    "Sau khi câu trả lời chính hoàn tất, hệ thống sinh thêm các câu hỏi gợi ý để hỗ trợ người dùng tiếp tục khai thác vấn đề pháp lý. Các câu hỏi này không tham gia vào quá trình tạo câu trả lời chính, mà được xử lý như một tác vụ phụ sau khi tin nhắn của trợ lý đã được lưu. Cách tổ chức này giúp luồng trả lời chính không bị chậm chỉ vì phải chờ sinh câu hỏi gợi ý.",
    "Nội dung đầu vào cho bước sinh câu hỏi gợi ý gồm câu hỏi ban đầu và phần tóm tắt ngắn của câu trả lời. Kết quả được giới hạn ở một số lượng nhỏ câu hỏi, tập trung vào những hướng hỏi tiếp có giá trị thực tế như điều kiện áp dụng, hồ sơ cần chuẩn bị, thời hạn xử lý, mức phạt liên quan hoặc quyền khiếu nại. Nếu tác vụ sinh gợi ý thất bại, hệ thống không làm lỗi toàn bộ luồng chat mà chỉ bỏ qua phần gợi ý. Điều này phù hợp với vai trò bổ trợ của chức năng này.",
    "Ở phía ứng dụng di động, câu hỏi gợi ý có thể xuất hiện muộn hơn câu trả lời chính. Sau khi nhận sự kiện done, ứng dụng tiếp tục kiểm tra trong một khoảng thời gian ngắn để lấy danh sách câu hỏi đã được tạo ở nền. Khi có kết quả, giao diện cập nhật lại tin nhắn tương ứng và hiển thị các gợi ý dưới câu trả lời. Nhờ vậy, người dùng vẫn nhận được câu trả lời nhanh, trong khi hệ thống có thêm thời gian tạo các gợi ý có chất lượng hơn.",
    "Nhìn chung, phần quản lý lịch sử hội thoại và câu hỏi gợi ý giúp luồng chat trở thành một trải nghiệm tư vấn liên tục thay vì các lượt hỏi đáp rời rạc. Lịch sử giúp hệ thống hiểu bối cảnh, còn câu hỏi gợi ý giúp người dùng mở rộng vấn đề theo hướng tự nhiên. Dù vậy, cả hai thành phần này đều được đặt dưới nguyên tắc an toàn: chúng hỗ trợ trải nghiệm và định hướng trao đổi, nhưng căn cứ pháp lý của câu trả lời vẫn phải đến từ pipeline truy hồi và kiểm chứng trong từng lượt xử lý.",
]


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc, prefix):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Không tìm thấy đoạn bắt đầu bằng: {prefix}")


def main():
    doc = Document(DOCX)
    idx_283 = find_paragraph(doc, "2.8.3.")
    idx_29 = find_paragraph(doc, "2.9.")

    original_heading_style = doc.paragraphs[idx_283].style
    for paragraph in list(doc.paragraphs[idx_283:idx_29]):
        remove_paragraph(paragraph)

    anchor = doc.paragraphs[find_paragraph(doc, "2.9.") - 1]
    current = anchor
    for offset, text in enumerate(reversed(NEW_SECTION)):
        style = original_heading_style if offset == len(NEW_SECTION) - 1 else "Normal"
        current = insert_paragraph_after(anchor, text, style=style)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
