from pathlib import Path
from docx import Document

docx_path = Path("Bao_Cao/Vương Văn Duy_CT060411_NGHIÊN CỨU PHÁT TRIỂN TRỢ LÝ ẢO PHÁP LUẬT CHO CHUYỂN ĐỔI SỐ.docx")
doc = Document(docx_path)

keywords = [
    "4.4.1. Thiết kế thực nghiệm",
    "4.4.2. Các chỉ số đánh giá",
    "4.4.3. Kết quả thực nghiệm",
    "Kết quả thực nghiệm cho thấy",
    "Temporal Conflict Detection",
    "khoảng cách giữa Accuracy@1",
    "Về Answer Quality",
    "Về latency",
]

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if any(k in text for k in keywords):
        print(f"\n--- HIT paragraph {i}: {text[:120]}")
        for j in range(max(0, i - 2), min(len(doc.paragraphs), i + 8)):
            print(f"{j}: {doc.paragraphs[j].text}")
