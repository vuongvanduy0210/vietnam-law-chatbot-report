from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "03_Chuong_3.docx"
DIAGRAM_DIR = ROOT / "diagrams" / "chapter3_drawio" / "png_preview"
SEQUENCE_DIR = ROOT / "diagrams" / "chapter3_sequence_general"

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"
BLACK = RGBColor(0, 0, 0)


DIAGRAMS: dict[str, str] = {
    "hinh_3_1_usecase_tong_quat": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
skinparam rectangleBorderThickness 1.4
skinparam usecaseBorderThickness 1.3
left to right direction
actor "Người dùng" as User
actor "Quản trị viên" as Admin
rectangle "Vietnam Law Chatbot" {
  usecase "Đăng ký,\nđăng nhập" as UCAuth
  usecase "Quản lý\nhội thoại" as UCConv
  usecase "Gửi câu hỏi\npháp luật" as UCChat
  usecase "Tư vấn\ncó hướng dẫn" as UCGuided
  usecase "Tra cứu thư viện\npháp luật" as UCLaw
  usecase "Tìm kiếm AI\nngữ nghĩa" as UCAI
  usecase "Quản trị\nvăn bản" as UCDoc
  usecase "Cập nhật văn bản\npháp luật" as UCUpload
  usecase "Theo dõi tiến trình\nxử lý tài liệu" as UCTrack
  usecase "Xem dashboard\nthống kê" as UCDash
}
User --> UCAuth
User --> UCConv
User --> UCChat
User --> UCGuided
User --> UCLaw
User --> UCAI
Admin --> UCAuth
Admin --> UCDoc
Admin --> UCUpload
Admin --> UCTrack
Admin --> UCDash
@enduml
""",
    "hinh_3_2_usecase_auth": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Người dùng" as User
rectangle "Nhóm xác thực và tài khoản" {
  usecase "Đăng ký tài khoản" as Signup
  usecase "Đăng nhập" as Login
  usecase "Làm mới token" as Refresh
  usecase "Đăng xuất" as Logout
  usecase "Xem/cập nhật\nthông tin cá nhân" as Profile
  usecase "Đổi mật khẩu" as Password
}
User --> Signup
User --> Login
User --> Logout
User --> Profile
User --> Password
Login ..> Refresh : <<include>>
@enduml
""",
    "hinh_3_3_usecase_conversation": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Người dùng" as User
rectangle "Quản lý hội thoại" {
  usecase "Tạo hội thoại" as NewConv
  usecase "Xem danh sách\nhội thoại" as ListConv
  usecase "Cập nhật hội thoại\n(đổi tên, ghim, lưu trữ)" as UpdateConv
  usecase "Xoá hội thoại" as DeleteConv
  usecase "Tìm kiếm\nlịch sử chat" as SearchChat
}
User --> NewConv
User --> ListConv
User --> UpdateConv
User --> DeleteConv
User --> SearchChat
@enduml
""",
    "hinh_3_4_usecase_chat_question": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Người dùng" as User
rectangle "Hỏi đáp pháp luật" {
  usecase "Nhập câu hỏi\npháp luật" as Input
  usecase "Gửi tin nhắn\nqua chat" as Send
  usecase "Nhận event ready" as Ready
  usecase "Theo dõi tiến trình\nSSE" as Progress
  usecase "Nhận câu trả lời\nvà nguồn dẫn" as Done
  usecase "Xem câu hỏi\ngợi ý" as Suggested
}
User --> Input
Input ..> Send : <<include>>
Send ..> Ready : <<include>>
Send ..> Progress : <<include>>
Send ..> Done : <<include>>
Done ..> Suggested : <<extend>>
@enduml
""",
    "hinh_3_4_usecase_guided": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Người dùng" as User
rectangle "Tư vấn có hướng dẫn" {
  usecase "Nhập câu hỏi\nban đầu" as Input
  usecase "Nhận câu hỏi\nlàm rõ" as Clarify
  usecase "Chọn đáp án\nngữ cảnh" as Select
  usecase "Nhận câu trả lời\ncó nguồn dẫn" as Answer
  usecase "Theo dõi tiến trình\nqua SSE" as Stream
}
User --> Input
Input ..> Clarify : <<include>>
User --> Select
Select ..> Answer : <<include>>
Answer ..> Stream : <<include>>
@enduml
""",
    "hinh_3_6_usecase_law_browse": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Người dùng" as User
rectangle "Tra cứu thư viện pháp luật" {
  usecase "Xem danh sách\nvăn bản" as List
  usecase "Lọc theo năm,\nchủ đề" as Filter
  usecase "Tìm kiếm theo\nnội dung" as Search
  usecase "Xem chi tiết\nđiều luật" as Detail
  usecase "Xem topic,\nnăm, keyword" as Meta
}
User --> List
User --> Filter
User --> Search
User --> Detail
List ..> Meta : <<include>>
@enduml
""",
    "hinh_3_7_usecase_ai_search": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Người dùng" as User
rectangle "Tìm kiếm AI ngữ nghĩa" {
  usecase "Nhập truy vấn\nngôn ngữ tự nhiên" as Query
  usecase "Gửi yêu cầu\nAI Search" as Send
  usecase "Truy hồi vector\ntrong ChromaDB" as Retrieve
  usecase "Xếp hạng lại\nkết quả" as Rerank
  usecase "Xem danh sách\nđiều luật liên quan" as Result
  usecase "Mở chi tiết\nđiều luật" as Detail
}
User --> Query
Query ..> Send : <<include>>
Send ..> Retrieve : <<include>>
Retrieve ..> Rerank : <<include>>
Rerank ..> Result : <<include>>
Result ..> Detail : <<extend>>
@enduml
""",
    "hinh_3_8_usecase_upload": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Quản trị viên" as Admin
rectangle "Cập nhật văn bản pháp luật" {
  usecase "Upload PDF\nvăn bản pháp luật" as Upload
  usecase "Xem trước kết quả\ntrích xuất" as Preview
  usecase "Kiểm tra trùng\nvăn bản" as Duplicate
  usecase "Validate dữ liệu\ntrích xuất" as Validate
  usecase "Lưu MongoDB\nvà ChromaDB" as Save
  usecase "Rollback khi lỗi" as Rollback
}
Admin --> Upload
Admin --> Preview
Upload ..> Duplicate : <<include>>
Upload ..> Validate : <<include>>
Validate ..> Save : <<include>>
Save ..> Rollback : <<extend>>
@enduml
""",
    "hinh_3_9_usecase_tracking": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
left to right direction
actor "Quản trị viên" as Admin
rectangle "Theo dõi xử lý tài liệu" {
  usecase "Kết nối WebSocket" as WS
  usecase "Nhận UPLOAD_PROGRESS" as Progress
  usecase "Nhận UPLOAD_STATUS" as Status
  usecase "Huỷ task\nđang chạy" as Cancel
  usecase "Xoá log task" as DeleteTask
  usecase "Xem dashboard\nthống kê" as Dash
}
Admin --> WS
WS ..> Progress : <<include>>
WS ..> Status : <<include>>
Admin --> Cancel
Admin --> DeleteTask
Admin --> Dash
@enduml
""",
    "hinh_3_7_sequence_chat": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
participant "Mobile App" as Mobile
participant "Main Service" as Main
database "PostgreSQL" as PG
participant "RAG Service" as RAG
database "ChromaDB" as Chroma
database "MongoDB" as Mongo
Mobile -> Main : POST /api/v1/chat/messages/stream
Main -> PG : tạo/lấy conversation
Main -> PG : lưu user message
Main --> Mobile : event ready
Main -> RAG : POST /api/v1/rag/agent-search/stream
RAG --> Main : event progress(validate/analyze)
Main --> Mobile : event progress
RAG -> Chroma : truy hồi vector chunks
RAG -> Mongo : đối chiếu metadata/nội dung nguồn
RAG --> Main : event progress(retrieve/synthesize/verify)
Main --> Mobile : event progress
RAG --> Main : event done(answer, sources)
Main -> PG : lưu assistant message
Main --> Mobile : event done
Mobile -> Main : GET /messages/{id}/suggested-questions
Main --> Mobile : suggested_questions
@enduml
""",
    "hinh_3_8_sequence_guided": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
participant "Mobile App" as Mobile
participant "Main Service" as Main
participant "RAG Service" as RAG
database "ChromaDB" as Chroma
Mobile -> Main : POST /api/v1/guided/clarify
Main -> RAG : POST /api/v1/rag/guided-clarify
RAG --> Main : status, topic, questions
Main --> Mobile : câu hỏi làm rõ
Mobile -> Main : POST /api/v1/guided/answer/stream
Main -> RAG : POST /api/v1/rag/guided-answer/stream
RAG --> Main : event progress(validate)
Main --> Mobile : event progress
RAG -> Chroma : truy hồi nguồn phù hợp
RAG --> Main : event progress(retrieve/synthesize/verify)
Main --> Mobile : event progress
RAG --> Main : event done(answer, sources)
Main --> Mobile : event done
@enduml
""",
    "hinh_3_9_sequence_law_search": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
participant "Mobile App" as Mobile
participant "Main Service" as Main
database "MongoDB" as Mongo
participant "RAG Service" as RAG
database "ChromaDB" as Chroma
Mobile -> Main : GET /api/v1/laws?q/year/topics
Main -> Mongo : aggregation group by law_id
Mongo --> Main : danh sách văn bản
Main --> Mobile : items, total, page
Mobile -> Main : POST /api/v1/laws/ai-search
Main -> RAG : POST /api/v1/rag/semantic-search
RAG -> Chroma : vector search + rerank
Chroma --> RAG : sources
RAG --> Main : results
Main --> Mobile : kết quả semantic search
@enduml
""",
    "hinh_3_10_sequence_upload": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
participant "Admin Web" as Admin
participant "Main Service" as Main
database "PostgreSQL" as PG
database "MongoDB" as Mongo
participant "RAG Service" as RAG
database "ChromaDB" as Chroma
Admin -> Main : POST /api/v1/documents/upload-v2
Main -> PG : tạo document_tasks
Main --> Admin : WebSocket UPLOAD_PROGRESS
Main -> Main : kiểm tra file, hash, law_id
Main -> Main : OCR + parse + validate
Main -> Mongo : lưu articles + laws_cache
Main -> RAG : POST /api/v1/ingest/articles
RAG -> Chroma : chunk, embedding, insert
Chroma --> RAG : success
RAG --> Main : ingest result
Main -> PG : status=completed, progress=100
Main --> Admin : WebSocket UPLOAD_STATUS completed
alt lỗi sau khi đã ghi dữ liệu
  Main -> Mongo : rollback articles/cache
  Main -> RAG : DELETE /api/v1/ingest/articles/{law_id}
  Main -> PG : status=failed
  Main --> Admin : WebSocket UPLOAD_STATUS failed
end
@enduml
""",
    "hinh_3_11_erd_postgresql": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
hide circle
entity "users" as users {
  * id : UUID
  --
  email : varchar
  hashed_password : varchar
  full_name : varchar
  phone_number : varchar
  role : varchar
  is_active : boolean
  last_login_at : timestamptz
}
entity "conversations" as conversations {
  * id : UUID
  --
  user_id : UUID
  title : varchar
  is_pinned : boolean
  is_archived : boolean
  message_count : integer
  last_message_at : timestamptz
}
entity "messages" as messages {
  * id : UUID
  --
  conversation_id : UUID
  question_id : UUID
  role : varchar
  content : text
  sources : JSONB
  metadata : JSONB
}
entity "refresh_tokens" as tokens {
  * id : UUID
  --
  user_id : UUID
  token : varchar
  expires_at : timestamptz
  is_revoked : boolean
}
entity "document_tasks" as tasks {
  * id : UUID
  --
  user_id : UUID
  filename : varchar
  status : enum
  progress : integer
  law_id : varchar
  article_count : integer
}
users ||--o{ conversations
users ||--o{ tokens
users ||--o{ tasks
conversations ||--o{ messages
messages ||--o| messages : question_id
@enduml
""",
    "hinh_3_12_data_stores": r"""
@startuml
skinparam monochrome true
skinparam shadowing false
skinparam defaultFontName "Arial"
rectangle "PostgreSQL" as PG {
  rectangle "users" as U
  rectangle "conversations" as C
  rectangle "messages.sources" as S
  rectangle "document_tasks.law_id" as T
}
rectangle "MongoDB" as M {
  rectangle "VietnamLawDB.articles" as A
  rectangle "laws_cache" as LC
}
rectangle "ChromaDB" as CH {
  rectangle "vietnamese_law" as V
  rectangle "chunk_id = law_id_article_id_chunkN" as CK
}
T --> A : law_id
A --> V : law_id + article_id
S --> V : source/chunk reference
LC --> A : materialized summary
@enduml
""",
}


def set_font(font, size=None, bold=None, italic=None, all_caps=None, name=BODY_FONT):
    font.name = name
    font.color.rgb = BLACK
    font._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if all_caps is not None:
        font.all_caps = all_caps


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    set_font(normal.font, 14)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h1 = doc.styles["Heading 1"]
    set_font(h1.font, 16, bold=True, all_caps=True)
    pf = h1.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing = 1.3

    h2 = doc.styles["Heading 2"]
    set_font(h2.font, 14, bold=True)
    pf = h2.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3

    h3 = doc.styles["Heading 3"]
    set_font(h3.font, 14, bold=True, italic=True)
    pf = h3.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3

    h4 = doc.styles["Heading 4"]
    set_font(h4.font, 14, italic=True)
    pf = h4.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.3

    for name, size, bold, italic, align, before, after in [
        ("Tên hình vẽ", 14, False, True, WD_ALIGN_PARAGRAPH.CENTER, 6, 12),
        ("Tên bảng", 14, True, False, WD_ALIGN_PARAGRAPH.LEFT, 6, 0),
    ]:
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        set_font(style.font, size, bold=bold, italic=italic)
        pf = style.paragraph_format
        pf.alignment = align
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.3


def set_run_font(run, size=14, bold=None, italic=None, name=BODY_FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


CODE_TOKENS = {
    "users", "conversations", "messages", "document_tasks", "refresh_tokens",
    "question_id", "sources", "metadata", "metadata_", "conversation_id",
    "user_message", "assistant_message", "suggested_questions", "ready",
    "progress", "done", "error", "timeout", "pending", "processing",
    "completed", "failed", "cancelled", "UPLOAD_PROGRESS", "UPLOAD_STATUS",
    "task_id", "law_id", "article_id", "VietnamLawDB.articles",
    "vietnamese_law", "laws_cache", "chunk_id", "retrieve_internal_law",
    "search_web_for_law", "SCORE_THRESHOLD_DISPLAY",
    "SCORE_THRESHOLD_GENERATION", "METADATA_FILTER_LIMIT",
    "/api/v1/auth/signup", "/api/v1/auth/login", "/api/v1/auth/refresh",
    "/api/v1/chat/messages/stream", "/api/v1/chat/conversations",
    "/api/v1/chat/messages/{message_id}/suggested-questions",
    "/api/v1/guided/clarify", "/api/v1/guided/answer/stream",
    "/api/v1/laws", "/api/v1/laws/ai-search", "/api/v1/laws/detail",
    "/api/v1/documents/upload-v2", "/api/v1/documents/ws",
    "/api/v1/documents/tasks", "/api/v1/ingest/articles",
    "/api/v1/rag/agent-search/stream", "/api/v1/rag/guided-answer/stream",
    "/api/v1/rag/semantic-search", "X-API-Key",
}


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if style is None or style == "Normal":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(0)
    if text:
        add_runs(p, text)
    return p


def add_runs(paragraph, text: str, size=14):
    if not text:
        return
    # Longest tokens first so endpoint paths are not split by smaller fragments.
    tokens = sorted(CODE_TOKENS, key=len, reverse=True)
    i = 0
    while i < len(text):
        matched = None
        for token in tokens:
            if text.startswith(token, i):
                matched = token
                break
        if matched:
            run = paragraph.add_run(matched)
            set_run_font(run, size=11, name=CODE_FONT)
            i += len(matched)
        else:
            j = i + 1
            while j < len(text) and not any(text.startswith(tok, j) for tok in tokens):
                j += 1
            run = paragraph.add_run(text[i:j])
            set_run_font(run, size=size, name=BODY_FONT)
            i = j


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, size=16 if level == 1 else 14, bold=True if level <= 3 else None, italic=True if level == 3 else None)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    cap = add_paragraph(doc, caption, style="Tên bảng")
    cap.paragraph_format.left_indent = Cm(1.0)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        widths = [16.0 / len(headers)] * len(headers)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(widths[i])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        add_runs(p, h, size=12)
        for run in p.runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Cm(widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, value, size=12)
    add_paragraph(doc, "")
    return table


def add_usecase_table(doc: Document, caption: str, fields: list[tuple[str, str]]):
    cap = add_paragraph(doc, caption, style="Tên bảng")
    cap.paragraph_format.left_indent = Cm(1.0)
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [4.2, 11.8]

    for row_idx, (label, value) in enumerate(fields):
        cells = table.rows[row_idx].cells
        for i, text in enumerate([label, value]):
            cell = cells[i]
            cell.width = Cm(widths[i])
            set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, text, size=12)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    add_paragraph(doc, "")
    return table


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap = add_paragraph(doc, caption, style="Tên hình vẽ")
    cap.paragraph_format.first_line_indent = Cm(0)


def render_diagrams():
    key_to_file = {
        "hinh_3_1_usecase_tong_quat": "hinh_3_1.png",
        "hinh_3_2_usecase_auth": "hinh_3_2.png",
        "hinh_3_3_usecase_conversation": "hinh_3_3.png",
        "hinh_3_4_usecase_chat_question": "hinh_3_4.png",
        "hinh_3_4_usecase_guided": "hinh_3_5.png",
        "hinh_3_6_usecase_law_browse": "hinh_3_6.png",
        "hinh_3_7_usecase_ai_search": "hinh_3_7.png",
        "hinh_3_8_usecase_upload": "hinh_3_8.png",
        "hinh_3_9_usecase_tracking": "hinh_3_9.png",
        "hinh_3_11_erd_postgresql": "hinh_3_14.png",
    }
    rendered = {key: DIAGRAM_DIR / filename for key, filename in key_to_file.items()}
    sequence_files = {
        "seq_auth_profile": "hinh_3_10_sequence_xac_thuc_tai_khoan.png",
        "seq_conversation": "hinh_3_11_sequence_quan_ly_hoi_thoai.png",
        "seq_chat_answer": "hinh_3_12_sequence_hoi_dap_phap_luat.png",
        "seq_guided": "hinh_3_13_sequence_tu_van_co_huong_dan.png",
        "seq_law_library": "hinh_3_14_sequence_thu_vien_phap_luat.png",
        "seq_semantic_search": "hinh_3_15_sequence_tim_kiem_ngu_nghia.png",
        "seq_dashboard_tracker": "hinh_3_16_sequence_dashboard_tracker.png",
        "seq_upload_processing": "hinh_3_17_sequence_upload_xu_ly_van_ban.png",
        "seq_admin_document": "hinh_3_18_sequence_quan_ly_van_ban_admin.png",
    }
    rendered.update({key: SEQUENCE_DIR / filename for key, filename in sequence_files.items()})
    missing = [str(path) for path in rendered.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Thiếu ảnh draw.io đã export: " + ", ".join(missing))
    return rendered


def add_intro(doc):
    heading(doc, "CHƯƠNG 3: PHÂN TÍCH, THIẾT KẾ HỆ THỐNG", 1)
    add_paragraph(doc, "Trên cơ sở phương pháp xây dựng hệ thống đã trình bày ở chương 2, chương này đi vào phân tích yêu cầu và thiết kế hệ thống Vietnam Law Chatbot ở mức chức năng. Mục tiêu của chương là làm rõ hệ thống cần đáp ứng những yêu cầu nào, các nhóm tác nhân tương tác với hệ thống ra sao, dữ liệu được tổ chức như thế nào và các thành phần trao đổi thông tin qua những API chính nào.")
    add_paragraph(doc, "Nội dung chương được triển khai theo tiến trình phân tích và thiết kế hệ thống phần mềm. Trước hết, chương xác định yêu cầu về hình thức sản phẩm, yêu cầu chức năng và yêu cầu phi chức năng. Tiếp theo, các ca sử dụng trọng tâm được đặc tả và minh hoạ bằng biểu đồ. Sau đó, chương trình bày các luồng tuần tự đại diện cho quá trình xử lý nghiệp vụ, cuối cùng là thiết kế cơ sở dữ liệu và API làm nền tảng cho phần xây dựng, thực nghiệm ở chương sau.")


def section_requirements(doc):
    heading(doc, "3.1. Yêu cầu của hệ thống", 2)
    heading(doc, "3.1.1. Yêu cầu về hình thức sản phẩm", 3)
    add_paragraph(doc, "Hệ thống được thiết kế theo mô hình nhiều thành phần, trong đó mỗi thành phần đảm nhận một vai trò riêng trong quá trình phục vụ người dùng và quản trị dữ liệu. Việc tách ứng dụng di động, trang quản trị và các dịch vụ backend giúp hệ thống dễ bảo trì hơn, đồng thời cho phép phần xử lý nghiệp vụ và phần xử lý tri thức pháp luật được mở rộng độc lập khi khối lượng dữ liệu hoặc số lượng người dùng tăng lên.")
    add_table(doc, "Bảng 3.1. Yêu cầu về hình thức sản phẩm", ["Thành phần", "Nền tảng", "Vai trò"], [
        ["Mobile App", "Android và iOS", "Cung cấp chức năng hỏi đáp, tư vấn có hướng dẫn, tra cứu thư viện pháp luật và quản lý hội thoại cho người dùng cuối."],
        ["Admin Web", "Trình duyệt web", "Cung cấp giao diện quản trị dữ liệu pháp luật, tải lên văn bản, theo dõi tiến trình xử lý và xem thống kê hệ thống."],
        ["Main Service", "Backend service", "Xác thực người dùng, quản lý hội thoại, điều phối yêu cầu hỏi đáp, quản lý văn bản và kết nối với các kho dữ liệu."],
        ["RAG Service", "Backend service nội bộ", "Thực hiện truy hồi ngữ nghĩa, xử lý Agentic RAG, tư vấn có hướng dẫn, tạo embedding và lưu vector vào ChromaDB."],
    ], [3.0, 3.0, 10.0])

    heading(doc, "3.1.2. Yêu cầu chức năng", 3)
    heading(doc, "3.1.2.1. Yêu cầu chức năng của hệ thống đối với người dùng", 4)
    add_paragraph(doc, "Đối với người dùng cuối, các yêu cầu chức năng được xây dựng xoay quanh quá trình đặt câu hỏi pháp luật, theo dõi lịch sử hội thoại và tra cứu lại căn cứ pháp lý khi cần. Ngoài chức năng hỏi đáp trực tiếp, hệ thống còn cần hỗ trợ hình thức tư vấn có hướng dẫn để xử lý các tình huống người dùng chưa cung cấp đủ ngữ cảnh.")
    add_table(doc, "Bảng 3.2. Yêu cầu chức năng đối với người dùng", ["Mã", "Chức năng", "Mô tả"], [
        ["F-01", "Đăng ký, đăng nhập và quản lý tài khoản", "Người dùng tạo tài khoản, đăng nhập, làm mới phiên đăng nhập, cập nhật thông tin cá nhân và đổi mật khẩu."],
        ["F-02", "Quản lý hội thoại", "Người dùng xem danh sách hội thoại, mở chi tiết hội thoại, đổi tên, ghim, lưu trữ hoặc xoá hội thoại."],
        ["F-03", "Gửi câu hỏi pháp luật", "Người dùng gửi câu hỏi tự nhiên và nhận câu trả lời từ hệ thống theo cơ chế trả tiến trình xử lý."],
        ["F-04", "Xem câu hỏi gợi ý", "Sau khi câu trả lời hoàn tất, người dùng có thể xem các câu hỏi tiếp theo được gợi ý dựa trên nội dung trả lời."],
        ["F-05", "Tư vấn có hướng dẫn", "Hệ thống hỏi thêm các thông tin làm rõ trước khi trả lời trong các trường hợp câu hỏi ban đầu còn thiếu ngữ cảnh."],
        ["F-06", "Tra cứu thư viện pháp luật", "Người dùng xem danh sách văn bản, lọc theo năm, chủ đề, từ khoá và xem chi tiết điều luật."],
        ["F-07", "Tìm kiếm ngữ nghĩa", "Người dùng nhập truy vấn tự nhiên, xem danh sách điều luật liên quan, mở nguồn dẫn hoặc điều chỉnh từ khoá để tìm lại."],
    ], [1.6, 4.1, 10.3])

    heading(doc, "3.1.2.2. Yêu cầu chức năng của hệ thống đối với quản trị viên", 4)
    add_paragraph(doc, "Đối với quản trị viên, trọng tâm của hệ thống là duy trì kho dữ liệu pháp luật luôn sẵn sàng cho các chức năng hỏi đáp và tra cứu. Các thao tác quản trị không chỉ dừng ở việc tải lên văn bản, mà còn bao gồm theo dõi tiến trình xử lý, kiểm tra kết quả cập nhật, quản lý các tác vụ phát sinh và quan sát tình trạng hoạt động chung của hệ thống.")
    add_table(doc, "Bảng 3.3. Yêu cầu chức năng đối với quản trị viên", ["Mã", "Chức năng", "Mô tả"], [
        ["F-08", "Đăng nhập quản trị", "Quản trị viên đăng nhập bằng tài khoản có quyền quản trị viên để truy cập trang tổng quan."],
        ["F-09", "Tải lên văn bản pháp luật", "Quản trị viên chọn tệp PDF, tải lên hệ thống, theo dõi kết quả xử lý và có thể huỷ yêu cầu khi quá trình chưa hoàn tất."],
        ["F-10", "Theo dõi tiến trình xử lý", "Quản trị viên xem danh sách tác vụ xử lý tài liệu, theo dõi tiến trình, xem trạng thái thành công/thất bại, huỷ tác vụ đang chạy và xem thống kê tổng quan."],
        ["F-11", "Quản lý văn bản pháp luật", "Quản trị viên xem danh sách văn bản, tìm kiếm, xem chi tiết, làm mới dữ liệu hiển thị hoặc xoá văn bản khỏi hệ thống."],
        ["F-12", "Quản lý tác vụ xử lý tài liệu", "Quản trị viên xem lịch sử tác vụ, huỷ tác vụ đang chạy hoặc xoá nhật ký tác vụ không còn cần thiết."],
        ["F-13", "Xem thống kê hệ thống", "Quản trị viên theo dõi số lượng người dùng, hội thoại, tin nhắn, văn bản và trạng thái các tác vụ gần đây."],
    ], [1.6, 4.2, 10.2])

    heading(doc, "3.1.3. Yêu cầu phi chức năng", 3)
    add_paragraph(doc, "Bên cạnh các yêu cầu chức năng, hệ thống cần đáp ứng một số yêu cầu phi chức năng để đảm bảo có thể sử dụng ổn định trong bối cảnh tư vấn pháp luật. Các yêu cầu này tập trung vào bảo mật, hiệu năng, tính nhất quán dữ liệu và khả năng mở rộng.")
    add_table(doc, "Bảng 3.4. Yêu cầu phi chức năng", ["Nhóm", "Yêu cầu", "Mức độ"], [
        ["Bảo mật", "Các API nghiệp vụ yêu cầu xác thực bằng access token; refresh token được lưu và có thể thu hồi.", "Bắt buộc"],
        ["Phân quyền", "Các chức năng tải lên văn bản, xoá văn bản, làm mới dữ liệu hiển thị và xem tác vụ tài liệu chỉ dành cho quản trị viên.", "Bắt buộc"],
        ["Hiệu năng", "Luồng hỏi đáp và tư vấn có hướng dẫn phải trả trạng thái xử lý sớm để người dùng nhận biết hệ thống đang thực hiện yêu cầu.", "Bắt buộc"],
        ["Tính nhất quán", "Khi cập nhật văn bản mới, MongoDB, ChromaDB và dữ liệu hiển thị phải được cập nhật đồng bộ; nếu lỗi cần hoàn tác dữ liệu đã ghi.", "Bắt buộc"],
        ["Khả năng mở rộng", "Main Service và RAG Service tách riêng để có thể mở rộng phần xử lý AI độc lập với phần nghiệp vụ.", "Khuyến nghị"],
        ["Trải nghiệm", "Giao diện hỏi đáp cần hiển thị tiến trình xử lý, nguồn tham chiếu và câu hỏi gợi ý sau câu trả lời.", "Khuyến nghị"],
    ], [3.0, 10.0, 3.0])


def section_analysis(doc, diagrams):
    heading(doc, "3.2. Phân tích hệ thống", 2)
    heading(doc, "3.2.1. Biểu đồ ca sử dụng tổng quát", 3)
    add_paragraph(doc, "Biểu đồ ca sử dụng tổng quát thể hiện phạm vi chức năng của hệ thống dưới góc nhìn người sử dụng. Hệ thống có hai nhóm tác nhân chính: người dùng cuối và quản trị viên. Người dùng cuối sử dụng các chức năng phục vụ nhu cầu hỏi đáp, tư vấn và tra cứu pháp luật; quản trị viên chịu trách nhiệm vận hành dữ liệu, cập nhật văn bản và theo dõi trạng thái xử lý tài liệu.")
    add_figure(doc, diagrams["hinh_3_1_usecase_tong_quat"], "Hình 3.1. Biểu đồ ca sử dụng tổng quát của hệ thống Vietnam Law Chatbot")

    heading(doc, "3.2.2. Đặc tả các ca sử dụng", 3)
    add_paragraph(doc, "Trong phạm vi đồ án, các ca sử dụng được đặc tả theo nhóm chức năng chính thay vì tách riêng toàn bộ thao tác con thành nhiều bảng nhỏ. Cách trình bày này phù hợp với hệ thống có nhiều thao tác liên quan chặt chẽ với nhau, giúp phần báo cáo giữ được mạch đọc liền mạch nhưng vẫn thể hiện đầy đủ mục tiêu sử dụng, tác nhân, luồng chính, luồng thay thế và kết quả của từng nhóm chức năng. Những thao tác chi tiết bên trong từng nhóm tiếp tục được làm rõ ở phần biểu đồ tuần tự.")
    usecases = [
        ("3.2.2.1. Ca sử dụng đăng ký, đăng nhập và quản lý tài khoản", "hinh_3_2_usecase_auth", "Hình 3.2. Ca sử dụng đăng ký, đăng nhập và quản lý tài khoản", "Bảng 3.5. Đặc tả ca sử dụng đăng ký, đăng nhập và quản lý tài khoản", [
            ("Tên ca sử dụng", "UC – Đăng ký, đăng nhập và quản lý tài khoản"),
            ("Tác nhân", "Người dùng"),
            ("Mô tả", "Ca sử dụng mô tả quá trình người dùng tạo tài khoản, đăng nhập vào ứng dụng, duy trì phiên sử dụng và cập nhật các thông tin cá nhân cơ bản trong hệ thống."),
            ("Luồng chính", "1. Người dùng mở ứng dụng và chọn chức năng đăng ký hoặc đăng nhập.\n2. Người dùng nhập email, mật khẩu và các thông tin cần thiết.\n3. Giao diện kiểm tra định dạng dữ liệu trước khi gửi yêu cầu đến hệ thống.\n4. Hệ thống kiểm tra tài khoản trong CSDL, xác thực mật khẩu hoặc tạo tài khoản mới.\n5. Nếu thông tin hợp lệ, hệ thống trả về phiên đăng nhập và thông tin người dùng.\n6. Ứng dụng lưu phiên đăng nhập, chuyển người dùng vào màn hình chính.\n7. Khi cần, người dùng có thể xem hồ sơ, cập nhật thông tin cá nhân hoặc đổi mật khẩu."),
            ("Luồng thay thế", "Nếu email đã tồn tại, mật khẩu không đúng, token hết hạn hoặc dữ liệu cập nhật không hợp lệ, hệ thống trả thông báo lỗi để người dùng nhập lại."),
            ("Kết quả", "Người dùng đăng nhập thành công hoặc thông tin tài khoản được cập nhật đúng trong hệ thống."),
            ("Yêu cầu đặc biệt", "Mật khẩu phải được kiểm tra và lưu trữ an toàn; các chức năng cập nhật tài khoản yêu cầu người dùng đã xác thực."),
        ]),
        ("3.2.2.2. Ca sử dụng quản lý danh sách cuộc trò chuyện", "hinh_3_3_usecase_conversation", "Hình 3.3. Ca sử dụng quản lý danh sách cuộc trò chuyện", "Bảng 3.6. Đặc tả ca sử dụng quản lý danh sách cuộc trò chuyện", [
            ("Tên ca sử dụng", "UC – Quản lý danh sách cuộc trò chuyện"),
            ("Tác nhân", "Người dùng"),
            ("Mô tả", "Ca sử dụng mô tả quá trình người dùng xem, mở và quản lý các cuộc trò chuyện đã tạo trong quá trình sử dụng chức năng hỏi đáp pháp luật."),
            ("Luồng chính", "1. Người dùng truy cập màn hình danh sách hội thoại.\n2. Hệ thống lấy danh sách hội thoại của người dùng, kèm tin nhắn gần nhất và trạng thái ghim/lưu trữ.\n3. Giao diện hiển thị danh sách hội thoại theo thứ tự thời gian hoặc theo trạng thái ưu tiên.\n4. Người dùng chọn một hội thoại để xem lại nội dung.\n5. Hệ thống trả về chi tiết hội thoại và danh sách tin nhắn tương ứng.\n6. Người dùng có thể đổi tên, ghim, lưu trữ hoặc xoá hội thoại.\n7. Hệ thống cập nhật trạng thái hội thoại và làm mới danh sách hiển thị."),
            ("Luồng thay thế", "Nếu hội thoại không tồn tại, đã bị xoá hoặc không thuộc người dùng hiện tại, hệ thống thông báo lỗi và không cho phép thao tác."),
            ("Kết quả", "Danh sách hội thoại và nội dung hội thoại được hiển thị đúng theo dữ liệu của người dùng."),
            ("Yêu cầu đặc biệt", "Mỗi người dùng chỉ được xem và thao tác với các hội thoại thuộc tài khoản của mình."),
        ]),
        ("3.2.2.3. Ca sử dụng gửi câu hỏi pháp luật", "hinh_3_4_usecase_chat_question", "Hình 3.4. Ca sử dụng gửi câu hỏi pháp luật", "Bảng 3.7. Đặc tả ca sử dụng gửi câu hỏi pháp luật", [
            ("Tên ca sử dụng", "UC – Gửi câu hỏi pháp luật"),
            ("Tác nhân", "Người dùng"),
            ("Mô tả", "Ca sử dụng mô tả quá trình người dùng gửi câu hỏi pháp luật bằng ngôn ngữ tự nhiên và nhận câu trả lời có kèm nguồn tham chiếu từ hệ thống."),
            ("Luồng chính", "1. Người dùng mở màn hình hội thoại và nhập câu hỏi pháp luật.\n2. Giao diện kiểm tra nội dung nhập và gửi yêu cầu đến hệ thống.\n3. Hệ thống tạo mới hoặc sử dụng hội thoại hiện có, lưu câu hỏi của người dùng.\n4. Hệ thống truy hồi dữ liệu pháp luật liên quan và xử lý câu trả lời.\n5. Trong quá trình xử lý, giao diện hiển thị tiến trình để người dùng biết yêu cầu đang được thực hiện.\n6. Hệ thống trả về câu trả lời, nguồn tham chiếu và lưu lại tin nhắn phản hồi.\n7. Giao diện hiển thị câu trả lời và có thể hiển thị thêm các câu hỏi gợi ý liên quan."),
            ("Luồng thay thế", "Nếu câu hỏi rỗng, hội thoại không hợp lệ hoặc hệ thống không xử lý được yêu cầu, giao diện hiển thị thông báo lỗi phù hợp."),
            ("Kết quả", "Người dùng nhận được câu trả lời pháp luật, nguồn tham chiếu và lịch sử hội thoại được cập nhật."),
            ("Yêu cầu đặc biệt", "Câu trả lời cần ưu tiên căn cứ từ dữ liệu pháp luật đã lưu và phải hiển thị nguồn để người dùng kiểm chứng."),
        ]),
        ("3.2.2.4. Ca sử dụng tư vấn có hướng dẫn", "hinh_3_4_usecase_guided", "Hình 3.5. Ca sử dụng tư vấn có hướng dẫn", "Bảng 3.8. Đặc tả ca sử dụng tư vấn có hướng dẫn", [
            ("Tên ca sử dụng", "UC – Tư vấn có hướng dẫn"),
            ("Tác nhân", "Người dùng"),
            ("Mô tả", "Ca sử dụng mô tả quá trình hệ thống hỗ trợ người dùng làm rõ tình huống pháp lý trước khi đưa ra nội dung tư vấn cuối cùng."),
            ("Luồng chính", "1. Người dùng mở chức năng tư vấn có hướng dẫn và nhập vấn đề cần tư vấn.\n2. Hệ thống phân tích câu hỏi ban đầu để xác định chủ đề và mức độ thiếu thông tin.\n3. Nếu cần thêm ngữ cảnh, hệ thống tạo danh sách câu hỏi làm rõ.\n4. Người dùng chọn phương án trả lời hoặc bổ sung thêm thông tin.\n5. Hệ thống sử dụng câu hỏi ban đầu và các thông tin làm rõ để truy hồi căn cứ pháp luật phù hợp.\n6. Hệ thống tạo câu trả lời tư vấn cuối cùng, kèm nguồn tham chiếu.\n7. Giao diện hiển thị kết quả tư vấn cho người dùng."),
            ("Luồng thay thế", "Nếu vấn đề không thuộc phạm vi pháp luật, thiếu thông tin nghiêm trọng hoặc phát sinh lỗi xử lý, hệ thống thông báo lý do và yêu cầu người dùng điều chỉnh thông tin."),
            ("Kết quả", "Người dùng nhận được câu trả lời tư vấn cụ thể hơn so với hỏi đáp trực tiếp."),
            ("Yêu cầu đặc biệt", "Các câu hỏi làm rõ cần ngắn gọn, dễ chọn và liên quan trực tiếp đến tình huống pháp lý người dùng đưa ra."),
        ]),
        ("3.2.2.5. Ca sử dụng tra cứu thư viện pháp luật", "hinh_3_6_usecase_law_browse", "Hình 3.6. Ca sử dụng tra cứu thư viện pháp luật", "Bảng 3.9. Đặc tả ca sử dụng tra cứu thư viện pháp luật", [
            ("Tên ca sử dụng", "UC – Tra cứu thư viện pháp luật"),
            ("Tác nhân", "Người dùng"),
            ("Mô tả", "Ca sử dụng mô tả quá trình người dùng xem danh sách văn bản pháp luật, lọc dữ liệu và mở chi tiết văn bản hoặc điều luật trong thư viện."),
            ("Luồng chính", "1. Người dùng mở màn hình thư viện pháp luật.\n2. Hệ thống trả về danh sách văn bản đã được lưu trong CSDL.\n3. Người dùng nhập từ khoá hoặc chọn bộ lọc theo năm, chủ đề, mã văn bản.\n4. Hệ thống truy vấn dữ liệu phù hợp với điều kiện lọc và trả danh sách kết quả.\n5. Người dùng chọn một văn bản hoặc điều luật muốn xem.\n6. Hệ thống trả nội dung chi tiết của văn bản hoặc điều luật.\n7. Giao diện hiển thị nội dung để người dùng đọc và đối chiếu."),
            ("Luồng thay thế", "Nếu không tìm thấy dữ liệu phù hợp, hệ thống hiển thị trạng thái rỗng và cho phép người dùng thay đổi điều kiện tìm kiếm."),
            ("Kết quả", "Người dùng tra cứu được văn bản hoặc điều luật cần xem trong thư viện pháp luật."),
            ("Yêu cầu đặc biệt", "Dữ liệu hiển thị cần giữ nguyên nội dung pháp luật và phân tách rõ văn bản, điều luật, chủ đề, năm ban hành."),
        ]),
        ("3.2.2.6. Ca sử dụng tìm kiếm ngữ nghĩa", "hinh_3_7_usecase_ai_search", "Hình 3.7. Ca sử dụng tìm kiếm ngữ nghĩa", "Bảng 3.10. Đặc tả ca sử dụng tìm kiếm ngữ nghĩa", [
            ("Tên ca sử dụng", "UC – Tìm kiếm ngữ nghĩa văn bản pháp luật"),
            ("Tác nhân", "Người dùng"),
            ("Mô tả", "Ca sử dụng mô tả quá trình người dùng tìm điều luật liên quan bằng truy vấn tự nhiên thay vì chỉ tìm theo từ khoá chính xác."),
            ("Luồng chính", "1. Người dùng mở thư viện pháp luật và bật chế độ tìm kiếm ngữ nghĩa.\n2. Người dùng nhập truy vấn bằng ngôn ngữ tự nhiên.\n3. Hệ thống phân tích truy vấn, truy hồi dữ liệu từ CSDL Vector và xếp hạng kết quả.\n4. Giao diện hiển thị danh sách điều luật liên quan cùng điểm phù hợp hoặc thông tin nguồn.\n5. Người dùng chọn một kết quả để xem chi tiết.\n6. Hệ thống trả nội dung điều luật tương ứng và giao diện hiển thị cho người dùng."),
            ("Luồng thay thế", "Nếu truy vấn quá ngắn, không có kết quả phù hợp hoặc hệ thống truy hồi thất bại, giao diện hiển thị thông báo và gợi ý người dùng thay đổi truy vấn."),
            ("Kết quả", "Người dùng nhận được danh sách điều luật liên quan đến nội dung truy vấn tự nhiên."),
            ("Yêu cầu đặc biệt", "Kết quả tìm kiếm cần ưu tiên điều luật có độ liên quan cao và cho phép người dùng mở nguồn chi tiết để kiểm chứng."),
        ]),
        ("3.2.2.7. Ca sử dụng cập nhật văn bản pháp luật", "hinh_3_8_usecase_upload", "Hình 3.8. Ca sử dụng cập nhật văn bản pháp luật", "Bảng 3.11. Đặc tả ca sử dụng cập nhật văn bản pháp luật", [
            ("Tên ca sử dụng", "UC – Cập nhật văn bản pháp luật"),
            ("Tác nhân", "Quản trị viên"),
            ("Mô tả", "Ca sử dụng mô tả quá trình quản trị viên tải lên tệp PDF văn bản pháp luật để hệ thống trích xuất, kiểm tra và cập nhật vào kho dữ liệu."),
            ("Luồng chính", "1. Quản trị viên đăng nhập vào trang quản trị.\n2. Quản trị viên mở chức năng cập nhật văn bản pháp luật và chọn tệp PDF.\n3. Giao diện kiểm tra định dạng, kích thước tệp và gửi yêu cầu tải lên.\n4. Hệ thống tạo tác vụ xử lý và thông báo trạng thái tiếp nhận.\n5. Hệ thống kiểm tra trùng lặp, trích xuất nội dung, chuẩn hoá thông tin văn bản và điều luật.\n6. Hệ thống lưu dữ liệu vào CSDL văn bản, tạo embedding và cập nhật CSDL Vector.\n7. Khi hoàn tất, hệ thống cập nhật trạng thái tác vụ và giao diện hiển thị kết quả xử lý."),
            ("Luồng thay thế", "Nếu tệp không hợp lệ, văn bản đã tồn tại, quá trình trích xuất lỗi hoặc lưu dữ liệu thất bại, hệ thống cập nhật trạng thái lỗi và thông báo nguyên nhân cho quản trị viên. Nếu quản trị viên hủy tác vụ, hệ thống chuyển tác vụ sang trạng thái đã hủy."),
            ("Kết quả", "Văn bản pháp luật mới được cập nhật vào hệ thống hoặc tác vụ xử lý được ghi nhận với trạng thái lỗi/hủy rõ ràng."),
            ("Yêu cầu đặc biệt", "Quá trình cập nhật phải bảo đảm đồng bộ giữa CSDL văn bản, dữ liệu hiển thị và CSDL Vector; khi lỗi cần tránh để dữ liệu bị lệch giữa các kho."),
        ]),
        ("3.2.2.8. Ca sử dụng theo dõi tiến trình xử lý tài liệu", "hinh_3_9_usecase_tracking", "Hình 3.9. Ca sử dụng theo dõi tiến trình xử lý tài liệu", "Bảng 3.12. Đặc tả ca sử dụng theo dõi tiến trình xử lý tài liệu", [
            ("Tên ca sử dụng", "UC – Theo dõi tiến trình xử lý tài liệu"),
            ("Tác nhân", "Quản trị viên"),
            ("Mô tả", "Ca sử dụng mô tả quá trình quản trị viên theo dõi trạng thái các tác vụ xử lý văn bản pháp luật trên trang quản trị."),
            ("Luồng chính", "1. Quản trị viên đăng nhập vào trang quản trị.\n2. Quản trị viên mở trang thống kê hoặc trang theo dõi tiến trình xử lý tài liệu.\n3. Hệ thống trả về danh sách tác vụ gần đây, trạng thái, tiến độ, tên tệp, mã văn bản và số điều luật nếu có.\n4. Giao diện thiết lập kênh nhận cập nhật theo thời gian thực cho các tác vụ đang chạy.\n5. Khi tác vụ thay đổi trạng thái, hệ thống gửi cập nhật để giao diện làm mới tiến trình.\n6. Quản trị viên có thể mở tác vụ để xem lỗi, quay lại trang tải lên, hủy tác vụ đang chạy hoặc xoá nhật ký tác vụ không cần thiết."),
            ("Luồng thay thế", "Nếu không có tác vụ nào, giao diện hiển thị trạng thái rỗng. Nếu tác vụ đã hoàn tất hoặc thất bại, hệ thống không cho phép hủy mà chỉ cho phép xem hoặc xoá nhật ký."),
            ("Kết quả", "Quản trị viên nắm được tình trạng xử lý tài liệu và có thể can thiệp kịp thời khi tác vụ lỗi hoặc chạy quá lâu."),
            ("Yêu cầu đặc biệt", "Các cập nhật theo thời gian thực cần phản ánh đúng trạng thái tác vụ, tránh hiển thị sai tiến trình khi có nhiều tab hoặc nhiều tác vụ cùng chạy."),
        ]),
    ]
    intro_by_figure = {
        "hinh_3_2_usecase_auth": "Nhóm ca sử dụng này tập trung vào các thao tác mở đầu phiên làm việc. Việc gom đăng ký, đăng nhập, làm mới phiên và cập nhật tài khoản vào cùng một đặc tả giúp thể hiện rõ vòng đời xác thực của người dùng mà không tách rời các bước có quan hệ phụ thuộc.",
        "hinh_3_3_usecase_conversation": "Quản lý hội thoại được đặt thành một nhóm riêng vì đây là dữ liệu gắn trực tiếp với trải nghiệm hỏi đáp của người dùng. Các thao tác xem danh sách, mở chi tiết, ghim, lưu trữ và xoá đều nhằm giúp người dùng kiểm soát lịch sử trao đổi của mình trong hệ thống.",
        "hinh_3_4_usecase_chat_question": "Ca sử dụng gửi câu hỏi pháp luật là nghiệp vụ trung tâm của ứng dụng di động. Phần đặc tả nhấn mạnh thao tác của người dùng và kết quả nhận được, còn các bước truy hồi, sinh câu trả lời và lưu nguồn tham chiếu được trình bày chi tiết hơn ở nhóm biểu đồ tuần tự.",
        "hinh_3_4_usecase_guided": "Tư vấn có hướng dẫn được tách khỏi hỏi đáp thông thường vì luồng sử dụng có thêm bước làm rõ thông tin. Cách tổ chức này phản ánh những tình huống pháp lý cần biết thêm ngữ cảnh trước khi hệ thống có thể đưa ra câu trả lời phù hợp.",
        "hinh_3_6_usecase_law_browse": "Tra cứu thư viện pháp luật phục vụ nhu cầu đọc và đối chiếu văn bản theo cấu trúc có sẵn. Ca sử dụng này khác với hỏi đáp ở chỗ người dùng chủ động lọc, mở danh sách văn bản và xem nội dung điều luật thay vì yêu cầu hệ thống tổng hợp câu trả lời.",
        "hinh_3_7_usecase_ai_search": "Tìm kiếm ngữ nghĩa mở rộng chức năng thư viện bằng cách cho phép nhập truy vấn tự nhiên. Đặc tả tập trung vào tương tác người dùng với danh sách kết quả, trong khi cơ chế embedding và truy hồi vector được trình bày ở phần thiết kế và biểu đồ tuần tự.",
        "hinh_3_8_usecase_upload": "Cập nhật văn bản pháp luật là nghiệp vụ quản trị làm thay đổi trực tiếp kho tri thức của hệ thống. Vì vậy, phần đặc tả nêu rõ các điều kiện kiểm tra, kết quả xử lý và yêu cầu đồng bộ dữ liệu sau khi quản trị viên tải lên tệp PDF.",
        "hinh_3_9_usecase_tracking": "Theo dõi tiến trình xử lý tài liệu hỗ trợ quản trị viên quan sát các tác vụ đang chạy và xử lý các tình huống lỗi. Nhóm ca sử dụng này bổ sung cho luồng cập nhật văn bản bằng cách mô tả phần giám sát, hủy tác vụ và quản lý lịch sử xử lý.",
    }

    for title, fig_key, fig_caption, table_caption, fields in usecases:
        heading(doc, title, 4)
        add_paragraph(doc, intro_by_figure[fig_key])
        add_figure(doc, diagrams[fig_key], fig_caption)
        add_usecase_table(doc, table_caption, fields)

    heading(doc, "3.2.3. Biểu đồ tuần tự", 3)
    add_paragraph(doc, "Biểu đồ tuần tự được sử dụng để mô tả thứ tự trao đổi giữa tác nhân, giao diện, hệ thống và lớp dữ liệu trong các nghiệp vụ chính. Ở chương này, biểu đồ không trình bày chi tiết đến mức tên lớp hoặc tên tệp mã nguồn, mà tập trung vào trình tự nghiệp vụ: tác nhân thực hiện thao tác gì, giao diện chuyển yêu cầu như thế nào, hệ thống truy cập dữ liệu ở thời điểm nào và kết quả được trả lại theo trình tự ra sao.")
    add_paragraph(doc, "Các luồng được lựa chọn bao phủ phạm vi sản phẩm hiện tại, gồm xác thực và tài khoản, quản lý hội thoại, hỏi đáp pháp luật, tư vấn có hướng dẫn, tra cứu thư viện pháp luật, tìm kiếm ngữ nghĩa, thống kê quản trị, cập nhật văn bản và quản lý văn bản pháp luật. Ở những luồng có đọc/ghi dữ liệu quan trọng, biểu đồ bổ sung thêm lifeline CSDL hoặc CSDL Vector để làm rõ quan hệ giữa thao tác nghiệp vụ và dữ liệu nền. Cách tách này giúp phần phân tích bám sát ca sử dụng đã trình bày, đồng thời vẫn đủ tổng quát để không phụ thuộc vào chi tiết triển khai cụ thể.")
    seqs = [
        ("3.2.3.1. Biểu đồ tuần tự luồng xác thực và quản lý tài khoản", "seq_auth_profile", "Hình 3.10. Biểu đồ tuần tự luồng xác thực và quản lý tài khoản", "Luồng xác thực bắt đầu khi người dùng nhập thông tin đăng ký hoặc đăng nhập. Giao diện chịu trách nhiệm kiểm tra định dạng dữ liệu ban đầu trước khi gửi yêu cầu đến hệ thống. Hệ thống truy vấn CSDL để kiểm tra tài khoản hoặc tạo tài khoản mới, sau đó trả về thông tin phiên đăng nhập cho giao diện. Với thao tác cập nhật hồ sơ hoặc đổi mật khẩu, dữ liệu người dùng tiếp tục được ghi lại vào CSDL, bảo đảm thông tin hiển thị ở các lần truy cập sau luôn nhất quán."),
        ("3.2.3.2. Biểu đồ tuần tự luồng quản lý hội thoại", "seq_conversation", "Hình 3.11. Biểu đồ tuần tự luồng quản lý hội thoại", "Luồng quản lý hội thoại thể hiện cách ứng dụng tải danh sách hội thoại, mở chi tiết một hội thoại và cập nhật trạng thái hội thoại. CSDL được bổ sung vào biểu đồ để làm rõ rằng danh sách hội thoại, tin nhắn gần nhất và các thông tin bổ sung như ghim/lưu trữ đều được đọc hoặc ghi thông qua hệ thống. Khi người dùng chọn đổi tên, ghim, lưu trữ hoặc xoá, hệ thống kiểm tra quyền sở hữu trước khi ghi thay đổi và trả lại trạng thái mới cho giao diện."),
        ("3.2.3.3. Biểu đồ tuần tự luồng hỏi đáp pháp luật", "seq_chat_answer", "Hình 3.12. Biểu đồ tuần tự luồng hỏi đáp pháp luật", "Luồng hỏi đáp là nghiệp vụ trọng tâm của ứng dụng. Sau khi người dùng gửi câu hỏi, hệ thống lưu câu hỏi và lịch sử hội thoại, truy hồi dữ liệu pháp luật liên quan, sau đó trả tiến trình xử lý cho giao diện. Khi câu trả lời hoàn tất, nội dung trả lời và nguồn tham chiếu được lưu lại để người dùng có thể mở lại hội thoại sau này. Việc biểu diễn CSDL trong luồng này giúp làm rõ rằng chức năng hỏi đáp không chỉ trả lời tức thời mà còn duy trì lịch sử, nguồn dẫn và câu hỏi gợi ý."),
        ("3.2.3.4. Biểu đồ tuần tự luồng tư vấn có hướng dẫn", "seq_guided", "Hình 3.13. Biểu đồ tuần tự luồng tư vấn có hướng dẫn", "Luồng tư vấn có hướng dẫn được thiết kế cho các tình huống câu hỏi ban đầu còn thiếu ngữ cảnh. Hệ thống trước hết kiểm tra dữ liệu pháp luật liên quan để tạo câu hỏi làm rõ phù hợp. Sau khi người dùng bổ sung thông tin, hệ thống tiếp tục truy hồi căn cứ pháp luật rồi mới tạo kết quả tư vấn cuối cùng. Cách xử lý này giúp câu trả lời cụ thể hơn, đặc biệt với những vấn đề pháp lý phụ thuộc nhiều vào hoàn cảnh thực tế."),
        ("3.2.3.5. Biểu đồ tuần tự luồng tra cứu thư viện pháp luật", "seq_law_library", "Hình 3.14. Biểu đồ tuần tự luồng tra cứu thư viện pháp luật", "Luồng tra cứu thư viện pháp luật phục vụ nhu cầu xem văn bản và điều luật theo hướng có cấu trúc. Người dùng có thể mở danh sách văn bản, lọc theo năm/chủ đề/từ khoá, sau đó chọn một văn bản hoặc điều luật để xem chi tiết. CSDL được đặt thành một thành phần riêng trong biểu đồ để thể hiện rõ dữ liệu trả về là dữ liệu pháp luật đã được lưu trữ, khác với luồng hỏi đáp nơi hệ thống cần tổng hợp câu trả lời mới."),
        ("3.2.3.6. Biểu đồ tuần tự luồng tìm kiếm ngữ nghĩa", "seq_semantic_search", "Hình 3.15. Biểu đồ tuần tự luồng tìm kiếm ngữ nghĩa", "Luồng tìm kiếm ngữ nghĩa cho phép người dùng nhập truy vấn tự nhiên thay vì chỉ tìm theo từ khoá chính xác. Hệ thống phân tích truy vấn, truy hồi dữ liệu từ CSDL Vector, xếp hạng lại kết quả rồi trả danh sách điều luật liên quan cho giao diện. Việc tách CSDL Vector trong biểu đồ giúp nhấn mạnh khác biệt giữa tra cứu thông thường và tìm kiếm ngữ nghĩa: dữ liệu được tìm theo độ tương đồng nội dung thay vì chỉ theo bộ lọc có cấu trúc."),
        ("3.2.3.7. Biểu đồ tuần tự luồng thống kê và theo dõi tác vụ", "seq_dashboard_tracker", "Hình 3.16. Biểu đồ tuần tự luồng thống kê và theo dõi tác vụ", "Luồng thống kê giúp quản trị viên nắm nhanh tình trạng dữ liệu và các tác vụ xử lý gần đây. Khi mở trang thống kê, hệ thống tổng hợp số liệu từ CSDL để trả về các chỉ số chính. Khi chuyển sang trang theo dõi tiến trình, giao diện thiết lập kênh nhận cập nhật theo thời gian thực và tải danh sách tác vụ hiện có. Mỗi khi tác vụ thay đổi trạng thái, hệ thống ghi trạng thái mới và gửi cập nhật để giao diện phản ánh ngay tiến trình xử lý."),
        ("3.2.3.8. Biểu đồ tuần tự luồng tải lên và xử lý văn bản pháp luật", "seq_upload_processing", "Hình 3.17. Biểu đồ tuần tự luồng tải lên và xử lý văn bản pháp luật", "Luồng tải lên văn bản là nghiệp vụ quản trị quan trọng vì làm thay đổi kho dữ liệu phục vụ hỏi đáp và tra cứu. Giao diện kiểm tra định dạng tệp trước khi gửi lên hệ thống. Hệ thống tạo tác vụ xử lý trong CSDL, kiểm tra trùng lặp, trích xuất nội dung, lưu văn bản và điều luật, sau đó vector hoá dữ liệu vào CSDL Vector. Nếu tệp không hợp lệ, xử lý thất bại hoặc quản trị viên hủy yêu cầu, trạng thái tác vụ được cập nhật để quản trị viên có thể theo dõi và xử lý tiếp."),
        ("3.2.3.9. Biểu đồ tuần tự luồng quản lý văn bản phía quản trị", "seq_admin_document", "Hình 3.18. Biểu đồ tuần tự luồng quản lý văn bản phía quản trị", "Sau khi văn bản đã được đưa vào hệ thống, quản trị viên cần có luồng quản lý riêng để tìm kiếm, lọc, xoá văn bản hoặc làm mới dữ liệu hiển thị. CSDL lưu nội dung văn bản và thông tin quản trị, trong khi CSDL Vector lưu dữ liệu phục vụ tìm kiếm ngữ nghĩa. Khi xoá văn bản, hệ thống phải xử lý đồng thời hai lớp dữ liệu này để tránh tình trạng giao diện còn hiển thị văn bản không còn khả dụng hoặc kết quả tìm kiếm vẫn trỏ đến dữ liệu đã bị xoá."),
    ]
    for h, key, cap, desc in seqs:
        heading(doc, h, 4)
        add_paragraph(doc, desc)
        add_figure(doc, diagrams[key], cap)


def section_design(doc, diagrams):
    heading(doc, "3.3. Thiết kế hệ thống", 2)
    heading(doc, "3.3.1. Thiết kế cơ sở dữ liệu", 3)
    add_paragraph(doc, "Dữ liệu của hệ thống được tổ chức theo đặc trưng sử dụng của từng nhóm nghiệp vụ. PostgreSQL lưu các dữ liệu có quan hệ chặt chẽ và cần đảm bảo tính toàn vẹn, chẳng hạn tài khoản, hội thoại, tin nhắn và tác vụ xử lý tài liệu. MongoDB lưu nội dung văn bản pháp luật ở dạng document linh hoạt, phù hợp với cấu trúc văn bản có nhiều thuộc tính mở rộng. ChromaDB lưu vector embedding của các đoạn văn bản để phục vụ truy hồi ngữ nghĩa trong các chức năng hỏi đáp và tìm kiếm.")
    add_figure(doc, diagrams["hinh_3_11_erd_postgresql"], "Hình 3.19. Sơ đồ quan hệ dữ liệu trong PostgreSQL")
    add_paragraph(doc, "Trong phạm vi PostgreSQL, các bảng được thiết kế xoay quanh ba nhóm dữ liệu chính. Nhóm thứ nhất là dữ liệu người dùng và phiên đăng nhập, phục vụ xác thực và phân quyền. Nhóm thứ hai là dữ liệu hội thoại, dùng để lưu lịch sử trao đổi giữa người dùng và hệ thống. Nhóm thứ ba là dữ liệu tác vụ xử lý tài liệu, dùng cho quá trình quản trị văn bản pháp luật. Cách tổ chức này giúp tách rõ dữ liệu phục vụ người dùng cuối với dữ liệu phục vụ vận hành hệ thống.")

    add_table(doc, "Bảng 3.13. Bảng users lưu thông tin người dùng", ["Trường", "Kiểu", "Mô tả"], [
        ["id", "UUID", "Khoá chính của người dùng."],
        ["email", "String", "Email đăng nhập, duy nhất trong hệ thống."],
        ["hashed_password", "String", "Mật khẩu đã được băm trước khi lưu."],
        ["full_name", "String", "Họ tên người dùng."],
        ["phone_number", "String", "Số điện thoại, có thể rỗng."],
        ["role", "String", "Phân biệt người dùng thường và quản trị viên."],
        ["is_active", "Boolean", "Trạng thái hoạt động của tài khoản."],
        ["last_login_at", "DateTime", "Thời điểm đăng nhập gần nhất."],
    ], [4.0, 3.0, 9.0])
    add_table(doc, "Bảng 3.14. Bảng conversations lưu hội thoại", ["Trường", "Kiểu", "Mô tả"], [
        ["id", "UUID", "Khoá chính của hội thoại."],
        ["user_id", "UUID", "Khoá ngoại liên kết tới users."],
        ["title", "String", "Tiêu đề hội thoại, có thể sinh từ chủ đề câu hỏi đầu tiên."],
        ["is_pinned", "Boolean", "Đánh dấu hội thoại được ghim."],
        ["is_archived", "Boolean", "Đánh dấu hội thoại đã lưu trữ."],
        ["message_count", "Integer", "Số lượng tin nhắn trong hội thoại."],
        ["last_message_at", "DateTime", "Thời điểm có tin nhắn cuối."],
    ], [4.0, 3.0, 9.0])
    add_table(doc, "Bảng 3.15. Bảng messages lưu tin nhắn", ["Trường", "Kiểu", "Mô tả"], [
        ["id", "UUID", "Khoá chính của tin nhắn."],
        ["conversation_id", "UUID", "Khoá ngoại liên kết tới conversations."],
        ["question_id", "UUID", "Tham chiếu tới tin nhắn câu hỏi khi cần ghép cặp hỏi - đáp."],
        ["role", "String", "Vai trò của tin nhắn: user hoặc assistant."],
        ["content", "Text", "Nội dung tin nhắn."],
        ["sources", "JSONB", "Danh sách nguồn pháp luật được dùng trong câu trả lời."],
        ["metadata", "JSONB", "Thông tin bổ sung như phân tích truy vấn, thời gian xử lý hoặc câu hỏi gợi ý."],
    ], [4.0, 3.0, 9.0])
    add_table(doc, "Bảng 3.16. Bảng document_tasks lưu tiến trình xử lý tài liệu", ["Trường", "Kiểu", "Mô tả"], [
        ["id", "UUID", "Khoá chính của tác vụ xử lý tài liệu."],
        ["user_id", "UUID", "Quản trị viên tạo tác vụ."],
        ["filename", "String", "Tên tệp PDF gốc."],
        ["file_size_bytes", "BigInteger", "Kích thước tệp được tải lên."],
        ["status", "Enum", "Trạng thái pending, processing, completed, failed hoặc cancelled."],
        ["progress", "Integer", "Phần trăm xử lý."],
        ["current_step", "String", "Bước xử lý hiện tại để hiển thị trên Admin Web."],
        ["law_id", "String", "Mã văn bản pháp luật sau khi phân tách dữ liệu."],
        ["article_count", "Integer", "Số điều luật được trích xuất."],
        ["error_message", "Text", "Thông tin lỗi nếu tác vụ thất bại."],
    ], [4.0, 3.0, 9.0])
    add_table(doc, "Bảng 3.17. Bảng refresh_tokens lưu phiên đăng nhập", ["Trường", "Kiểu", "Mô tả"], [
        ["id", "UUID", "Khoá chính của refresh token."],
        ["user_id", "UUID", "Khoá ngoại liên kết tới users."],
        ["token", "String", "Refresh token được lưu để kiểm tra phiên đăng nhập."],
        ["expires_at", "DateTime", "Thời điểm hết hạn."],
        ["is_revoked", "Boolean", "Đánh dấu token đã bị thu hồi."],
    ], [4.0, 3.0, 9.0])

    add_table(doc, "Bảng 3.18. Collection VietnamLawDB.articles trong MongoDB", ["Trường", "Mô tả"], [
        ["_id", "Khoá duy nhất của điều luật, thường được ghép từ law_id và article_id."],
        ["law_id", "Mã văn bản pháp luật."],
        ["article_id", "Số hoặc mã điều luật trong văn bản."],
        ["title", "Tiêu đề điều luật."],
        ["text", "Nội dung nguyên văn của điều luật."],
        ["source_url", "Đường dẫn tệp PDF gốc nếu có."],
        ["full_content_search", "Nội dung ghép phục vụ tìm kiếm toàn văn."],
        ["metadata", "Thông tin chủ đề, từ khoá, tóm tắt và năm ban hành."],
    ], [4.5, 11.5])
    add_paragraph(doc, "Đối với dữ liệu văn bản pháp luật, MongoDB và ChromaDB được sử dụng bổ trợ cho nhau. MongoDB giữ nội dung có thể đọc và hiển thị lại cho người dùng, trong khi ChromaDB giữ các đoạn văn bản đã được vector hoá để phục vụ truy hồi theo ngữ nghĩa. Việc tách hai lớp lưu trữ này giúp hệ thống vừa đáp ứng nhu cầu tra cứu văn bản nguyên bản, vừa hỗ trợ các chức năng hỏi đáp dựa trên nội dung liên quan.")
    add_table(doc, "Bảng 3.19. Collection vietnamese_law trong ChromaDB", ["Trường", "Mô tả"], [
        ["id", "Mã đoạn văn bản theo dạng law_id_article_id_chunkN."],
        ["document", "Nội dung đoạn văn bản gồm tiêu đề điều luật và phần văn bản."],
        ["embedding", "Vector biểu diễn ngữ nghĩa của đoạn văn bản."],
        ["metadata.law_id", "Mã văn bản pháp luật."],
        ["metadata.article_id", "Mã điều luật."],
        ["metadata.chunk_index", "Thứ tự đoạn trong điều luật."],
        ["metadata.total_chunks", "Tổng số đoạn của điều luật."],
        ["metadata.year", "Năm ban hành dùng cho lọc hoặc ưu tiên văn bản mới."],
        ["metadata.topics", "Danh sách chủ đề được lưu dạng chuỗi JSON."],
        ["metadata.keywords", "Danh sách từ khoá được lưu dạng chuỗi JSON."],
    ], [4.5, 11.5])

    heading(doc, "3.3.2. Thiết kế các API", 3)
    add_paragraph(doc, "Các API của hệ thống được chia thành hai nhóm theo vai trò sử dụng. Main Service cung cấp API cho ứng dụng di động và Admin Web, đảm nhiệm xác thực, quản lý hội thoại, quản lý văn bản và điều phối các yêu cầu nghiệp vụ. RAG Service chủ yếu đóng vai trò dịch vụ nội bộ, nhận yêu cầu từ Main Service để thực hiện truy hồi, sinh câu trả lời, tư vấn có hướng dẫn và cập nhật dữ liệu vector vào ChromaDB.")
    add_paragraph(doc, "Cách phân chia API như trên giúp giao diện người dùng không phải làm việc trực tiếp với dịch vụ xử lý AI hoặc CSDL Vector. Các yêu cầu từ phía giao diện được tập trung qua Main Service, sau đó Main Service mới điều phối sang RAG Service khi nghiệp vụ cần truy hồi hoặc sinh câu trả lời. Thiết kế này giúp kiểm soát xác thực, phân quyền và ghi nhận lịch sử xử lý ở một điểm thống nhất.")
    add_table(doc, "Bảng 3.20. Nhóm API xác thực", ["API", "Phương thức", "Chức năng"], [
        ["/api/v1/auth/signup", "POST", "Đăng ký tài khoản mới."],
        ["/api/v1/auth/login", "POST", "Đăng nhập và nhận token."],
        ["/api/v1/auth/refresh", "POST", "Làm mới access token."],
        ["/api/v1/auth/logout", "POST", "Đăng xuất một phiên."],
        ["/api/v1/auth/me", "GET/PUT", "Xem hoặc cập nhật thông tin người dùng."],
        ["/api/v1/auth/change-password", "PUT", "Đổi mật khẩu."],
    ], [6.0, 3.0, 7.0])
    add_table(doc, "Bảng 3.21. Nhóm API hội thoại và hỏi đáp", ["API", "Phương thức", "Chức năng"], [
        ["/api/v1/chat/conversations", "GET", "Lấy danh sách hội thoại."],
        ["/api/v1/chat/conversations/{id}", "GET/PUT/DELETE", "Xem, cập nhật hoặc xoá hội thoại."],
        ["/api/v1/chat/conversations/{id}/messages", "GET", "Lấy tin nhắn của hội thoại."],
        ["/api/v1/chat/messages", "POST", "Gửi tin nhắn và nhận câu trả lời dạng thường."],
        ["/api/v1/chat/messages/stream", "POST", "Gửi tin nhắn và nhận tiến trình/câu trả lời qua SSE."],
        ["/api/v1/chat/messages/{message_id}/suggested-questions", "GET", "Lấy câu hỏi gợi ý cho tin nhắn assistant."],
    ], [6.8, 3.2, 6.0])
    add_table(doc, "Bảng 3.22. Nhóm API tư vấn có hướng dẫn", ["API", "Phương thức", "Chức năng"], [
        ["/api/v1/guided/clarify", "POST", "Sinh câu hỏi làm rõ từ câu hỏi ban đầu."],
        ["/api/v1/guided/answer", "POST", "Tạo câu trả lời cuối cùng sau khi người dùng chọn đáp án."],
        ["/api/v1/guided/answer/stream", "POST", "Tạo câu trả lời cuối cùng và trả tiến trình qua SSE."],
    ], [6.5, 3.0, 6.5])
    add_table(doc, "Bảng 3.23. Nhóm API tra cứu pháp luật", ["API", "Phương thức", "Chức năng"], [
        ["/api/v1/laws", "GET", "Lấy danh sách văn bản pháp luật."],
        ["/api/v1/laws/search", "GET/POST", "Tìm kiếm văn bản theo nội dung, mã văn bản, điều, năm, chủ đề hoặc từ khoá."],
        ["/api/v1/laws/ai-search", "POST", "Tìm kiếm ngữ nghĩa bằng RAG Service."],
        ["/api/v1/laws/topics", "GET", "Lấy danh sách chủ đề."],
        ["/api/v1/laws/years", "GET", "Lấy danh sách năm ban hành."],
        ["/api/v1/laws/detail", "GET", "Lấy chi tiết một điều luật."],
    ], [6.2, 3.0, 6.8])
    add_table(doc, "Bảng 3.24. Nhóm API quản trị tài liệu", ["API", "Phương thức", "Chức năng"], [
        ["/api/v1/documents/ws", "WebSocket", "Theo dõi tiến trình xử lý tài liệu theo thời gian thực."],
        ["/api/v1/documents/upload-v2", "POST", "Tải lên và xử lý PDF văn bản pháp luật."],
        ["/api/v1/documents/tasks", "GET", "Lấy lịch sử tác vụ xử lý tài liệu."],
        ["/api/v1/documents/tasks/{task_id}/cancel", "POST", "Huỷ tác vụ đang chạy hoặc đang treo."],
        ["/api/v1/documents/tasks/{task_id}", "DELETE", "Xoá nhật ký tác vụ."],
        ["/api/v1/laws/{law_id}", "DELETE", "Xoá văn bản khỏi MongoDB, ChromaDB và dữ liệu hiển thị."],
    ], [6.5, 3.0, 6.5])
    add_table(doc, "Bảng 3.25. Nhóm API nội bộ của RAG Service", ["API", "Phương thức", "Chức năng"], [
        ["/api/v1/rag/agent-search/stream", "POST", "Chạy luồng Agentic RAG và trả tiến trình qua SSE."],
        ["/api/v1/rag/guided-clarify", "POST", "Phân tích câu hỏi và sinh câu hỏi làm rõ."],
        ["/api/v1/rag/guided-answer/stream", "POST", "Chạy luồng tư vấn có hướng dẫn và trả tiến trình qua SSE."],
        ["/api/v1/rag/semantic-search", "POST", "Tìm kiếm ngữ nghĩa trong ChromaDB."],
        ["/api/v1/ingest/articles", "POST", "Chia đoạn, tạo embedding và lưu điều luật vào ChromaDB."],
        ["/api/v1/ingest/articles/{law_id}", "DELETE", "Xoá các đoạn vector của một văn bản khi hoàn tác hoặc xoá văn bản."],
    ], [6.8, 3.0, 6.2])


def section_summary(doc):
    heading(doc, "3.4. Tổng kết chương 3", 2)
    add_paragraph(doc, "Chương 3 đã trình bày quá trình phân tích và thiết kế hệ thống Vietnam Law Chatbot theo các nhóm nội dung chính: yêu cầu hệ thống, ca sử dụng, luồng tuần tự, cơ sở dữ liệu và API. Thông qua phần phân tích, hệ thống được xác định có hai nhóm tác nhân chính là người dùng cuối và quản trị viên. Các chức năng trọng tâm tập trung vào hỏi đáp pháp luật, tư vấn có hướng dẫn, tra cứu thư viện văn bản, tìm kiếm ngữ nghĩa và cập nhật dữ liệu pháp luật.")
    add_paragraph(doc, "Kết quả thiết kế cho thấy hệ thống được tổ chức theo mô hình nhiều thành phần, có sự phân tách rõ giữa xử lý nghiệp vụ, lưu trữ dữ liệu và xử lý tri thức pháp luật. PostgreSQL đảm nhiệm dữ liệu tài khoản, hội thoại và tác vụ xử lý; MongoDB lưu nội dung văn bản pháp luật ở dạng linh hoạt; ChromaDB phục vụ truy hồi vector cho các chức năng dựa trên ngữ nghĩa. Các API cùng cơ chế cập nhật tiến trình theo thời gian thực là cơ sở để triển khai giao diện, kiểm thử chức năng và đánh giá hệ thống trong chương tiếp theo.")


def main():
    diagrams = render_diagrams()
    doc = Document()
    configure_styles(doc)
    add_intro(doc)
    section_requirements(doc)
    section_analysis(doc, diagrams)
    section_design(doc, diagrams)
    section_summary(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
