from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "02_Chuong_2_ban_bo_sung_phan_tich_rag.docx"
OUT = ROOT / "02_Chuong_2_ban_dinh_dang_theo_mau.docx"


BODY_FONT = "Times New Roman"


def ensure_paragraph_style(doc: Document, name: str) -> _ParagraphStyle:
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_font(font, size=None, bold=None, italic=None, all_caps=None):
    font.name = BODY_FONT
    font._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if all_caps is not None:
        font.all_caps = all_caps


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_font(normal.font, 14)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = None
    pf.right_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h1 = doc.styles["Heading 1"]
    set_font(h1.font, 16, bold=True, all_caps=True)
    pf = h1.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h2 = doc.styles["Heading 2"]
    set_font(h2.font, 14, bold=True)
    pf = h2.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h3 = doc.styles["Heading 3"]
    set_font(h3.font, 14, bold=True, italic=True)
    pf = h3.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h4 = doc.styles["Heading 4"]
    set_font(h4.font, 14, italic=True)
    pf = h4.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)

    fig = ensure_paragraph_style(doc, "Tên hình vẽ")
    fig.base_style = normal
    set_font(fig.font, 14, italic=True)
    pf = fig.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    tbl = ensure_paragraph_style(doc, "Tên bảng")
    tbl.base_style = normal
    set_font(tbl.font, 14, bold=True)
    pf = tbl.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    caption = doc.styles["Caption"]
    set_font(caption.font, 10, bold=True)
    caption.paragraph_format.first_line_indent = Cm(0)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def set_run_fonts(paragraph, size=14):
    for run in paragraph.runs:
        set_font(run.font, size)


def apply_paragraph_formatting(doc: Document):
    for p in doc.paragraphs:
        text = p.text.strip()

        if not text:
            continue

        if paragraph_has_drawing(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            continue

        if text.startswith("CHƯƠNG 2"):
            p.style = doc.styles["Heading 1"]
            p.alignment = None
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(18)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            set_run_fonts(p, 16)
            for run in p.runs:
                run.bold = True
            continue

        if re.match(r"^2\.\d+\.\s", text):
            p.style = doc.styles["Heading 2"]
            p.alignment = None
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            set_run_fonts(p, 14)
            for run in p.runs:
                run.bold = True
                run.italic = False
            continue

        if re.match(r"^2\.\d+\.\d+\.\s", text):
            p.style = doc.styles["Heading 3"]
            p.alignment = None
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            set_run_fonts(p, 14)
            for run in p.runs:
                run.bold = True
                run.italic = True
            continue

        if text.startswith("Hình 2."):
            p.style = doc.styles["Tên hình vẽ"]
            p.alignment = None
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            set_run_fonts(p, 14)
            for run in p.runs:
                run.italic = True
            continue

        if text.startswith("Bảng 2.") and re.match(r"^Bảng 2\.\d+\.", text):
            p.style = doc.styles["Tên bảng"]
            p.alignment = None
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            set_run_fonts(p, 14)
            for run in p.runs:
                run.bold = True
            continue

        p.style = doc.styles["Normal"]
        set_run_fonts(p, 14)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        if text.startswith("•") or re.match(r"^\d+\.\s", text):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
        else:
            p.paragraph_format.first_line_indent = Cm(1.25)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_tables(doc: Document):
    for table in doc.tables:
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.style = doc.styles["Normal"]
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.left_indent = Cm(0)
                    p.paragraph_format.right_indent = Cm(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        set_font(run.font, 12, bold=(True if row_idx == 0 else run.bold))


def format_images(doc: Document):
    usable_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    max_width = usable_width
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = shape.height / shape.width
            shape.width = max_width
            shape.height = int(max_width * ratio)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    configure_styles(doc)
    apply_paragraph_formatting(doc)
    format_images(doc)
    format_tables(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
