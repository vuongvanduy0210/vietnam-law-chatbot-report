from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "Bao_Cao" / "02_Chuong_2_ban_bo_sung_phan_tich.docx"
OUT = ROOT / "Bao_Cao" / "02_Chuong_2_ban_bo_sung_phan_tich_rag.docx"


SECTION_241 = [
    "Kho pháp luật nội bộ là nền tảng quan trọng nhất của pipeline RAG, vì đây là nơi lưu các điều luật đã được chuẩn hóa và có thể truy xuất lại theo ngữ nghĩa. Thay vì để LLMs tự suy luận từ kiến thức có sẵn, hệ thống buộc quá trình trả lời phải bắt đầu từ các đoạn văn bản pháp luật được lưu trong kho tri thức. Cách tổ chức này giúp câu trả lời có căn cứ rõ ràng hơn, đồng thời cho phép hệ thống trích dẫn lại số hiệu văn bản, điều khoản và nội dung gốc khi cần.",
    "Đơn vị dữ liệu cơ bản trong kho tri thức không phải là toàn bộ văn bản pháp luật, mà là từng điều luật hoặc từng đoạn được tách ra từ điều luật dài. Mỗi bản ghi trong kho vector gồm ba phần chính: nội dung văn bản dùng để tính embedding, vector embedding biểu diễn ngữ nghĩa của nội dung đó và metadata pháp lý đi kèm. Metadata bao gồm số hiệu văn bản, mã điều, tiêu đề điều, năm ban hành, nhóm chủ đề, từ khóa, tóm tắt, chỉ số đoạn và tổng số đoạn của điều luật. Nhờ đó, khi một đoạn được truy hồi, hệ thống không chỉ biết đoạn đó nói gì mà còn biết nó thuộc văn bản nào, điều nào và nằm ở vị trí nào trong cấu trúc nguồn.",
    "Trước khi đưa dữ liệu vào ChromaDB, hệ thống áp dụng chiến thuật chia nhỏ theo hướng lai. Các điều luật ngắn được giữ nguyên để tránh làm vỡ mạch pháp lý; chỉ những điều luật dài vượt ngưỡng mới được tách thành nhiều đoạn. Ngưỡng tách được đặt khoảng 1.000 từ, với phần chồng lấn khoảng 150 từ giữa hai đoạn liền kề. Phần chồng lấn này có vai trò giữ lại ngữ cảnh ở ranh giới đoạn, đặc biệt với các điều luật có nhiều khoản, điểm hoặc điều kiện áp dụng nối tiếp nhau. Khi cắt đoạn, hệ thống ưu tiên điểm ngắt tự nhiên theo dấu câu thay vì cắt cứng theo số từ, nhờ đó nội dung trong mỗi chunk vẫn tương đối hoàn chỉnh và dễ được mô hình hiểu đúng.",
    "Nội dung được đưa vào embedding không chỉ là phần thân điều luật mà còn ghép thêm tiêu đề điều ở phía trước. Đây là một chi tiết quan trọng vì nhiều điều luật có phần thân dùng đại từ hoặc cụm diễn đạt phụ thuộc vào tiêu đề. Khi tiêu đề được đưa vào cùng nội dung, vector của chunk thể hiện rõ hơn chủ đề pháp lý của đoạn, giúp tăng khả năng tìm đúng kết quả đối với những câu hỏi ngắn như \"mức phạt khi vượt đèn đỏ\" hoặc \"điều kiện đăng ký kết hôn\".",
    "Sau bước chia nhỏ, mỗi chunk được mã hóa bằng mô hình embedding tiếng Việt để tạo vector ngữ nghĩa. Vector này được lưu vào collection ChromaDB cùng với nội dung gốc và metadata. Về mặt logic, collection có thể được xem như một bảng tri thức chuyên biệt cho truy hồi ngữ nghĩa: mỗi dòng là một chunk, mỗi chunk có mã định danh duy nhất, một vector, phần văn bản có thể trích dẫn và tập thuộc tính phục vụ lọc, hiển thị hoặc kiểm tra nguồn. Việc tách rõ nội dung, vector và metadata giúp hệ thống vừa tìm kiếm nhanh bằng vector, vừa giữ được khả năng giải thích kết quả bằng thông tin pháp lý có cấu trúc.",
    "Khi người dùng đặt câu hỏi, hệ thống không đưa nguyên văn câu hỏi vào kho vector một cách trực tiếp trong mọi trường hợp. Trước đó, bước phân tích câu hỏi tạo ra một truy vấn nội bộ ngắn gọn hơn, tập trung vào vấn đề pháp lý cần tra cứu. Truy vấn này được mã hóa thành embedding và gửi đến ChromaDB để tìm các chunk có khoảng cách ngữ nghĩa gần nhất. Kết quả trả về ở giai đoạn đầu là một tập ứng viên rộng, thường lớn hơn số nguồn cuối cùng cần dùng, vì mục tiêu của bước này là tránh bỏ sót điều luật liên quan.",
    "Trong triển khai hiện tại, hệ thống ưu tiên tìm kiếm vector trên phạm vi toàn bộ collection thay vì lọc metadata quá sớm. Lý do là câu hỏi tự nhiên của người dùng thường không nêu đầy đủ số hiệu văn bản, lĩnh vực hoặc năm ban hành; nếu áp dụng bộ lọc cứng ngay từ đầu, hệ thống có thể loại nhầm những điều luật thật sự liên quan. Metadata vẫn được lưu đầy đủ, nhưng được dùng chủ yếu để giải thích nguồn, nhận diện văn bản, xử lý xung đột thời gian và hỗ trợ các bước kiểm tra sau truy hồi.",
    "Sau khi ChromaDB trả về tập ứng viên, pipeline tiếp tục xếp hạng lại bằng cross-encoder. Khác với bi-encoder, cross-encoder đánh giá trực tiếp từng cặp câu hỏi và đoạn văn bản, vì vậy có khả năng phân biệt tốt hơn giữa các đoạn chỉ giống về từ khóa và các đoạn thật sự trả lời đúng vấn đề. Điểm cuối cùng được kết hợp từ điểm tìm kiếm vector và điểm xếp hạng lại, trong đó điểm của cross-encoder được đặt trọng số cao hơn. Với dữ liệu pháp luật, hệ thống còn cộng thêm một mức ưu tiên nhỏ cho văn bản mới hơn và giảm ưu tiên với văn bản quá cũ, nhằm hạn chế trường hợp câu trả lời dựa vào quy định đã bị thay thế.",
    "Các đoạn vượt qua ngưỡng tin cậy mới được đưa vào ngữ cảnh cho Agent. Mỗi đoạn được đóng gói kèm số hiệu văn bản, điều, tiêu đề, năm ban hành, lĩnh vực và nội dung nguyên văn. Việc đóng gói theo ranh giới rõ ràng giúp Agent biết đâu là phần có thể trích dẫn, đâu chỉ là thông tin điều phối của hệ thống. Nếu không tìm thấy kết quả đủ tin cậy trong kho nội bộ, pipeline không cố tạo câu trả lời từ dữ liệu yếu mà chuyển sang hướng kiểm tra nguồn cập nhật bên ngoài.",
    "Như vậy, luồng tra cứu nội bộ gồm bốn lớp xử lý liên tiếp: chuẩn hóa dữ liệu thành các chunk có ngữ cảnh, mã hóa chunk thành embedding và lưu trong ChromaDB, truy hồi tập ứng viên rộng bằng vector search, sau đó xếp hạng lại và lọc theo ngưỡng tin cậy trước khi cung cấp bằng chứng cho Agent. Thiết kế này phù hợp với bài toán tư vấn pháp luật vì nó không chỉ cần tìm đoạn văn bản có vẻ giống câu hỏi, mà còn cần giữ được nguồn gốc điều khoản, hạn chế nhiễu, ưu tiên văn bản còn phù hợp và tạo cơ sở cho bước kiểm chứng câu trả lời ở cuối pipeline.",
]

SECTION_274_BRIDGE = (
    "Riêng phần ghi dữ liệu vào kho vector tuân theo chiến lược đã phân tích ở mục 2.4.1: "
    "điều luật được chuẩn hóa thành các chunk có ngữ cảnh, mã hóa embedding và lưu kèm metadata "
    "để phục vụ truy hồi. Vì vậy, trong mục cập nhật tri thức, trọng tâm còn lại là bảo đảm tính "
    "nhất quán giữa kho quản trị và kho truy hồi khi một tác vụ cập nhật thành công hoặc thất bại."
)


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc, prefix):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Không tìm thấy đoạn bắt đầu bằng: {prefix}")


def main():
    doc = Document(SRC)

    # Replace current short 2.4.1 body with a detailed analysis grounded in code.
    idx_241 = find_paragraph(doc, "2.4.1.")
    idx_242 = find_paragraph(doc, "2.4.2.")
    for paragraph in list(doc.paragraphs[idx_241 + 1 : idx_242]):
        text = paragraph.text.strip()
        if text.startswith("Hình 2.6.") or not text:
            continue
        remove_paragraph(paragraph)

    anchor = doc.paragraphs[idx_241]
    for text in reversed(SECTION_241):
        insert_paragraph_after(anchor, text, style="Normal")

    # Remove the duplicated chunking subsection from 2.7. The strategy is now
    # explained in 2.4.1, where the internal knowledge source is analyzed.
    idx_275 = find_paragraph(doc, "2.7.5.")
    idx_28 = find_paragraph(doc, "2.8.")
    for paragraph in list(doc.paragraphs[idx_275:idx_28]):
        remove_paragraph(paragraph)

    idx_28_after_delete = find_paragraph(doc, "2.8.")
    paragraph_before_28 = doc.paragraphs[idx_28_after_delete - 1]
    insert_paragraph_after(paragraph_before_28, SECTION_274_BRIDGE, style="Normal")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
