from pathlib import Path

from docx import Document
from docx.shared import RGBColor


DOCX = Path(__file__).resolve().parents[1] / "02_Chuong_2_ban_dinh_dang_theo_mau.docx"
BLACK = RGBColor(0, 0, 0)


REPLACEMENTS = {
    "• ": "",
    "—": "-",
    "→": "đến",
    "$$\\text{FinalScore} = 0.3 \\times \\text{BiEncoderScore} + 0.7 \\times \\text{CrossEncoderScore}$$":
        "FinalScore = 0,3 x BiEncoderScore + 0,7 x CrossEncoderScore",
}


def replace_text(paragraph):
    text = paragraph.text
    new_text = text
    for old, new in REPLACEMENTS.items():
        new_text = new_text.replace(old, new)

    # Clean accidental doubled spaces caused by removing symbols.
    while "  " in new_text:
        new_text = new_text.replace("  ", " ")

    if new_text != text:
        paragraph.text = new_text.strip()


def force_paragraph_black(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = BLACK


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def main():
    doc = Document(DOCX)

    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            try:
                style.font.color.rgb = BLACK
            except Exception:
                pass

    for paragraph in iter_all_paragraphs(doc):
        replace_text(paragraph)
        force_paragraph_black(paragraph)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
