from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TARGETS = [
    ROOT / "Vương Văn Duy_Báo cáo.docx",
    ROOT / "Bao_Cao" / "02_Chuong_2_ban_dinh_dang_theo_mau.docx",
]
BLACK = RGBColor(0, 0, 0)


def set_common_font(run, size: int) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), "Times New Roman")


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph, False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph, True


def target_size(paragraph, in_table: bool) -> int:
    if in_table:
        return 12
    style = paragraph.style.name if paragraph.style else ""
    if style.startswith("Heading 1"):
        return 16
    return 14


def fix(path: Path) -> int:
    doc = Document(path)
    changed = 0
    for paragraph, in_table in iter_paragraphs(doc):
        size = target_size(paragraph, in_table)
        for run in paragraph.runs:
            if "Next.js" in run.text:
                set_common_font(run, size)
                changed += 1
    if changed:
        doc.save(path)
    return changed


def main() -> None:
    for target in TARGETS:
        if target.exists():
            print(target, fix(target))


if __name__ == "__main__":
    main()
