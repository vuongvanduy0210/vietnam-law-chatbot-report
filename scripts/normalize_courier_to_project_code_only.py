from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.run import Run


ROOT = Path(__file__).resolve().parents[2]
TARGETS = [
    ROOT / "Vương Văn Duy_Báo cáo.docx",
    ROOT / "Bao_Cao" / "02_Chuong_2_ban_dinh_dang_theo_mau.docx",
]

BLACK = RGBColor(0, 0, 0)

# Only identifiers, event names, routes, payload fields, table/collection names,
# and config constants that are direct project/code artifacts should use code font.
CODE_RE = re.compile(
    "|".join(
        [
            r"/api/[A-Za-z0-9_/:.-]+",
            r"\{[^{}]{1,120}\}",
            r"\bData\.[A-Za-z_][A-Za-z0-9_]*\b",
            r"\b[A-Z][A-Z0-9]+(?:_[A-Z0-9]+)+\b",
            r"\b[a-z][a-z0-9]+(?:_[a-z0-9]+)+\b",
            r"\bVietnamLawDB\.articles\b",
            r"\bvietnamese_law\b",
            r"\b(?:users|conversations|messages|document_tasks|refresh_tokens)\b",
            r"\b(?:retrieve_internal_law|search_web_for_law)\b",
            r"\b(?:ready|progress|done|error|timeout|pending|processing|completed|failed|cancelled)\b",
            r"\b(?:admin_token|sign_in)\b",
            r"\bef_construction=100, ef_search=100, max_neighbors=16\b",
        ]
    )
)


def set_font(run: Run, name: str, size: int) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), name)


def normal_size(paragraph, in_table: bool) -> int:
    if in_table:
        return 12
    style = paragraph.style.name if paragraph.style else ""
    if style.startswith("Heading 1"):
        return 16
    return 14


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph, False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph, True


def split_segments(text: str):
    pos = 0
    for match in CODE_RE.finditer(text):
        if match.start() > pos:
            yield text[pos : match.start()], False
        yield match.group(0), True
        pos = match.end()
    if pos < len(text):
        yield text[pos:], False


def rewrite_paragraph(paragraph, in_table: bool) -> int:
    text = paragraph.text
    if not text:
        return 0
    if not CODE_RE.search(text):
        changed = 0
        size = normal_size(paragraph, in_table)
        for run in paragraph.runs:
            if run.text:
                set_font(run, "Times New Roman", size)
                changed += 1
        return changed

    # Rebuild only paragraphs that contain project-code tokens, preserving
    # paragraph-level style while avoiding inherited Courier on prose.
    template = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs:
        run.text = ""

    size = normal_size(paragraph, in_table)
    used = False
    current = template
    changed = 0
    for segment, is_code in split_segments(text):
        if not segment:
            continue
        if used:
            new_r = deepcopy(template._r)
            template._r.addnext(new_r)
            template = Run(new_r, paragraph)
            current = template
        current.text = segment
        if is_code:
            set_font(current, "Courier New", 11)
        else:
            set_font(current, "Times New Roman", size)
        used = True
        changed += 1
    return changed


def fix(path: Path) -> int:
    doc = Document(path)
    changed = 0
    for paragraph, in_table in iter_paragraphs(doc):
        has_courier = any(
            run.text
            and (
                run.font.name == "Courier New"
                or (run.style and run.style.font and run.style.font.name == "Courier New")
            )
            for run in paragraph.runs
        )
        if has_courier or CODE_RE.search(paragraph.text):
            changed += rewrite_paragraph(paragraph, in_table)
    if changed:
        doc.save(path)
    return changed


def main() -> None:
    for target in TARGETS:
        if target.exists():
            print(target, fix(target))


if __name__ == "__main__":
    main()
