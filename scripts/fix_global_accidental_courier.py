from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TARGETS = [
    ROOT / "Vương Văn Duy_Báo cáo.docx",
    ROOT / "Bao_Cao" / "02_Chuong_2_ban_dinh_dang_theo_mau.docx",
]

INLINE_CODE_RE = re.compile(
    r"Next\.js|React|Vercel|SSR|SSG|WebSocket|Server-Sent Events|SSE|HTTP|HTTPS|REST|JWT"
)


def set_font(run, name: str, size: int) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), name)


def rewrite(paragraph, normal_size: int) -> None:
    text = paragraph.text.replace("Vercel Vercel.", "Vercel.")
    for run in paragraph.runs:
        run.text = ""

    first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    pos = 0
    used = False
    current = first
    for match in INLINE_CODE_RE.finditer(text):
        if match.start() > pos:
            if used:
                current = paragraph.add_run()
            current.text = text[pos : match.start()]
            set_font(current, "Times New Roman", normal_size)
            used = True
        if used:
            current = paragraph.add_run()
        current.text = match.group(0)
        set_font(current, "Courier New", 11)
        used = True
        pos = match.end()
    if pos < len(text):
        if used:
            current = paragraph.add_run()
        current.text = text[pos:]
        set_font(current, "Times New Roman", normal_size)


def iter_table_paragraphs(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def fix(path: Path) -> int:
    doc = Document(path)
    changed = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if (
            text.startswith("Next.js là framework")
            or text.startswith("Next.js hỗ trợ")
            or text.startswith("Next.js dễ dàng")
            or "Next.js hỗ trợ Server-Side Rendering" in text
            or "Next.js dễ dàng tích hợp với WebSocket" in text
        ):
            rewrite(paragraph, 14)
            changed += 1

    for paragraph in iter_table_paragraphs(doc):
        if paragraph.text.strip() == "HTTP nội bộ":
            rewrite(paragraph, 12)
            changed += 1

    if changed:
        doc.save(path)
    return changed


def main() -> None:
    for target in TARGETS:
        if target.exists():
            print(target, fix(target))


if __name__ == "__main__":
    main()
