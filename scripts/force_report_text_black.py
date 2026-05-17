from pathlib import Path

from docx import Document
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
TARGET = ROOT / "Vương Văn Duy_Báo cáo.docx"
BLACK = RGBColor(0, 0, 0)


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def main() -> None:
    doc = Document(TARGET)
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            try:
                style.font.color.rgb = BLACK
            except Exception:
                pass
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
