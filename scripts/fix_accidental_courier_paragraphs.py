from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TARGETS = [
    ROOT / "Vương Văn Duy_Báo cáo.docx",
    ROOT / "Bao_Cao" / "02_Chuong_2_ban_dinh_dang_theo_mau.docx",
]

CODE_RE = re.compile(r"\b(?:SSE|WebSocket|HTTP|HTTPS|REST|JWT)\b")


def set_font(run, name: str, size: int) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), name)


def rewrite_with_inline_code(paragraph) -> None:
    text = paragraph.text
    for run in paragraph.runs:
        run.text = ""

    first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    pos = 0
    current = first
    used_first = False
    for match in CODE_RE.finditer(text):
        if match.start() > pos:
            if used_first:
                current = paragraph.add_run()
            current.text = text[pos : match.start()]
            set_font(current, "Times New Roman", 14)
            used_first = True
        if used_first:
            current = paragraph.add_run()
        current.text = match.group(0)
        set_font(current, "Courier New", 11)
        used_first = True
        pos = match.end()

    if pos < len(text):
        if used_first:
            current = paragraph.add_run()
        current.text = text[pos:]
        set_font(current, "Times New Roman", 14)


def fix(path: Path) -> int:
    doc = Document(path)
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("SSE phù hợp với luồng chat"):
            rewrite_with_inline_code(paragraph)
            changed += 1
    if changed:
        doc.save(path)
    return changed


def main() -> None:
    for target in TARGETS:
        if target.exists():
            print(target, fix(target))


if __name__ == "__main__":
    main()
