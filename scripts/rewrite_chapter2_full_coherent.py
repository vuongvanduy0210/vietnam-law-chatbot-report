from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "Bao_Cao" / "02_Chuong_2_ban_van_phong_mau.docx"
OUT = ROOT / "Bao_Cao" / "02_Chuong_2_ban_viet_lai_day_du_theo_mau.docx"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph, text: str):
    from docx.oxml import OxmlElement

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    new_para.text = text
    return new_para


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    paragraphs = list(doc.paragraphs)

    replacements = {
        1: (
            "Chương này trình bày phương pháp xây dựng hệ thống Vietnam Law Chatbot. Trọng tâm của chương là làm rõ "
            "cách tổ chức kiến trúc, cách xây dựng tác nhân AI, phương pháp truy hồi tri thức pháp luật, luồng tư vấn "
            "có hướng dẫn và quy trình cập nhật văn bản mới từ phía quản trị. Nội dung trong chương được trình bày theo "
            "hướng phương pháp luận: trước hết xác định mô hình tổng quát, sau đó phân tích các luồng xử lý chính và cuối "
            "cùng giải thích cách các lớp ứng dụng hỗ trợ người dùng tương tác với hệ thống."
        ),
        5: (
            "Hệ thống Vietnam Law Chatbot được xây dựng theo kiến trúc Microservices. Thay vì gom toàn bộ chức năng vào "
            "một ứng dụng duy nhất, hệ thống tách các nhóm nghiệp vụ thành những dịch vụ có vai trò riêng. Cách tổ chức "
            "này phù hợp với bài toán chatbot pháp luật vì hệ thống vừa cần xử lý các nghiệp vụ thông thường như xác thực, "
            "quản lý hội thoại, quản trị văn bản, vừa cần thực hiện các tác vụ AI có độ trễ và chi phí tính toán lớn hơn."
        ),
        6: (
            "Lý do đầu tiên là yêu cầu tách riêng phần xử lý nặng. Các bước tạo embedding, truy hồi vector, xếp hạng lại "
            "kết quả và suy luận bằng LLMs thường mất nhiều thời gian hơn các thao tác nghiệp vụ thông thường. Khi đặt "
            "những bước này trong một dịch vụ riêng, các chức năng như đăng nhập, xem lịch sử hội thoại hoặc mở dashboard "
            "quản trị không bị phụ thuộc trực tiếp vào tải xử lý của pipeline AI."
        ),
        7: (
            "Lý do thứ hai là bảo vệ phần lõi tri thức và suy luận. RAG Service chứa quy trình truy hồi, tổng hợp và kiểm "
            "chứng câu trả lời nên không nên được mở trực tiếp cho client. Người dùng cuối chỉ gửi yêu cầu đến Main Service, "
            "sau đó Main Service mới điều phối sang RAG Service khi cần xử lý câu hỏi pháp luật."
        ),
        8: (
            "Lý do thứ ba là khả năng mở rộng và bảo trì. Main Service có nhịp thay đổi chủ yếu theo nghiệp vụ người dùng "
            "và quản trị, trong khi RAG Service có thể cần điều chỉnh thường xuyên hơn về chiến lược truy hồi, ngưỡng tin "
            "cậy, cách kiểm chứng hoặc cách sử dụng LLMs. Việc tách hai phần này giúp hệ thống dễ thử nghiệm và nâng cấp "
            "mà không làm ảnh hưởng toàn bộ ứng dụng."
        ),
        12: (
            "Dựa trên các yêu cầu trên, hệ thống được tổ chức thành bốn lớp. Lớp giao diện phục vụ hai nhóm người dùng "
            "khác nhau; lớp dịch vụ backend chịu trách nhiệm điều phối nghiệp vụ; lớp dữ liệu lưu trữ thông tin lâu dài; "
            "lớp dịch vụ bên ngoài cung cấp các năng lực hỗ trợ như LLMs, tìm kiếm cập nhật và lưu trữ file."
        ),
        24: (
            "Nhìn vào mô hình tổng quan, luồng tương tác của hệ thống đi từ giao diện đến Main Service, sau đó mới được "
            "chuyển tiếp sang RAG Service hoặc các kho dữ liệu phù hợp. Cách đi này giúp hệ thống có một điểm kiểm soát "
            "chung cho xác thực, phân quyền, ghi nhận lịch sử và điều phối tác vụ."
        ),
        25: (
            "Lớp điều phối nghiệp vụ không trực tiếp thay thế các dịch vụ phía sau, mà đóng vai trò trung gian để giảm "
            "độ phức tạp cho frontend. Mobile App và Admin Web không cần biết chi tiết hệ thống đang sử dụng bao nhiêu "
            "cơ sở dữ liệu hoặc bao nhiêu bước AI, mà chỉ cần làm việc với các chức năng nghiệp vụ đã được Main Service "
            "chuẩn hóa."
        ),
        26: (
            "Lớp dữ liệu nền bao gồm dữ liệu người dùng, hội thoại, văn bản pháp luật và vector tri thức. Các dịch vụ bên "
            "ngoài chỉ đóng vai trò bổ sung năng lực cho pipeline, ví dụ hỗ trợ suy luận ngôn ngữ, đối chiếu thông tin cập "
            "nhật hoặc lưu trữ file gốc của văn bản pháp luật."
        ),
        30: (
            "Main Service là cửa ngõ nghiệp vụ của hệ thống. Thành phần này tiếp nhận yêu cầu từ Mobile App và Admin Web, "
            "kiểm tra quyền truy cập, ghi nhận các dữ liệu cần thiết và chuyển yêu cầu đến nhóm xử lý phù hợp. Đối với "
            "người dùng cuối, Main Service quản lý hội thoại, gửi câu hỏi sang RAG Service và lưu lại câu trả lời. Đối với "
            "quản trị viên, Main Service tiếp nhận văn bản mới, tạo tác vụ xử lý và theo dõi tiến trình cập nhật dữ liệu."
        ),
        31: (
            "Vai trò của Main Service vì vậy là chuẩn hóa các thao tác từ nhiều giao diện thành các luồng nghiệp vụ thống "
            "nhất. Một yêu cầu chat và một yêu cầu cập nhật văn bản có bản chất rất khác nhau, nhưng đều cần được xác thực, "
            "ghi nhận trạng thái và điều phối qua các thành phần phía sau."
        ),
        32: (
            "Việc đặt Main Service ở giữa cũng giúp giảm sự phụ thuộc của frontend vào cấu trúc nội bộ. Nếu sau này thay "
            "đổi cách truy hồi, thay đổi kho vector hoặc bổ sung bước kiểm chứng mới, lớp giao diện vẫn có thể giữ nguyên "
            "luồng tương tác chính với người dùng."
        ),
        36: (
            "RAG Service là thành phần quyết định chất lượng trả lời pháp luật. Dịch vụ này không chỉ gọi LLMs để tạo sinh "
            "văn bản, mà tổ chức cả quá trình phân tích câu hỏi, truy hồi căn cứ, đối chiếu nguồn và kiểm chứng nội dung "
            "trước khi trả về kết quả cho Main Service."
        ),
        37: (
            "Trong miền pháp luật, câu trả lời cần dựa trên văn bản có thể kiểm tra lại. Vì vậy, RAG Service được thiết kế "
            "để LLMs làm việc với các nguồn đã truy hồi thay vì trả lời hoàn toàn từ tri thức sẵn có. Đây là điểm khác biệt "
            "quan trọng giữa một chatbot thông thường và hệ thống tư vấn pháp luật có căn cứ."
        ),
        38: (
            "Có thể xem RAG Service là nơi kết hợp ba nhóm kỹ thuật: truy hồi thông tin để tìm nguồn, suy luận ngôn ngữ để "
            "diễn giải nội dung, và kiểm chứng để giảm sai lệch ở đầu ra. Các phần tiếp theo của chương sẽ lần lượt phân "
            "tích cách hệ thống tổ chức ba nhóm kỹ thuật này."
        ),
        43: (
            "Sau khi xác định kiến trúc tổng thể, bước tiếp theo là xác định tác nhân AI hoạt động trong môi trường nào. "
            "Với bài toán tư vấn pháp luật, tác nhân không chỉ sinh câu trả lời mà còn phải biết sử dụng nguồn dữ liệu, "
            "đặt giới hạn phạm vi, chủ động tra cứu khi thiếu căn cứ và từ chối những yêu cầu không phù hợp."
        ),
        51: "Từ các nguồn dữ liệu và phạm vi đã xác định, tác nhân AI trong hệ thống có các nhóm hành động chính sau:",
        52: "1. Tra cứu tri thức nội bộ để tìm văn bản, điều luật và đoạn nội dung liên quan đến câu hỏi.",
        53: "2. Đối chiếu nguồn cập nhật bên ngoài khi câu hỏi có khả năng phụ thuộc vào văn bản mới hoặc tình trạng hiệu lực.",
        54: "3. Tổng hợp câu trả lời dựa trên các căn cứ đã thu thập và diễn giải theo ngôn ngữ dễ hiểu.",
        55: "4. Từ chối hoặc giới hạn câu trả lời khi câu hỏi nằm ngoài phạm vi tư vấn pháp luật.",
        56: "5. Tra cứu bổ sung khi bằng chứng hiện có chưa đủ để đưa ra câu trả lời đáng tin cậy.",
        59: (
            "Từ mô hình tác nhân ở trên, hệ thống xây dựng luồng RAG có kiểm chứng cho một lượt hỏi đáp. Luồng này không "
            "để LLMs trả lời ngay sau khi nhận câu hỏi, mà chia quá trình xử lý thành các bước có trách nhiệm rõ ràng. "
            "Nhờ đó, hệ thống có thể kiểm soát đầu vào, xác định hướng tìm kiếm, thu thập bằng chứng và kiểm tra lại câu "
            "trả lời trước khi hiển thị cho người dùng."
        ),
        60: (
            "Điểm cốt lõi của phương pháp này là tách phần suy luận khỏi phần kiểm chứng. LLMs có thể diễn giải và tổng "
            "hợp tốt, nhưng trong lĩnh vực pháp luật, kết luận cuối cùng cần được ràng buộc bởi nguồn. Vì vậy, mỗi bước "
            "trong pipeline đều phục vụ mục tiêu đưa câu hỏi của người dùng về gần hơn với các căn cứ pháp lý phù hợp."
        ),
        62: (
            "Trong một lượt hỏi đáp, hệ thống cần duy trì các thông tin trung gian như câu hỏi hiện tại, ngữ cảnh hội thoại, "
            "kết quả phân tích truy vấn, các nguồn đã truy hồi và trạng thái hợp lệ của câu hỏi. Những thông tin này giúp "
            "pipeline biết khi nào cần tìm thêm bằng chứng và khi nào có thể chuyển sang kiểm chứng."
        ),
        66: (
            "Luồng tổng quan gồm bốn bước chính. Đầu tiên, hệ thống kiểm tra xem câu hỏi có thuộc phạm vi pháp luật hay "
            "không. Tiếp theo, câu hỏi hợp lệ được phân tích để tạo hướng tìm kiếm rõ hơn. Sau đó, tác nhân AI thu thập "
            "bằng chứng từ các nguồn phù hợp. Cuối cùng, câu trả lời dự thảo được kiểm chứng trước khi trả về người dùng."
        ),
        67: (
            "So với RAG tuyến tính, luồng này cho phép tác nhân linh hoạt hơn trong quá trình tìm nguồn. Nếu bằng chứng "
            "chưa đủ, hệ thống có thể tạo truy vấn bổ sung hoặc đối chiếu thêm nguồn cập nhật. Điều này phù hợp với câu "
            "hỏi pháp luật vì một câu hỏi ngắn có thể phụ thuộc vào loại đối tượng, tình huống áp dụng và thời điểm hiệu lực."
        ),
        68: (
            "Như vậy, Hình 2.4 không chỉ thể hiện thứ tự xử lý mà còn thể hiện hai lớp kiểm soát chất lượng: kiểm soát "
            "đầu vào để loại bỏ yêu cầu ngoài phạm vi và kiểm chứng đầu ra để hạn chế nội dung thiếu căn cứ."
        ),
        70: (
            "Bước kiểm soát đầu vào xác định câu hỏi có thuộc phạm vi tư vấn pháp luật hay không. Nếu người dùng đặt câu "
            "hỏi lạc đề, hệ thống phản hồi từ chối lịch sự và không chạy các bước truy hồi phía sau. Nhờ vậy, tài nguyên "
            "xử lý được tập trung cho đúng mục tiêu của hệ thống."
        ),
        72: (
            "Ví dụ, với một yêu cầu lập trình không liên quan đến pháp luật, hệ thống có thể dừng ngay ở bước này. Đây là "
            "một kiểm soát cần thiết vì chatbot có thể nhận nhiều loại câu hỏi khác nhau, trong khi phạm vi đồ án chỉ tập "
            "trung vào tư vấn và tra cứu pháp luật Việt Nam."
        ),
        74: (
            "Khi câu hỏi hợp lệ, hệ thống chuyển sang bước phân tích truy vấn. Mục tiêu của bước này là tách câu hỏi tự "
            "nhiên của người dùng thành các thông tin có ích cho truy hồi: chủ đề pháp lý, thuật ngữ chính, phạm vi tìm "
            "kiếm và các dấu hiệu cần đối chiếu."
        ),
        75: (
            "Phân tích truy vấn cũng quyết định cách sử dụng từng nguồn. Truy vấn gửi vào kho nội bộ cần cô đọng và gần "
            "với thuật ngữ pháp lý; truy vấn dành cho nguồn cập nhật có thể cần thêm dấu hiệu về văn bản mới hoặc năm hiệu "
            "lực. Nhờ vậy, các bước sau không phải xử lý trực tiếp một câu hỏi còn quá đời thường hoặc thiếu trọng tâm."
        ),
        80: (
            "Ở bước suy luận, tác nhân AI vận hành theo vòng Re-Act. Trước hết, tác nhân xác định thông tin còn thiếu; "
            "sau đó chọn hành động tra cứu phù hợp; cuối cùng quan sát kết quả nhận được để quyết định tiếp tục tìm kiếm "
            "hay chuyển sang tổng hợp câu trả lời."
        ),
        81: (
            "Cơ chế này giúp hệ thống không tạo câu trả lời quá sớm. Khi căn cứ chưa đủ, tác nhân có thể quay lại tìm thêm "
            "nguồn. Khi căn cứ đã đủ, pipeline mới chuyển sang tổng hợp. Điều này làm cho câu trả lời phụ thuộc nhiều hơn "
            "vào dữ liệu truy hồi và ít phụ thuộc hơn vào khả năng suy đoán của LLMs."
        ),
        82: (
            "Trong miền pháp luật, một câu hỏi có thể cần nhiều loại bằng chứng: điều luật gốc, điều kiện áp dụng, mức xử "
            "phạt, thời điểm hiệu lực hoặc văn bản thay thế. Vòng Re-Act giúp quá trình thu thập bằng chứng diễn ra theo "
            "từng bước thay vì gom tất cả vào một lần tìm kiếm duy nhất."
        ),
        83: (
            "Để tránh kéo dài không cần thiết, hệ thống vẫn đặt giới hạn số vòng lặp và luôn chuyển sang bước kiểm chứng "
            "trước khi trả lời. Đây là điểm cân bằng giữa chất lượng bằng chứng và độ trễ chấp nhận được của trải nghiệm chat."
        ),
        86: (
            "Bước kiểm chứng đối chiếu câu trả lời dự thảo với các nguồn đã truy hồi. Các nội dung như số điều, mức phạt, "
            "điều kiện áp dụng hoặc thời hạn phải xuất hiện trong bằng chứng. Nếu phát hiện phần diễn giải vượt quá nguồn, "
            "hệ thống cần sửa hoặc loại bỏ trước khi trả về kết quả cuối cùng."
        ),
        91: (
            "Tác nhân AI chỉ có thể trả lời chính xác khi có nguồn tri thức đáng tin cậy. Trong hệ thống này, nguồn tri "
            "thức được chia thành hai nhóm: kho pháp luật nội bộ đã được chuẩn hóa và nguồn cập nhật bên ngoài dùng để "
            "đối chiếu tính hiện hành của quy định."
        ),
        93: (
            "Kho nội bộ cung cấp nguyên văn điều luật, metadata và liên kết đến văn bản gốc. Đây là nguồn chính để tạo "
            "câu trả lời có cấu trúc và có trích dẫn. Vì dữ liệu đã được chuẩn hóa trước, hệ thống có thể truy hồi theo "
            "điều khoản, chủ đề, từ khóa và ngữ nghĩa của câu hỏi."
        ),
        94: (
            "Truy vấn gửi vào kho nội bộ không nhất thiết là nguyên văn câu hỏi ban đầu. Sau bước phân tích, hệ thống có "
            "thể dùng phiên bản truy vấn ngắn gọn hơn, tập trung vào ý định pháp lý chính. Cách làm này giúp giảm nhiễu "
            "và tăng khả năng tìm đúng nhóm văn bản liên quan."
        ),
        99: (
            "Nguồn bên ngoài không thay thế kho dữ liệu nội bộ mà đóng vai trò đối chiếu. Khi câu hỏi liên quan đến quy "
            "định mới, mức phạt mới hoặc văn bản có khả năng đã thay đổi, hệ thống có thể kiểm tra thêm nguồn cập nhật "
            "trước khi đưa ra kết luận."
        ),
        100: (
            "Khi hai nhóm nguồn cùng được sử dụng, hệ thống ưu tiên nguyên tắc an toàn: chỉ khẳng định nội dung pháp lý "
            "khi có căn cứ rõ ràng. Nếu nguồn nội bộ và nguồn cập nhật có dấu hiệu khác nhau, hệ thống cần ưu tiên kiểm "
            "tra tính hiệu lực trước khi tổng hợp câu trả lời."
        ),
        101: (
            "Như vậy, kho nội bộ đảm nhiệm vai trò cung cấp tri thức đã chuẩn hóa, còn nguồn bên ngoài giúp giảm rủi ro "
            "sử dụng văn bản không còn phù hợp. Hai nguồn này bổ sung cho nhau thay vì trùng lặp chức năng."
        ),
        105: (
            "Sau khi xác định nguồn tri thức, vấn đề tiếp theo là cách chọn đúng đoạn văn bản từ một kho dữ liệu lớn. "
            "Nếu chỉ dựa vào LLMs, hệ thống không thể đọc toàn bộ kho pháp luật trong mỗi câu hỏi. Nếu chỉ tìm kiếm từ "
            "khóa, hệ thống dễ bỏ sót những đoạn có ý nghĩa tương đồng nhưng cách diễn đạt khác. Vì vậy, đồ án sử dụng "
            "pipeline truy xuất hai giai đoạn."
        ),
        107: (
            "Giai đoạn đầu có nhiệm vụ lọc rộng bằng tìm kiếm vector để lấy ra tập ứng viên có khả năng liên quan. Giai "
            "đoạn sau đánh giá sâu hơn từng ứng viên để chọn các đoạn văn bản phù hợp nhất với câu hỏi. Cách làm này vừa "
            "giữ được tốc độ xử lý, vừa tăng chất lượng căn cứ đưa vào LLMs."
        ),
        108: (
            "Trong thiết kế này, LLMs không quyết định toàn bộ việc tìm nguồn. Kho vector và bước xếp hạng lại chịu trách "
            "nhiệm chọn bằng chứng; tác nhân AI chịu trách nhiệm sử dụng bằng chứng đó để suy luận và trình bày câu trả lời. "
            "Sự phân vai này giúp pipeline rõ ràng hơn và tránh để LLMs tự suy đoán nguồn."
        ),
        121: (
            "2.5.4. Ưu tiên văn bản mới"
        ),
        122: (
            "Đặc thù của dữ liệu pháp luật là nhiều nghị định, thông tư có thể được sửa đổi hoặc thay thế bởi văn bản mới. "
            "Nếu chỉ dựa vào độ tương đồng ngữ nghĩa, hai văn bản cùng chủ đề có thể có điểm số gần nhau dù một văn bản đã "
            "không còn phù hợp. Vì vậy, hệ thống bổ sung cơ chế ưu tiên văn bản mới trước khi chọn căn cứ cuối cùng."
        ),
        128: (
            "Ngoài độ liên quan, dữ liệu pháp luật còn có yếu tố thời gian. Một văn bản cũ có thể diễn đạt gần với câu "
            "hỏi hơn, nhưng văn bản mới hơn mới là căn cứ đang có hiệu lực. Do đó, hệ thống cần xét thêm ngày ban hành, "
            "ngày hiệu lực và quan hệ sửa đổi, thay thế giữa các văn bản."
        ),
        130: (
            "Quy trình xử lý xung đột gồm ba bước: nhận diện các nguồn có khả năng cùng chủ đề, đối chiếu thông tin hiệu "
            "lực và gắn trạng thái cho nguồn hiện hành hoặc nguồn cũ. Nguồn hiện hành được ưu tiên làm căn cứ chính; nguồn "
            "cũ chỉ nên dùng để giải thích sự thay đổi khi cần."
        ),
        131: (
            "Nhờ bổ sung tiêu chí thời gian, câu trả lời của hệ thống phù hợp hơn với bản chất của bài toán pháp luật. "
            "Người dùng không chỉ cần câu trả lời có vẻ liên quan, mà cần câu trả lời dựa trên quy định đang có giá trị áp dụng."
        ),
        139: (
            "Luồng chat tự do phù hợp khi câu hỏi của người dùng đã có đủ ngữ cảnh. Tuy nhiên, trong thực tế, nhiều câu "
            "hỏi pháp luật được đặt rất ngắn, ví dụ \"vượt đèn đỏ phạt bao nhiêu\". Câu hỏi này chưa cho biết phương tiện, "
            "tình huống cụ thể hoặc các yếu tố có thể làm thay đổi mức xử phạt. Vì vậy, hệ thống bổ sung luồng tư vấn có "
            "hướng dẫn để chủ động hỏi lại trước khi trả lời."
        ),
        140: (
            "Luồng tư vấn có hướng dẫn tách quá trình tư vấn thành hai bước. Bước đầu thu thập thêm dữ kiện còn thiếu bằng "
            "các câu hỏi làm rõ. Bước sau sử dụng câu hỏi ban đầu cùng các lựa chọn của người dùng để chạy truy hồi, suy "
            "luận và kiểm chứng. Cách tổ chức này giúp câu trả lời cuối cùng gắn với tình huống cụ thể hơn."
        ),
        142: (
            "Trong thực tế tư vấn pháp luật, thiếu ngữ cảnh là nguyên nhân phổ biến khiến câu trả lời trở nên quá rộng "
            "hoặc không đúng trường hợp của người hỏi. Chẳng hạn, cùng là hành vi vượt đèn đỏ nhưng mức phạt có thể khác "
            "nhau giữa xe máy và ô tô. Nếu hệ thống trả lời ngay, nó buộc phải bao quát nhiều trường hợp hoặc tự giả định "
            "thông tin còn thiếu."
        ),
        143: "Nếu chỉ dùng luồng hỏi đáp một bước, hệ thống thường gặp hai khó khăn:",
        144: "1. Câu trả lời dễ dài và phân tán vì phải bao quát nhiều trường hợp có thể xảy ra.",
        145: "2. Hệ thống có thể suy đoán dữ kiện còn thiếu, làm giảm độ chính xác của tư vấn.",
        146: (
            "Luồng tư vấn có hướng dẫn giải quyết vấn đề này bằng cách biến câu hỏi mơ hồ thành một tình huống rõ ràng hơn "
            "trước khi chạy pipeline trả lời."
        ),
        147: (
            "Về mặt kỹ thuật, bước làm rõ giúp truy vấn trở nên cụ thể hơn, kết quả truy hồi tập trung hơn và bước kiểm chứng "
            "có phạm vi đối chiếu rõ ràng hơn. Đây là lý do chức năng này được tách thành một luồng riêng thay vì chỉ là một "
            "câu hỏi phụ trong chat tự do."
        ),
        149: (
            "Luồng tư vấn có hướng dẫn được thiết kế theo hướng không lưu trạng thái phiên trên máy chủ. Sau khi hệ thống "
            "sinh câu hỏi làm rõ, các lựa chọn của người dùng sẽ được gửi kèm trong yêu cầu trả lời ở bước tiếp theo."
        ),
        155: (
            "Ở bước đầu tiên, hệ thống phân tích câu hỏi ban đầu để xác định những dữ kiện còn thiếu. Đầu ra của bước này "
            "là danh sách câu hỏi làm rõ, thường được biểu diễn dưới dạng các lựa chọn ngắn để người dùng thao tác nhanh trên "
            "giao diện."
        ),
        156: (
            "Điểm cần lưu ý là bước này chưa tạo câu trả lời pháp lý cuối cùng. Nó chỉ chuẩn bị ngữ cảnh cho bước sau. Nhờ "
            "tách riêng như vậy, hệ thống có thể phản hồi nhanh ở bước làm rõ nhưng vẫn giữ yêu cầu truy hồi và kiểm chứng "
            "đối với câu trả lời chính thức."
        ),
        158: (
            "Ở bước thứ hai, hệ thống kết hợp câu hỏi ban đầu với các lựa chọn làm rõ của người dùng. Yêu cầu sau khi được "
            "bổ sung ngữ cảnh sẽ cụ thể hơn, từ đó giúp pipeline truy hồi đúng căn cứ và hạn chế việc trả lời quá rộng."
        ),
        160: (
            "Ví dụ, nếu câu hỏi ban đầu là \"Vượt đèn đỏ phạt bao nhiêu?\" và người dùng chọn phương tiện là xe máy, hệ "
            "thống sẽ ưu tiên truy hồi các quy định liên quan đến xe mô tô, xe gắn máy thay vì trộn lẫn với ô tô hoặc các "
            "phương tiện khác."
        ),
        161: (
            "Sau khi có ngữ cảnh rõ ràng, luồng trả lời vẫn tuân theo nguyên tắc của RAG có kiểm chứng: tìm nguồn, tổng hợp "
            "trên căn cứ và kiểm tra lại trước khi hiển thị cho người dùng."
        ),
        164: (
            "Câu trả lời trong luồng tư vấn có hướng dẫn được truyền về giao diện theo từng trạng thái xử lý. Người dùng "
            "không cần chờ trong im lặng, mà có thể thấy hệ thống đang chuẩn bị truy hồi, tra cứu nguồn, tổng hợp và kiểm tra "
            "độ chính xác."
        ),
        165: (
            "Cơ chế này giữ trải nghiệm tương tự luồng chat thời gian thực, nhưng dữ liệu đầu vào đã rõ hơn nhờ bước làm rõ "
            "trước đó. Vì vậy, chức năng tư vấn có hướng dẫn vừa cải thiện trải nghiệm sử dụng, vừa hỗ trợ chất lượng câu trả lời."
        ),
        168: (
            "Các trạng thái hiển thị cho người dùng được giữ ở mức dễ hiểu, không đi sâu vào chi tiết kỹ thuật nội bộ. Mục "
            "tiêu là cho người dùng biết hệ thống đang xử lý có kiểm soát, đồng thời chỉ hiển thị nội dung pháp lý cuối cùng "
            "khi pipeline đã hoàn tất."
        ),
        173: (
            "Chất lượng tư vấn phụ thuộc trực tiếp vào chất lượng kho dữ liệu pháp luật. Vì vậy, ngoài luồng hỏi đáp, hệ "
            "thống cần một quy trình cập nhật văn bản mới từ phía quản trị. Quy trình này phải đảm bảo văn bản được kiểm tra, "
            "trích xuất, chuẩn hóa và đồng bộ vào kho truy hồi trước khi được sử dụng để trả lời người dùng."
        ),
        174: (
            "Về tổng thể, luồng cập nhật tri thức bắt đầu từ thao tác tải lên file PDF của quản trị viên. Sau đó hệ thống "
            "thực hiện tiền kiểm, trích xuất nội dung, chuẩn hóa thành cấu trúc điều luật, kiểm tra hợp lệ, lưu dữ liệu gốc "
            "và tạo vector phục vụ tìm kiếm. Kết quả cuối cùng không chỉ là một file được lưu, mà là một văn bản pháp luật "
            "sẵn sàng tham gia vào pipeline hỏi đáp."
        ),
        176: (
            "Bước tiền kiểm được đặt ở đầu quy trình để giảm rủi ro đưa dữ liệu sai vào hệ thống. Ở bước này, hệ thống có "
            "thể phát hiện tài liệu không hợp lệ, văn bản trùng lặp hoặc file không phù hợp trước khi chạy các tác vụ tốn "
            "nhiều tài nguyên như nhận dạng nội dung và tạo vector."
        ),
        177: (
            "Sau tiền kiểm, dữ liệu được xử lý theo hai hướng. Một hướng tạo tri thức gốc để quản trị viên có thể xem, tìm "
            "kiếm và đối chiếu lại văn bản. Hướng còn lại tạo các vector phục vụ truy hồi ngữ nghĩa trong quá trình hỏi đáp. "
            "Hai hướng này phải đồng bộ với nhau để tránh tình trạng giao diện quản trị đã có văn bản nhưng chatbot chưa thể tìm được."
        ),
        178: (
            "Như vậy, cập nhật văn bản pháp luật không chỉ là chức năng tải file. Đây là một pipeline dữ liệu hoàn chỉnh, "
            "có kiểm tra đầu vào, xử lý nội dung, ghi dữ liệu và kiểm soát trạng thái thành công hoặc thất bại."
        ),
        179: (
            "Tiền kiểm cũng giúp tiết kiệm chi phí vận hành. Nếu quản trị viên tải nhầm một văn bản đã tồn tại, hệ thống có "
            "thể dừng sớm thay vì tiếp tục trích xuất, chuẩn hóa và tạo vector cho một dữ liệu trùng lặp."
        ),
        181: (
            "Pipeline xử lý PDF được chia thành hai nhiệm vụ khác nhau: trích xuất nội dung chữ từ tài liệu và chuẩn hóa "
            "nội dung đó thành cấu trúc pháp lý. Việc tách hai nhiệm vụ giúp hệ thống tận dụng đúng vai trò của từng thành phần."
        ),
        182: (
            "Ở nhiệm vụ thứ nhất, hệ thống cố gắng lấy lớp chữ có sẵn trong PDF. Nếu văn bản là bản scan hoặc chất lượng "
            "thấp, hệ thống sử dụng OCR để nhận dạng ký tự. Kết quả của bước này là phần văn bản thô phục vụ cho bước chuẩn hóa."
        ),
        183: (
            "Ở nhiệm vụ thứ hai, LLMs được dùng để hiểu cấu trúc của văn bản pháp luật, chẳng hạn tên văn bản, chương, mục, "
            "điều, khoản và các metadata liên quan. LLMs không thay thế bước trích xuất chữ, mà tập trung vào việc chuyển nội "
            "dung đã đọc được thành dữ liệu có cấu trúc."
        ),
        184: "Cách tách này có ba lợi ích chính:",
        185: "• Giảm chi phí xử lý vì phần lớn PDF pháp luật đã có lớp chữ để trích xuất trực tiếp.",
        186: "• Cho phép kiểm soát tốt hơn chất lượng nhận dạng tiếng Việt ở những tài liệu scan.",
        187: "• Giúp LLMs tập trung vào nhiệm vụ hiểu cấu trúc pháp lý thay vì phải xử lý toàn bộ file từ đầu.",
        188: (
            "Kết quả sau chuẩn hóa bao gồm thông tin nhận diện văn bản, ngày hiệu lực, danh sách điều khoản và metadata "
            "phục vụ tìm kiếm, trích dẫn và quản trị."
        ),
        190: (
            "Nếu bước trích xuất hoặc chuẩn hóa thất bại, hệ thống mới chuyển sang đường xử lý dự phòng để đọc trực tiếp "
            "tài liệu bằng LLMs. Cách thiết kế này giúp ưu tiên đường xử lý tiết kiệm hơn trong trường hợp thông thường, "
            "nhưng vẫn có phương án thay thế với những file khó đọc."
        ),
        193: (
            "Một số tác vụ trong quy trình cập nhật văn bản có thể chạy song song. Chẳng hạn, việc lưu file PDF gốc lên "
            "kho lưu trữ và việc trích xuất nội dung từ PDF không phụ thuộc trực tiếp vào nhau. Tổ chức song song các bước "
            "này giúp giảm thời gian chờ của quản trị viên."
        ),
        194: (
            "File gốc được lưu lại để mỗi điều luật có thể tham chiếu về nguồn khi cần kiểm tra. Đây là yêu cầu quan trọng "
            "đối với hệ thống pháp luật vì người dùng hoặc quản trị viên có thể cần đối chiếu câu trả lời với văn bản ban đầu."
        ),
        195: (
            "Trong khi đó, phần trích xuất nội dung tạo dữ liệu đầu vào cho bước chuẩn hóa cấu trúc. Khi hai kết quả này "
            "đã có đủ, hệ thống kết hợp chúng thành bản ghi văn bản hoàn chỉnh."
        ),
        197: (
            "Nhờ xử lý song song, hệ thống rút ngắn thời gian cập nhật so với cách làm tuần tự. Khi trích xuất thành công, "
            "hệ thống không cần gửi toàn bộ PDF sang LLMs; khi trích xuất thất bại, đường dự phòng vẫn giúp tác vụ có cơ hội hoàn tất."
        ),
        198: (
            "Về mặt vận hành, các tác vụ tốn CPU được tách khỏi luồng xử lý chính để không ảnh hưởng đến các chức năng đang "
            "phục vụ người dùng như chat hoặc dashboard."
        ),
        200: (
            "Sau khi nội dung được chuẩn hóa, hệ thống phải ghi dữ liệu sang nhiều kho khác nhau. Kho văn bản phục vụ quản "
            "trị và hiển thị chi tiết nguồn; kho vector phục vụ truy hồi ngữ nghĩa. Do các kho này không nằm trong một giao "
            "dịch duy nhất, hệ thống cần cơ chế hoàn tác khi có lỗi ở giữa quy trình."
        ),
        203: (
            "Nếu chỉ một phần dữ liệu được ghi thành công, hệ thống có thể rơi vào trạng thái không nhất quán. Ví dụ, quản "
            "trị viên thấy văn bản đã tồn tại nhưng tác nhân AI chưa tìm được văn bản đó trong kho vector. Vì vậy, một lượt "
            "cập nhật chỉ được xem là thành công khi các kho liên quan đều đã đồng bộ."
        ),
        204: (
            "Khi một bước thất bại, tác vụ chuyển sang trạng thái lỗi và những phần dữ liệu đã ghi trước đó được thu hồi. "
            "Nguyên tắc này giúp kho tri thức luôn phản ánh trạng thái nhất quán giữa phần quản trị và phần truy hồi."
        ),
        208: (
            "Trước khi lưu vào kho vector, mỗi điều luật được chia thành các đoạn có kích thước phù hợp. Việc chia nhỏ giúp "
            "hệ thống tìm kiếm chính xác hơn vì mỗi vector đại diện cho một đơn vị nội dung vừa đủ, không quá rộng và cũng "
            "không quá rời rạc."
        ),
        209: "• Kích thước tối đa: khoảng 1.000 từ mỗi đoạn, nhằm giữ đủ ngữ cảnh nhưng không làm đoạn quá dài.",
        210: "• Phần chồng lấn: khoảng 150 từ giữa hai đoạn liền kề để hạn chế mất ngữ cảnh ở ranh giới.",
        211: (
            "Mỗi đoạn sau khi chia nhỏ vẫn giữ thông tin nhận diện văn bản và vị trí điều khoản gốc. Nhờ vậy, khi hệ thống "
            "trích dẫn một đoạn trong câu trả lời, người dùng có thể truy ngược về văn bản nguồn."
        ),
        212: (
            "Với cách chia này, phần lớn điều luật ngắn được giữ nguyên, chỉ những điều dài mới cần tách thành nhiều đoạn. "
            "Điều đó giúp kho vector đủ chi tiết cho tìm kiếm nhưng không tăng số lượng đoạn một cách không cần thiết."
        ),
        213: (
            "Nếu đoạn quá ngắn, một điều khoản có thể bị tách khỏi điều kiện áp dụng; nếu đoạn quá dài, vector đại diện sẽ "
            "kém đặc trưng. Vì vậy, kích thước đoạn và phần chồng lấn được chọn để cân bằng giữa độ đầy đủ của ngữ cảnh và "
            "độ chính xác của truy hồi."
        ),
        218: (
            "Bên cạnh chất lượng câu trả lời, trải nghiệm chat cũng cần cho người dùng biết hệ thống đang xử lý đến đâu. "
            "Trong luồng chat thời gian thực, câu hỏi được ghi nhận vào hội thoại, pipeline RAG được chạy ở phía sau, các "
            "trạng thái xử lý được gửi dần về giao diện và câu trả lời cuối cùng được hiển thị kèm nguồn tham chiếu."
        ),
        219: (
            "Thiết kế này tách rõ tiến trình xử lý trung gian và nội dung pháp lý cuối cùng. Các trạng thái như đang truy "
            "hồi hoặc đang kiểm chứng chỉ giúp người dùng theo dõi tiến độ; còn câu trả lời chính thức chỉ xuất hiện sau khi "
            "pipeline hoàn tất."
        ),
        221: (
            "Luồng chat tuân theo nguyên tắc lưu dữ liệu quan trọng trước khi hiển thị kết quả cuối cùng. Câu hỏi của người "
            "dùng được ghi nhận vào lịch sử hội thoại, sau đó hệ thống mới xử lý bằng pipeline RAG. Khi có câu trả lời, kết "
            "quả tiếp tục được lưu để người dùng có thể xem lại về sau."
        ),
        222: (
            "Nguyên tắc này phù hợp với hệ thống tư vấn pháp luật vì lịch sử hội thoại không chỉ phục vụ trải nghiệm, mà còn "
            "giúp người dùng theo dõi lại bối cảnh tư vấn và các nguồn đã được sử dụng."
        ),
        225: (
            "Các cập nhật trong quá trình streaming chỉ biểu diễn trạng thái xử lý, không phải nội dung pháp lý cuối cùng. "
            "Nhờ đó, giao diện tránh cảm giác chờ thụ động nhưng vẫn giữ nguyên tắc chỉ hiển thị kết luận khi đã qua truy "
            "hồi và kiểm chứng."
        ),
        227: (
            "Lịch sử hội thoại được dùng để hiểu các câu hỏi nối tiếp, ví dụ khi người dùng hỏi tiếp sau một câu trả lời "
            "trước đó. Tuy nhiên, lịch sử không được xem là nguồn pháp lý. Mỗi câu trả lời vẫn cần dựa trên căn cứ được "
            "truy hồi trong lượt xử lý hiện tại."
        ),
        229: (
            "Câu hỏi gợi ý được sinh sau câu trả lời chính để hỗ trợ người dùng tiếp tục mạch tư vấn. Chức năng này không "
            "làm chậm luồng trả lời chính vì nó chỉ đóng vai trò bổ trợ trải nghiệm."
        ),
        232: (
            "Sau khi thiết kế phần lõi backend và pipeline RAG, hệ thống cần lớp ứng dụng để người dùng cuối và quản trị "
            "viên có thể thao tác với các chức năng tương ứng. Vì mục tiêu sử dụng khác nhau, đồ án tách lớp frontend thành "
            "Mobile App cho người dùng cuối và Admin Web cho quản trị viên."
        ),
        234: (
            "Mobile App được xây dựng bằng Kotlin Multiplatform kết hợp Compose Multiplatform. Lựa chọn này phù hợp với "
            "định hướng phát triển phần mềm di động vì hệ thống cần chạy trên cả Android và iOS, đồng thời vẫn muốn chia "
            "sẻ các phần logic chung như mô hình dữ liệu, gọi API, xử lý dòng sự kiện và quản lý trạng thái."
        ),
        235: (
            "Thay vì phát triển hai ứng dụng độc lập hoàn toàn, Kotlin Multiplatform cho phép chia sẻ phần logic nghiệp vụ "
            "giữa hai nền tảng:"
        ),
        236: (
            "Phần dùng chung bao gồm mô hình dữ liệu, lớp truy cập API, ánh xạ dữ liệu, xử lý stream và quản lý trạng thái. "
            "Nhờ đó, luồng chat và luồng tư vấn có hướng dẫn có thể hoạt động nhất quán trên nhiều thiết bị."
        ),
        237: (
            "Các phần phụ thuộc đặc thù nền tảng như giao diện native, quyền hệ thống hoặc vòng đời ứng dụng vẫn được tách "
            "riêng. Cách làm này giữ được trải nghiệm phù hợp cho từng nền tảng nhưng không làm lặp lại toàn bộ logic nghiệp vụ."
        ),
        240: (
            "Mobile App sử dụng mô hình MVI để tổ chức trạng thái giao diện. Mô hình này tách rõ dữ liệu đang hiển thị, hành "
            "động của người dùng và các phản hồi một lần của hệ thống."
        ),
        246: (
            "Cách tổ chức này phù hợp với ứng dụng chat AI vì giao diện thay đổi liên tục: người dùng gửi tin nhắn, hệ thống "
            "nhận trạng thái xử lý, câu trả lời được cập nhật và câu hỏi gợi ý có thể xuất hiện sau đó."
        ),
        248: (
            "Với MVI, các sự kiện streaming từ server được chuyển thành thay đổi trạng thái rõ ràng. Màn hình chỉ cần hiển "
            "thị theo trạng thái hiện tại, nhờ đó tránh trộn logic xử lý dữ liệu trực tiếp vào phần giao diện."
        ),
        249: (
            "Đặc biệt với màn hình chat, tại cùng một thời điểm có thể tồn tại nhiều trạng thái như tin nhắn vừa gửi, tiến "
            "trình đang chạy, câu trả lời chưa hoàn thành hoặc danh sách câu hỏi gợi ý. MVI giúp gom các trạng thái này vào "
            "một mô hình quản lý thống nhất."
        ),
        251: (
            "Một vấn đề thường gặp của chatbot dùng LLMs là thời gian chờ. Với luồng RAG có kiểm chứng, hệ thống còn phải "
            "kiểm tra câu hỏi, phân tích truy vấn, truy hồi nguồn, tổng hợp và kiểm chứng. Nếu giao diện chỉ hiển thị trạng "
            "thái chờ tĩnh, người dùng khó biết hệ thống còn đang xử lý hay đã gặp lỗi."
        ),
        252: "Đồ án giải quyết vấn đề này bằng cơ chế truyền tiến trình theo thời gian thực:",
        258: (
            "Admin Web được xây dựng để phục vụ quản trị kho tri thức pháp luật. Khác với Mobile App, giao diện này không "
            "tập trung vào hội thoại mà tập trung vào các thao tác vận hành như xem thống kê, quản lý văn bản, tải văn bản "
            "mới và theo dõi tiến trình xử lý."
        ),
        266: (
            "Do các thao tác quản trị có thể ảnh hưởng trực tiếp đến kho tri thức của hệ thống, Admin Web ưu tiên sự rõ ràng "
            "và khả năng kiểm soát. Quản trị viên cần biết văn bản nào đã được xử lý, tác vụ nào đang chạy và lỗi xảy ra ở "
            "bước nào nếu quá trình cập nhật thất bại."
        ),
        269: (
            "Pipeline tải lên PDF là tác vụ nền kéo dài, gồm nhiều bước và không sinh văn bản liên tục như chat. Vì vậy, "
            "Admin Web sử dụng WebSocket để nhận trạng thái xử lý từ backend trong suốt quá trình cập nhật tài liệu."
        ),
        278: (
            "Kênh cập nhật được giới hạn theo tài khoản quản trị để mỗi người chỉ nhìn thấy tiến trình của tác vụ do mình "
            "khởi tạo. Điều này cần thiết khi hệ thống có nhiều quản trị viên cùng vận hành."
        ),
        279: (
            "So với việc hỏi trạng thái định kỳ, WebSocket phù hợp hơn vì tiến trình xử lý không thay đổi theo chu kỳ cố "
            "định. Backend chỉ gửi cập nhật khi có trạng thái mới, nhờ đó giao diện phản hồi kịp thời hơn và giảm số lượng "
            "request không cần thiết."
        ),
        281: (
            "Cả Mobile App và Admin Web đều giao tiếp với hệ thống thông qua Main Service. Cách thiết kế này giữ RAG Service "
            "ở phía nội bộ, đồng thời tập trung xác thực, phân quyền và điều phối nghiệp vụ tại một điểm thống nhất."
        ),
        287: (
            "Như vậy, lớp ứng dụng không chỉ là phần hiển thị giao diện. Nó phản ánh các quyết định kiến trúc của hệ thống: "
            "tách trải nghiệm người dùng cuối khỏi giao diện quản trị, dùng dòng sự kiện cho chat, dùng WebSocket cho tác vụ "
            "nền và chia sẻ logic mobile đa nền tảng."
        ),
        290: (
            "Chương 2 đã trình bày phương pháp xây dựng hệ thống Vietnam Law Chatbot từ kiến trúc tổng thể đến các luồng "
            "xử lý trọng tâm. Hệ thống được tổ chức theo kiến trúc microservices, trong đó Main Service đảm nhiệm điều phối "
            "nghiệp vụ, RAG Service đảm nhiệm truy hồi và suy luận, còn các kho dữ liệu lưu trữ tri thức pháp luật, vector "
            "tìm kiếm và lịch sử sử dụng."
        ),
        291: (
            "Đối với chức năng hỏi đáp, báo cáo đã phân tích cách hệ thống kết hợp tác nhân AI, truy hồi hai giai đoạn, "
            "đối chiếu nguồn cập nhật và kiểm chứng đầu ra để tạo câu trả lời có căn cứ. Đối với chức năng tư vấn có hướng "
            "dẫn, hệ thống bổ sung bước làm rõ ngữ cảnh trước khi trả lời nhằm hạn chế suy đoán trong các câu hỏi thiếu dữ kiện."
        ),
        292: (
            "Đối với chức năng cập nhật văn bản pháp luật, báo cáo đã làm rõ quy trình từ tải PDF, tiền kiểm, trích xuất, "
            "chuẩn hóa, kiểm tra hợp lệ, lưu trữ, chia nhỏ văn bản đến cập nhật kho vector. Đây là nền tảng để hệ thống duy "
            "trì kho tri thức có cấu trúc và có thể mở rộng."
        ),
        293: (
            "Các nội dung trên tạo cơ sở cho chương 3, nơi báo cáo chuyển sang phân tích và thiết kế hệ thống ở mức chi tiết "
            "hơn, bao gồm yêu cầu chức năng, mô hình dữ liệu, API và các biểu đồ xử lý."
        ),
        120: (
            "Trong công thức trên, bước xếp hạng lại được đặt trọng số cao hơn vì nó đánh giá trực tiếp mối quan hệ giữa "
            "câu hỏi và đoạn văn bản. Điểm tìm kiếm vector vẫn được giữ lại để bảo đảm các đoạn có tương đồng ngữ nghĩa cao "
            "không bị loại bỏ hoàn toàn."
        ),
    }

    # Các đoạn bị xoá chủ yếu là phần nói lại cùng một ý ở ngay đoạn liền trước.
    delete_indices = {87, 105, 132, 159, 196, 205, 206}

    for index, text in replacements.items():
        if index < len(paragraphs):
            paragraphs[index].text = text

    for index in sorted(delete_indices, reverse=True):
        if index < len(paragraphs):
            delete_paragraph(paragraphs[index])

    duplicate_texts = {
        'Ví dụ, với câu hỏi ban đầu "Vượt đèn đỏ phạt bao nhiêu?", nếu người dùng chọn "Xe máy", hệ thống sẽ ưu tiên các căn cứ liên quan đến xe mô tô, xe gắn máy và chỉ mở rộng sang nguồn khác khi cần đối chiếu tính cập nhật.',
        "Kết quả chuẩn hoá bao gồm mã văn bản, tên văn bản, ngày hiệu lực, danh sách điều khoản và các metadata cần thiết để phục vụ tìm kiếm, trích dẫn và quản trị.",
    }
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text in duplicate_texts or text.startswith("Đặc thù pháp luật Việt Nam:"):
            delete_paragraph(paragraph)

    # Sửa lại phần công thức nếu các chỉ mục bị lệch sau khi xoá đoạn.
    paragraphs_after_delete = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs_after_delete):
        text = paragraph.text.strip()
        if text == "Chiến lược Two-Stage giải quyết trade-off:":
            paragraph.text = "Chiến lược hai giai đoạn được tổ chức như sau:"
        if text == "Điểm số cuối cùng kết hợp cả hai nguồn:" and index + 1 < len(paragraphs_after_delete):
            formula_paragraph = paragraphs_after_delete[index + 1]
            if "$$\\text{FinalScore}" not in formula_paragraph.text:
                formula_paragraph.text = "$$\\text{FinalScore} = 0.3 \\times \\text{BiEncoderScore} + 0.7 \\times \\text{CrossEncoderScore}$$"
                insert_after(
                    formula_paragraph,
                    "Trong công thức trên, bước xếp hạng lại được đặt trọng số cao hơn vì nó đánh giá trực tiếp mối quan hệ giữa câu hỏi và đoạn văn bản. Điểm tìm kiếm vector vẫn được giữ lại để bảo đảm các đoạn có tương đồng ngữ nghĩa cao không bị loại bỏ hoàn toàn.",
                )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
