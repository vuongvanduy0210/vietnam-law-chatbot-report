from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run


ROOT = Path(__file__).resolve().parents[2]
TARGETS = [
    ROOT / "Vương Văn Duy_Báo cáo.docx",
    ROOT / "Bao_Cao" / "02_Chuong_2_ban_dinh_dang_theo_mau.docx",
]

CODE_RE = re.compile(
    "|".join(
        [
            r"/api/[A-Za-z0-9_/:.-]+",
            r"\{[^{}]{1,120}\}",
            r"\b[A-Za-z_][A-Za-z0-9_]*(?:\.[A-Za-z_][A-Za-z0-9_]*)+\b",
            r"\b[A-Za-z_][A-Za-z0-9_]*=\d+(?:\.\d+)?\b",
            r"\b[A-Z][A-Z0-9]+(?:_[A-Z0-9]+)+\b",
            r"\b[a-z][a-z0-9]+(?:_[a-z0-9]+)+\b",
            r"\b(?:ready|progress|done|error|pending|processing|completed|failed|cancelled|timeout|message)\b",
            r"\b(?:JWT|SSE|REST|HTTPS|HTTP|WebSocket|Bearer|X-API-Key|API Key|admin_token)\b",
            r"\b(?:asyncpg|vietnamese_law)\b",
            r"\b(?:BiEncoderScore|CrossEncoderScore|FinalScore)\b",
        ]
    )
)

WHOLE_CODE_RE = re.compile(
    r"^\s*(?:Thought\s+\d+|Act\s+\d+|Observation\s+\d+|Finish\b|\.\.\.|\[Lặp lại)",
    re.IGNORECASE,
)

SKIP_STYLE_PREFIXES = (
    "Heading",
    "Tên hình",
    "Tên bảng",
    "Caption",
    "Title",
    "TOC",
)


def iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def set_code_font(run: Run) -> None:
    run.font.name = "Courier New"
    run.font.size = Pt(11)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), "Courier New")

    sz = r_pr.sz
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), "22")

    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), "22")


def style_name(paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def should_skip(paragraph) -> bool:
    name = style_name(paragraph)
    return any(name.startswith(prefix) for prefix in SKIP_STYLE_PREFIXES)


def split_segments(text: str):
    pos = 0
    for match in CODE_RE.finditer(text):
        if match.start() > pos:
            yield text[pos : match.start()], False
        yield match.group(0), True
        pos = match.end()
    if pos < len(text):
        yield text[pos:], False


def replace_run_with_segments(paragraph, run: Run, segments) -> int:
    segments = [(text, is_code) for text, is_code in segments if text]
    if not segments:
        return 0

    first_text, first_is_code = segments[0]
    run.text = first_text
    if first_is_code:
        set_code_font(run)
    changed = int(first_is_code)

    anchor = run._r
    for text, is_code in segments[1:]:
        new_r = deepcopy(run._r)
        anchor.addnext(new_r)
        new_run = Run(new_r, paragraph)
        new_run.text = text
        if is_code:
            set_code_font(new_run)
            changed += 1
        anchor = new_r
    return changed


def apply_to_doc(path: Path) -> int:
    doc = Document(path)
    changed = 0

    for paragraph in iter_paragraphs(doc):
        text = paragraph.text
        if not text.strip():
            continue

        if "$$" in text or WHOLE_CODE_RE.search(text):
            for run in paragraph.runs:
                if run.text:
                    set_code_font(run)
                    changed += 1
            continue

        if should_skip(paragraph) or not CODE_RE.search(text):
            continue

        # Copy the run list first because splitting inserts new runs.
        for run in list(paragraph.runs):
            if run.text and CODE_RE.search(run.text):
                changed += replace_run_with_segments(paragraph, run, split_segments(run.text))

    doc.save(path)
    return changed


def main() -> None:
    for target in TARGETS:
        if not target.exists():
            print(f"skip missing: {target}")
            continue
        backup = target.with_name(f"{target.stem}_backup_truoc_font_code{target.suffix}")
        if not backup.exists():
            shutil.copy2(target, backup)
        changed = apply_to_doc(target)
        print(f"formatted {changed} code runs: {target}")


if __name__ == "__main__":
    main()
