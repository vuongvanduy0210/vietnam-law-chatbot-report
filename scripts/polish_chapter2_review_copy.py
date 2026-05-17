from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "Bao_Cao" / "02_Chuong_2.docx"
OUT = ROOT / "Bao_Cao" / "02_Chuong_2_ban_hoan_chinh_de_kiem_tra.docx"


def set_cell(table, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = text


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    paragraphs = list(doc.paragraphs)

    replacements = {
        1: (
            "Chương này trình bày các phương pháp và kiến trúc kỹ thuật được áp dụng để xây dựng hệ thống "
            "Vietnam Law Chatbot. Khác với chương 1 tập trung vào cơ sở lý thuyết, chương 2 đi sâu vào cách "
            "tổ chức hệ thống, phương pháp truy hồi tri thức, quy trình suy luận có kiểm chứng, luồng tư vấn "
            "có hướng dẫn, cơ chế cập nhật văn bản pháp luật và cách xây dựng lớp ứng dụng phục vụ người dùng."
        ),
        6: (
            "Thứ nhất, hệ thống cần tách riêng các tác vụ tính toán nặng. Các bước như tạo embedding, xếp hạng "
            "lại kết quả truy hồi và suy luận bằng LLMs thường có độ trễ cao hơn nhiều so với các API nghiệp vụ "
            "thông thường. Nếu đặt toàn bộ trong một khối ứng dụng duy nhất, các chức năng đăng nhập, xem lịch sử "
            "hội thoại hoặc quản trị dữ liệu có thể bị ảnh hưởng khi pipeline AI đang xử lý tải lớn."
        ),
        14: (
            "• Mobile App (Kotlin Multiplatform): ứng dụng dành cho người dùng cuối, chạy trên Android và iOS từ "
            "một nền tảng mã nguồn chung. Ứng dụng gửi yêu cầu tới backend qua HTTPS/JWT và nhận câu trả lời theo "
            "dòng sự kiện thời gian thực."
        ),
        43: (
            "Từ kiến trúc tổng thể trên, thành phần cần phân tích sâu nhất là RAG Service, vì đây là nơi quyết định "
            "chất lượng tư vấn pháp luật. Trước khi thiết kế quy trình suy luận, cần xác định rõ tác nhân AI hoạt "
            "động trong môi trường nào, được phép sử dụng nguồn dữ liệu nào và phải bị giới hạn bởi các nguyên tắc "
            "nào để tránh trả lời vượt quá phạm vi của hệ thống."
        ),
        51: "Tác nhân AI được thiết kế với năm nhóm hành động chính:",
        52: "1. Tra cứu tri thức nội bộ: tìm điều luật, văn bản và đoạn nội dung liên quan trong kho tri thức của hệ thống.",
        53: "2. Đối chiếu nguồn cập nhật bên ngoài: kiểm tra thêm thông tin pháp luật hiện hành khi câu hỏi có khả năng phụ thuộc vào văn bản mới.",
        54: "3. Tổng hợp câu trả lời: kết hợp các nguồn bằng chứng, diễn giải theo ngôn ngữ dễ hiểu và nêu rõ căn cứ pháp lý.",
        55: "4. Từ chối phù hợp: nhận diện câu hỏi ngoài phạm vi pháp luật hoặc không đủ điều kiện để tư vấn.",
        56: "5. Lặp lại quá trình tra cứu: bổ sung truy vấn phụ khi bằng chứng thu được chưa đủ rõ để hình thành câu trả lời đáng tin cậy.",
        58: "2.3. Luồng RAG có kiểm chứng",
        69: "2.3.3. Bước 1 — Kiểm soát đầu vào",
        71: "Bảng 2.6. Dữ liệu đầu vào và đầu ra của bước kiểm soát đầu vào",
        72: (
            'Ví dụ minh hoạ: với câu hỏi "Viết cho tôi một đoạn code Python sắp xếp mảng", hệ thống xác định đây '
            "không phải yêu cầu tư vấn pháp luật và từ chối ngay từ bước đầu. Nhờ vậy, pipeline AI phía sau không "
            "phải tiêu tốn tài nguyên cho một yêu cầu nằm ngoài phạm vi đồ án."
        ),
        73: "2.3.4. Bước 2 — Phân tích và tái cấu trúc truy vấn",
        87: "Bảng 2.7. Dữ liệu đầu vào và đầu ra của bước phân tích truy vấn",
        88: "2.3.5. Bước 3 — Suy luận và thu thập bằng chứng",
        95: "Bảng 2.8. Dữ liệu đầu vào và đầu ra của bước suy luận",
        96: "2.3.6. Bước 4 — Kiểm chứng câu trả lời",
        99: "Bảng 2.9. Dữ liệu đầu vào và đầu ra của bước kiểm chứng",
        110: (
            "Nguồn tìm kiếm bên ngoài được sử dụng như một kênh bổ sung khi hệ thống cần kiểm tra tính cập nhật của "
            "quy định pháp luật. Trong phạm vi báo cáo, điểm quan trọng không nằm ở tên công cụ cụ thể, mà ở vai trò "
            "phương pháp: nguồn bên ngoài giúp đối chiếu thông tin mới trước khi hệ thống đưa ra kết luận cuối cùng."
        ),
        118: (
            'Hình 2.6 mô tả pipeline tìm kiếm theo hướng "lọc rộng trước, chấm sâu sau". Ở đầu pipeline, hệ thống dùng '
            "tìm kiếm vector để thu hẹp không gian tìm kiếm từ kho văn bản lớn xuống một tập ứng viên nhỏ hơn. Sau đó, "
            "bước xếp hạng lại đánh giá mức phù hợp giữa câu hỏi và từng đoạn văn bản để chọn ra các căn cứ có giá trị nhất."
        ),
        120: "2.5.1. Hạn chế của tìm kiếm vector đơn lẻ",
        121: (
            "Tìm kiếm vector có ưu điểm lớn về tốc độ vì các đoạn văn bản trong kho tri thức có thể được mã hoá trước "
            "và lưu trong cơ sở dữ liệu vector. Khi người dùng đặt câu hỏi, hệ thống chỉ cần mã hoá câu hỏi và so sánh "
            "độ tương đồng để lấy ra các đoạn có khả năng liên quan."
        ),
        122: (
            "Tuy nhiên, cách tiếp cận này chưa luôn hiểu được mối quan hệ ngữ nghĩa tinh tế giữa câu hỏi và văn bản. "
            'Ví dụ, với câu hỏi "thủ tục đăng ký kết hôn với người nước ngoài", hệ thống có thể trả về nhiều đoạn chỉ '
            'chứa cụm "đăng ký kết hôn" nhưng không trực tiếp trả lời đúng trường hợp có yếu tố nước ngoài.'
        ),
        123: "2.5.2. Xếp hạng lại kết quả truy hồi",
        124: (
            "Để khắc phục hạn chế trên, hệ thống bổ sung bước xếp hạng lại. Thay vì chỉ so sánh hai vector đã được mã "
            "hoá độc lập, bước này đánh giá trực tiếp từng cặp câu hỏi và đoạn văn bản, nhờ đó nhận diện tốt hơn mức "
            "liên quan thực sự giữa yêu cầu của người dùng và căn cứ pháp lý."
        ),
        126: "• Giai đoạn 1: tìm kiếm vector nhanh để lấy một tập ứng viên rộng từ kho tri thức.",
        127: "• Giai đoạn 2: xếp hạng lại tập ứng viên để loại bỏ các đoạn chỉ liên quan bề mặt và ưu tiên đoạn trả lời trực tiếp hơn.",
        128: (
            "Cách thiết kế hai giai đoạn tạo ra sự cân bằng giữa tốc độ và chất lượng. Nếu lấy quá ít ứng viên ở giai "
            "đoạn đầu, hệ thống có nguy cơ bỏ sót điều luật phù hợp; nếu lấy quá nhiều, bước xếp hạng lại sẽ làm tăng "
            "độ trễ. Vì vậy, pipeline được cấu hình để lấy đủ rộng ở bước đầu và chỉ chấm sâu trên một tập đã được thu hẹp."
        ),
        129: "2.5.3. Công thức kết hợp điểm phù hợp",
        132: (
            "Trọng số cao hơn được dành cho bước xếp hạng lại vì bước này đánh giá quan hệ trực tiếp giữa câu hỏi và "
            "nội dung văn bản. Điểm tìm kiếm vector vẫn được giữ lại để bảo đảm các đoạn có tương đồng ngữ nghĩa cao "
            "không bị loại bỏ hoàn toàn khi bước chấm sâu có sai lệch."
        ),
        133: "2.5.4. Ưu tiên văn bản mới",
        136: "2.5.5. Giải quyết mâu thuẫn pháp lý theo thời gian",
        144: "2.5.6. Ngưỡng tin cậy",
        152: "2.6.1. Vấn đề và động lực thiết kế",
        153: (
            "Trong thực tế tư vấn pháp luật, nhiều câu hỏi ban đầu của người dùng thường thiếu dữ kiện quan trọng. "
            'Ví dụ, câu hỏi "vượt đèn đỏ phạt bao nhiêu" chưa cho biết người dùng đi xe máy hay ô tô, hành vi có gây '
            "tai nạn hay không, và bối cảnh xảy ra là gì. Nếu hệ thống trả lời ngay, câu trả lời có thể đúng về mặt "
            "văn bản nhưng chưa đúng với tình huống cụ thể."
        ),
        154: "Nếu chỉ sử dụng luồng hỏi đáp một bước, hệ thống thường gặp hai vấn đề:",
        155: "1. Hệ thống phải tự giả định thông tin còn thiếu, làm tăng nguy cơ trả lời không phù hợp với tình huống thực tế.",
        156: "2. Nếu hỏi lại bằng văn bản tự do, trải nghiệm người dùng kém rõ ràng và khó chuẩn hoá dữ liệu đầu vào cho bước trả lời.",
        157: (
            "Vì vậy, đồ án thiết kế thêm luồng tư vấn có hướng dẫn. Luồng này tách quá trình tư vấn thành hai bước: "
            "thu thập thêm ngữ cảnh cần thiết và trả lời dựa trên ngữ cảnh đã được làm rõ."
        ),
        159: "2.6.2. Thiết kế không lưu trạng thái phiên",
        160: (
            "Luồng tư vấn có hướng dẫn được thiết kế theo hướng không phụ thuộc vào trạng thái phiên trên máy chủ. "
            "Sau bước làm rõ, các lựa chọn của người dùng được gửi kèm trong yêu cầu trả lời. Cách làm này giúp hệ "
            "thống dễ mở rộng, giảm ràng buộc giữa các request và phù hợp với kiến trúc microservices."
        ),
        161: "Quy trình gồm ba thành phần thông tin chính:",
        162: "1. Câu hỏi ban đầu của người dùng.",
        163: "2. Danh sách câu hỏi làm rõ do hệ thống sinh ra.",
        164: "3. Các lựa chọn hoặc thông tin bổ sung mà người dùng cung cấp trước khi nhận câu trả lời cuối cùng.",
        165: "2.6.3. Bước 1 — Thu thập ngữ cảnh",
        166: (
            "Ở bước đầu tiên, hệ thống nhận câu hỏi ban đầu và xác định những dữ kiện còn thiếu để trả lời chính xác. "
            "Kết quả của bước này không phải là câu trả lời pháp lý ngay lập tức, mà là một tập câu hỏi làm rõ được "
            "trình bày dưới dạng lựa chọn dễ thao tác trên giao diện."
        ),
        181: "2.6.4. Bước 2 — Trả lời theo ngữ cảnh đã làm rõ",
        182: (
            "Ở bước thứ hai, hệ thống nhận lại câu hỏi ban đầu cùng các lựa chọn làm rõ của người dùng. Từ đó, yêu cầu "
            "được chuyển thành một truy vấn cụ thể hơn để phục vụ quá trình truy hồi, suy luận và kiểm chứng."
        ),
        189: (
            "Thông tin bổ sung từ người dùng được dùng để ràng buộc phạm vi trả lời. Ví dụ, cùng là hành vi vượt đèn "
            "đỏ, mức xử phạt có thể khác nhau giữa xe máy và ô tô; do đó lựa chọn phương tiện là dữ kiện quan trọng "
            "để hệ thống truy hồi đúng căn cứ pháp lý."
        ),
        191: (
            "Luồng trả lời vẫn tuân thủ nguyên tắc truy hồi, suy luận và kiểm chứng như luồng chat thông thường. Điểm "
            "khác biệt là câu hỏi đầu vào đã được làm giàu ngữ cảnh, giúp giảm nhu cầu suy đoán và tăng khả năng trích "
            "dẫn đúng văn bản."
        ),
        192: (
            'Ví dụ, với câu hỏi ban đầu "Vượt đèn đỏ phạt bao nhiêu?", nếu người dùng chọn "Xe máy", hệ thống sẽ ưu '
            "tiên các căn cứ liên quan đến xe mô tô, xe gắn máy và chỉ mở rộng sang nguồn khác khi cần đối chiếu tính cập nhật."
        ),
        194: "2.6.5. Cơ chế trả lời theo dòng sự kiện",
        195: (
            "Câu trả lời trong luồng tư vấn có hướng dẫn được truyền về giao diện theo từng bước xử lý. Người dùng có "
            "thể thấy hệ thống đang kiểm tra câu hỏi, truy hồi tài liệu, tổng hợp căn cứ và hoàn thiện câu trả lời."
        ),
        196: (
            "Cách truyền kết quả theo dòng sự kiện giúp giao diện phản hồi sớm hơn, đặc biệt với các câu hỏi cần nhiều "
            "bước truy hồi và kiểm chứng. Thay vì chờ toàn bộ pipeline kết thúc, người dùng liên tục nhận được tín hiệu "
            "về tiến trình xử lý."
        ),
        197: "Bảng 2.13. Các loại sự kiện trong luồng tư vấn có hướng dẫn",
        201: "2.7. Quy trình cập nhật tri thức từ văn bản pháp luật",
        210: (
            "Bước tiền kiểm có ý nghĩa lớn trong vận hành thực tế. Nếu không có bước này, mỗi khi quản trị viên tải "
            "nhầm một văn bản đã tồn tại, hệ thống vẫn phải thực hiện các bước tốn kém như trích xuất nội dung, chuẩn "
            "hoá cấu trúc và cập nhật vector trước khi phát hiện trùng lặp. Tiền kiểm giúp loại bỏ sớm các trường hợp "
            "không hợp lệ và giảm đáng kể chi phí xử lý."
        ),
        211: "2.7.2. Trích xuất và cấu trúc hóa văn bản PDF",
        212: (
            "Pipeline xử lý PDF được chia thành hai nhóm thao tác chính: trích xuất lớp văn bản và chuẩn hoá nội dung "
            "thành cấu trúc pháp lý. Việc tách hai nhiệm vụ này giúp hệ thống tận dụng đúng thế mạnh của từng thành phần."
        ),
        213: (
            "Thứ nhất, hệ thống trích xuất nội dung từ file PDF. Với văn bản có lớp text, nội dung có thể được lấy trực "
            "tiếp; với văn bản scan hoặc chất lượng thấp, hệ thống cần sử dụng OCR để nhận dạng ký tự."
        ),
        214: (
            "Thứ hai, LLMs được sử dụng để hiểu cấu trúc văn bản pháp luật và chuẩn hoá thành các đơn vị như văn bản, "
            "chương, mục, điều và khoản. LLMs không thay thế bước trích xuất chữ, mà tập trung vào nhiệm vụ hiểu cấu trúc "
            "và chuẩn hoá nội dung."
        ),
        219: "• LLMs tập trung vào nhiệm vụ mạnh nhất: hiểu cấu trúc pháp lý và chuẩn hóa nội dung.",
        220: (
            "Kết quả chuẩn hoá bao gồm mã văn bản, tên văn bản, ngày hiệu lực, danh sách điều khoản và các metadata cần "
            "thiết để phục vụ tìm kiếm, trích dẫn và quản trị."
        ),
        238: (
            "Kết quả chuẩn hoá không được lưu ngay mà phải đi qua bước kiểm tra hợp lệ. Bước này bảo đảm các trường quan "
            "trọng như mã văn bản, mã điều, tiêu đề, nội dung và thông tin nguồn không bị thiếu. Đây là yêu cầu cần thiết "
            "vì đầu ra do LLMs sinh ra vẫn có thể thiếu trường, sai định dạng hoặc chia tách điều khoản chưa ổn định."
        ),
        239: "2.7.3. Xử lý song song",
        240: (
            "Một số tác vụ trong quá trình cập nhật văn bản có thời gian xử lý dài, chẳng hạn lưu file gốc lên kho lưu "
            "trữ và trích xuất nội dung từ PDF. Hệ thống tổ chức các tác vụ này theo hướng song song để giảm thời gian "
            "chờ của quản trị viên và tránh làm nghẽn luồng xử lý chính."
        ),
        246: (
            "Lưu file gốc: hệ thống lưu bản PDF ban đầu để bảo đảm mỗi điều luật có thể tham chiếu lại nguồn văn bản khi "
            "cần kiểm tra hoặc đối chiếu."
        ),
        247: (
            "Trích xuất nội dung: hệ thống lấy nội dung từ PDF và chuẩn bị dữ liệu đầu vào cho bước chuẩn hoá cấu trúc."
        ),
        248: (
            "Sau khi hai tác vụ hoàn thành, kết quả được kết hợp để tạo thành một bản ghi văn bản đầy đủ, vừa có nội "
            "dung đã được chuẩn hoá, vừa có liên kết tới file nguồn."
        ),
        250: (
            "Về mặt vận hành, các tác vụ tốn CPU được đưa ra khỏi luồng xử lý chính để không ảnh hưởng đến các chức năng "
            "khác như chat hoặc dashboard. Cách thiết kế này giúp hệ thống ổn định hơn khi có nhiều yêu cầu đồng thời."
        ),
        251: "2.7.4. Cơ chế hoàn tác để giữ dữ liệu nhất quán",
        258: (
            "Trong hệ thống này, một lượt cập nhật chỉ được xem là thành công khi dữ liệu đã sẵn sàng ở cả ba phía: "
            "kho lưu trữ văn bản cho quản trị viên, cơ sở dữ liệu vector cho truy hồi, và lớp dữ liệu phục vụ hiển thị. "
            "Nếu một bước thất bại, hệ thống phải hoàn tác các phần đã ghi trước đó để tránh tình trạng dữ liệu bị lệch."
        ),
        259: "2.7.5. Chiến lược chia nhỏ văn bản",
        263: (
            "Mỗi đoạn sau khi chia nhỏ vẫn giữ thông tin nhận diện văn bản và vị trí điều khoản gốc. Nhờ vậy, khi hệ "
            "thống trích dẫn một đoạn trong câu trả lời, người dùng và quản trị viên có thể truy ngược về văn bản nguồn."
        ),
        272: "2.8.1. Cơ chế trả lời theo dòng sự kiện cho chat",
        275: "2.8.2. Các trạng thái trong luồng chat",
        280: "2.8.4. Câu hỏi gợi ý sau câu trả lời",
        287: (
            "Thay vì xây dựng hai ứng dụng độc lập cho Android và iOS, Kotlin Multiplatform cho phép chia sẻ phần lớn "
            "logic nghiệp vụ giữa hai nền tảng, bao gồm:"
        ),
        289: (
            "Trong bối cảnh đồ án chuyên ngành mobile, lựa chọn Kotlin Multiplatform không chỉ nhằm giảm khối lượng lặp "
            "lại, mà còn thể hiện cách thiết kế một ứng dụng đa nền tảng có cấu trúc rõ ràng. Những phần phụ thuộc đặc "
            "thù nền tảng như giao diện native, quyền hệ thống hoặc vòng đời ứng dụng vẫn được tách riêng để bảo đảm "
            "trải nghiệm phù hợp trên từng thiết bị."
        ),
        301: (
            "MVI đặc biệt phù hợp với màn hình chat vì tại cùng một thời điểm có thể tồn tại nhiều trạng thái: tin nhắn "
            "người dùng vừa gửi, pipeline đang xử lý, tiến trình đang hiển thị, câu trả lời cuối cùng chưa hoàn thành "
            "và danh sách câu hỏi gợi ý đang chờ cập nhật. Tất cả được biểu diễn trong một trạng thái giao diện duy nhất."
        ),
        304: "Đồ án giải quyết yêu cầu này bằng cơ chế truyền tiến trình theo thời gian thực:",
        305: "1. Khi người dùng gửi câu hỏi, giao diện lập tức hiển thị tin nhắn và chuyển sang trạng thái đang xử lý.",
        306: "2. Trong quá trình xử lý, backend gửi các bước tiến trình dễ hiểu như kiểm tra câu hỏi, truy hồi nguồn và kiểm chứng.",
        307: "3. Khi câu trả lời hoàn tất, giao diện thay thế trạng thái chờ bằng nội dung trả lời và danh sách câu hỏi gợi ý.",
        318: (
            "Admin Web được thiết kế theo hướng dashboard vận hành thay vì giao diện người dùng phổ thông. Các thao tác "
            "như tải văn bản, theo dõi tác vụ đang chạy, kiểm tra số điều luật và xoá văn bản đều ảnh hưởng trực tiếp "
            "đến kho tri thức của hệ thống. Vì vậy, giao diện quản trị ưu tiên tính rõ ràng, khả năng kiểm soát và phản "
            "hồi trạng thái hơn là các hiệu ứng trình bày."
        ),
        320: "2.9.5. WebSocket cho tiến trình xử lý tài liệu",
        323: "1. Quản trị viên chọn file PDF và xác nhận tải lên.",
        324: "2. Backend tạo tác vụ xử lý để theo dõi tiến trình.",
        325: "3. Giao diện mở kênh nhận cập nhật trạng thái của tác vụ.",
        326: "4. Backend gửi các mốc tiến trình như tiền kiểm, trích xuất, chuẩn hoá, lưu dữ liệu và tạo vector.",
        327: "5. Giao diện cập nhật thanh tiến trình, trạng thái và thông báo lỗi nếu có.",
        328: "6. Khi hoàn tất, văn bản mới xuất hiện trong danh sách quản trị và sẵn sàng phục vụ tra cứu.",
        333: (
            "Cả hai ứng dụng frontend đều giao tiếp với Main Service thông qua các nhóm API nghiệp vụ. Main Service "
            "đóng vai trò lớp điều phối, che giấu độ phức tạp của RAG Service và các kho dữ liệu phía sau."
        ),
        337: (
            "Mobile App tập trung vào các nhóm chức năng phục vụ người dùng cuối, bao gồm xác thực, quản lý hội thoại, "
            "tra cứu thư viện pháp luật và luồng tư vấn có hướng dẫn."
        ),
        342: (
            "Admin Web tập trung vào các nhóm chức năng phục vụ quản trị, bao gồm theo dõi thống kê hệ thống, quản lý "
            "văn bản pháp luật, tải văn bản mới và giám sát tiến trình xử lý tài liệu."
        ),
        351: (
            "Chương 2 đã trình bày phương pháp xây dựng hệ thống Vietnam Law Chatbot từ góc nhìn kiến trúc và luồng xử "
            "lý. Hệ thống được tổ chức theo kiến trúc microservices, trong đó Main Service đảm nhiệm điều phối nghiệp vụ, "
            "RAG Service đảm nhiệm truy hồi và suy luận, còn các kho dữ liệu lưu trữ tri thức pháp luật và lịch sử sử dụng."
        ),
        352: (
            "Đối với luồng hỏi đáp, chương này làm rõ cách hệ thống kết hợp RAG, tác nhân AI, tìm kiếm hai giai đoạn và "
            "bước kiểm chứng để giảm rủi ro trả lời sai căn cứ. Đối với luồng tư vấn có hướng dẫn, hệ thống tách quá "
            "trình tư vấn thành bước làm rõ ngữ cảnh và bước trả lời theo thông tin đã được bổ sung."
        ),
        353: (
            "Đối với luồng cập nhật văn bản pháp luật, chương này trình bày các bước tiền kiểm, trích xuất, chuẩn hoá, "
            "kiểm tra hợp lệ, lưu trữ, chia nhỏ văn bản và cập nhật cơ sở dữ liệu vector. Đây là phần quan trọng để hệ "
            "thống duy trì kho tri thức có cấu trúc và có thể mở rộng."
        ),
        354: (
            "Các nội dung trên tạo nền tảng cho chương 3, nơi báo cáo chuyển từ phương pháp thiết kế sang mô tả chi tiết "
            "quá trình cài đặt, kiểm thử và đánh giá hệ thống."
        ),
    }

    replacements.update(
        {
            7: (
                "Thứ hai, kiến trúc tách dịch vụ giúp tăng mức độ an toàn. Thành phần RAG Service chứa các mô hình "
                "và quy trình AI quan trọng nên được đặt ở lớp nội bộ, chỉ nhận yêu cầu thông qua Main Service. Người "
                "dùng cuối vì thế không truy cập trực tiếp vào phần suy luận và kho tri thức của hệ thống."
            ),
            8: (
                "Thứ ba, kiến trúc này hỗ trợ phát triển độc lập. Main Service xử lý các nghiệp vụ ổn định như xác "
                "thực, hội thoại và quản trị dữ liệu; trong khi RAG Service có thể thay đổi thường xuyên hơn do cần "
                "điều chỉnh mô hình, chiến lược truy hồi và cách kiểm chứng câu trả lời."
            ),
            18: (
                "• RAG Service (port 8001): dịch vụ nội bộ phụ trách truy hồi tri thức, suy luận bằng LLMs, kiểm chứng "
                "câu trả lời và xử lý cập nhật văn bản."
            ),
            22: "• ChromaDB: kho vector phục vụ tìm kiếm ngữ nghĩa trên các đoạn văn bản đã được chia nhỏ.",
            23: (
                "Lớp dịch vụ bên ngoài: nhà cung cấp LLMs, dịch vụ tìm kiếm web, dịch vụ đối chiếu thông tin cập nhật "
                "và kho lưu trữ file PDF."
            ),
            158: (
                "Về mặt kỹ thuật, bước làm rõ ngữ cảnh giúp giảm rủi ro suy đoán. Khi các dữ kiện chính đã được người "
                "dùng xác nhận, câu hỏi gửi vào pipeline trở nên cụ thể hơn, truy hồi chính xác hơn và bước kiểm chứng "
                "có phạm vi đối chiếu rõ ràng hơn."
            ),
            180: (
                "Ở bước này, hệ thống chỉ tạo câu hỏi làm rõ, chưa đưa ra tư vấn pháp lý cuối cùng. Đây là ranh giới "
                "quan trọng: câu hỏi làm rõ có thể được sinh nhanh, còn câu trả lời pháp lý cuối cùng bắt buộc phải đi "
                "qua truy hồi căn cứ và kiểm chứng trước khi hiển thị."
            ),
            198: (
                "Ứng dụng nhận dòng sự kiện và cập nhật vùng hiển thị tiến trình theo từng trạng thái xử lý. Khi nhận "
                "được kết quả cuối cùng, giao diện hiển thị câu trả lời hoàn chỉnh cùng danh sách nguồn tham chiếu. "
                "Thiết kế này minh bạch hơn trạng thái chờ tĩnh vì người dùng thấy hệ thống đang thực sự tra cứu và kiểm chứng."
            ),
            213: (
                "Thứ nhất, hệ thống trích xuất nội dung từ file PDF. Với văn bản có sẵn lớp chữ, nội dung có thể được "
                "lấy trực tiếp; với văn bản scan hoặc chất lượng thấp, hệ thống sử dụng OCR để nhận dạng ký tự."
            ),
            214: (
                "Thứ hai, LLMs được sử dụng để hiểu cấu trúc văn bản pháp luật và chuẩn hoá thành các đơn vị như văn "
                "bản, chương, mục, điều và khoản. LLMs không thay thế bước trích xuất chữ, mà tập trung vào nhiệm vụ "
                "hiểu cấu trúc và chuẩn hoá nội dung."
            ),
            216: "• Giảm chi phí xử lý vì phần lớn PDF pháp luật đã có sẵn lớp chữ để trích xuất.",
            217: "• Chủ động kiểm soát chất lượng nhận dạng tiếng Việt bằng các quy tắc sửa lỗi phổ biến.",
            218: "• Dễ theo dõi số trang xử lý, số ký tự trích xuất và thời gian thực hiện từng bước.",
            237: (
                "Nếu bước nhận dạng hoặc chuẩn hoá nội dung thất bại, hệ thống mới chuyển sang đường xử lý dự phòng "
                "để đọc trực tiếp tài liệu PDF bằng LLMs. Thiết kế này giúp cân bằng giữa chi phí, tốc độ và độ ổn định: "
                "đường chính ưu tiên xử lý văn bản đã trích xuất, còn đường dự phòng chỉ dùng khi cần thiết."
            ),
            249: (
                "Nhờ lưu file gốc song song với quá trình trích xuất nội dung, tổng thời gian xử lý giảm so với cách "
                "làm tuần tự. Khi trích xuất thành công, hệ thống không cần gửi toàn bộ file PDF sang LLMs; khi trích "
                "xuất thất bại, đường xử lý dự phòng vẫn giúp pipeline có cơ hội hoàn tất thay vì dừng ngay."
            ),
            253: "• MongoDB: lưu văn bản pháp luật, nội dung điều khoản và metadata phục vụ quản trị.",
            254: "• ChromaDB: lưu vector của các đoạn văn bản phục vụ truy hồi ngữ nghĩa.",
            257: (
                "Trường hợp một kho dữ liệu đã ghi thành công nhưng bước sau thất bại, hệ thống phải thu hồi phần dữ "
                "liệu đã ghi để tránh lệch trạng thái. Nguyên tắc này bảo đảm màn hình quản trị, kho văn bản và kho "
                "vector luôn phản ánh cùng một trạng thái của văn bản pháp luật."
            ),
            260: "Trước khi lưu vào kho vector, mỗi điều luật được chia thành các đoạn có kích thước vừa đủ:",
            261: "• Kích thước tối đa: khoảng 1.000 từ mỗi đoạn, nhằm giữ đủ ngữ cảnh nhưng không làm đoạn quá dài.",
            262: "• Phần chồng lấn: khoảng 150 từ giữa hai đoạn liền kề để tránh mất ngữ cảnh ở ranh giới.",
            264: (
                "Với cách chia này, phần lớn điều luật ngắn được giữ nguyên, chỉ những điều dài mới cần tách thành "
                "nhiều đoạn. Điều này giúp kho vector vừa đủ chi tiết cho tìm kiếm, vừa không làm tăng số lượng đoạn "
                "một cách không cần thiết."
            ),
            265: (
                "Kích thước đoạn được chọn để cân bằng giữa độ đầy đủ của ngữ cảnh và độ chính xác của truy hồi. Nếu "
                "đoạn quá ngắn, một điều khoản có thể bị tách khỏi điều kiện áp dụng; nếu đoạn quá dài, nội dung đại "
                "diện trong vector sẽ kém đặc trưng và kết quả tìm kiếm dễ bị nhiễu."
            ),
            292: (
                "Mobile App sử dụng mô hình MVI (Model - View - Intent) để tách rõ trạng thái giao diện, hành động của "
                "người dùng và phản hồi một lần của hệ thống. Trong đó:"
            ),
            294: "• Intent mô tả hành động từ người dùng hoặc sự kiện vòng đời của màn hình.",
            295: "• Effect mô tả các phản hồi một lần như điều hướng hoặc thông báo.",
            296: "• Lớp xử lý trung gian tiếp nhận Intent, gọi nghiệp vụ cần thiết và cập nhật State.",
            297: "• Màn hình chỉ hiển thị theo State hiện tại và gửi Intent khi người dùng thao tác.",
            298: (
                "Cách tổ chức này phù hợp với ứng dụng chat AI vì trạng thái giao diện thay đổi liên tục: đang gửi tin "
                "nhắn, đang nhận dữ liệu theo dòng, đang hiển thị tiến trình, đang nhận câu trả lời cuối cùng, đang tải "
                "câu hỏi gợi ý hoặc đang hiển thị lỗi kết nối."
            ),
            300: (
                "Với MVI, các luồng phức tạp như streaming không bị trộn trực tiếp vào giao diện. Sự kiện từ server "
                "được chuyển thành thay đổi trạng thái rõ ràng; màn hình chỉ cần hiển thị theo trạng thái hiện tại. "
                "Cách tổ chức này giúp giao diện chat ổn định hơn khi cùng lúc có nhiều dữ liệu thay đổi."
            ),
            312: "• Đăng nhập với tài khoản quản trị.",
            313: (
                "• Xem dashboard thống kê: tổng số văn bản, tổng số điều luật, số tác vụ thành công, thất bại hoặc "
                "đang xử lý, và các chủ đề phổ biến."
            ),
            314: "• Quản lý văn bản pháp luật: phân trang, tìm kiếm, xem chi tiết, mở rộng từng điều luật và xoá văn bản.",
            315: "• Tải lên file PDF văn bản mới bằng thao tác kéo thả hoặc chọn file.",
            316: "• Theo dõi tiến trình xử lý tài liệu theo thời gian thực.",
            317: "• Khôi phục trạng thái tải lên khi người dùng tải lại trang hoặc gián đoạn kết nối.",
            329: "Bảng 2.19. Các trạng thái xử lý tài liệu",
            330: (
                "Việc giới hạn kênh cập nhật theo tài khoản quản trị giúp mỗi quản trị viên chỉ nhìn thấy tiến trình "
                "của chính mình, tránh lộ thông tin tác vụ giữa các tài khoản khi hệ thống mở rộng."
            ),
            331: (
                "So với cơ chế hỏi trạng thái định kỳ, WebSocket phù hợp hơn cho pipeline tải lên tài liệu vì tiến trình "
                "không thay đổi theo chu kỳ cố định. Backend chỉ gửi cập nhật khi có trạng thái mới, nhờ đó giảm số lượng "
                "request không cần thiết và giúp giao diện phản hồi kịp thời hơn."
            ),
            334: "• Bảo vệ RAG Service khỏi truy cập trực tiếp từ phía client.",
            335: "• Tập trung xác thực, phân quyền và ghi nhận nhật ký nghiệp vụ ở Main Service.",
            336: "• Giữ cho lớp frontend chỉ cần làm việc với một cổng API thống nhất.",
        }
    )

    delete_indices = {
        76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86,
        167, 168, 169, 170, 171, 172, 173, 174, 175, 176, 177, 178, 179,
        183, 184, 185, 186, 187, 188,
        193,
        221, 222, 223, 224, 225, 226, 227, 228, 229, 230, 231, 232, 233, 234, 235, 236,
        241, 242, 243, 244, 245,
        338, 339, 340, 341,
        343, 344, 345, 346, 347,
    }

    for index, text in replacements.items():
        if index < len(paragraphs):
            paragraphs[index].text = text

    for index in sorted(delete_indices, reverse=True):
        if index < len(paragraphs):
            delete_paragraph(paragraphs[index])

    # Bảng 2.2: diễn giải theo khái niệm thay vì thuật ngữ công cụ.
    t = doc.tables[1]
    set_cell(t, 1, 2, "Tác nhân AI chỉ nhìn thấy câu hỏi, lịch sử liên quan và các kết quả truy hồi trong từng lượt xử lý.")
    set_cell(t, 5, 2, "Các hành động của tác nhân là những lựa chọn rời rạc như tra cứu, đối chiếu, tổng hợp hoặc từ chối.")
    set_cell(t, 6, 0, "Số tác nhân")
    set_cell(t, 6, 1, "Một tác nhân chính")
    set_cell(t, 6, 2, "Một tác nhân chịu trách nhiệm điều phối quá trình trả lời cho từng yêu cầu độc lập.")
    set_cell(t, 7, 2, "Môi trường gồm người dùng, kho dữ liệu pháp luật, nguồn đối chiếu bên ngoài và LLMs.")

    # Bảng 2.5: trạng thái xử lý ở mức báo cáo.
    t = doc.tables[4]
    set_cell(t, 0, 0, "Nhóm dữ liệu")
    set_cell(t, 0, 1, "Dạng thông tin")
    set_cell(t, 0, 2, "Được cập nhật trong bước")
    set_cell(t, 0, 3, "Mục đích")
    rows = [
        ("Lịch sử xử lý", "Danh sách thông điệp", "Toàn bộ pipeline", "Lưu câu hỏi, kết quả truy hồi và các phản hồi trung gian."),
        ("Kết quả phân tích truy vấn", "Thông tin có cấu trúc", "Phân tích truy vấn", "Làm rõ ý định, phạm vi pháp lý và hướng truy hồi phù hợp."),
        ("Tài liệu truy hồi", "Danh sách căn cứ", "Suy luận và tra cứu", "Lưu các đoạn văn bản được dùng làm bằng chứng cho câu trả lời."),
        ("Số vòng suy luận", "Số nguyên", "Suy luận", "Giới hạn số lần tra cứu bổ sung để kiểm soát độ trễ."),
        ("Kết quả kiểm soát đầu vào", "Đúng/Sai", "Kiểm soát đầu vào", "Cho biết câu hỏi có được tiếp tục xử lý hay phải từ chối."),
        ("Lý do từ chối", "Văn bản ngắn", "Kiểm soát đầu vào", "Giải thích lý do hệ thống không trả lời một yêu cầu ngoài phạm vi."),
    ]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell(t, r, c, value)

    # Bảng 2.6.
    t = doc.tables[5]
    set_cell(t, 0, 1, "Thông tin")
    set_cell(t, 0, 2, "Giá trị ví dụ")
    set_cell(t, 1, 1, "Câu hỏi người dùng")
    set_cell(t, 2, 0, "Đầu ra hợp lệ")
    set_cell(t, 2, 1, "Kết quả kiểm tra")
    set_cell(t, 2, 2, "Câu hỏi thuộc phạm vi tư vấn pháp luật.")
    set_cell(t, 3, 0, "Đầu ra bị từ chối")
    set_cell(t, 3, 1, "Kết quả kiểm tra")
    set_cell(t, 3, 2, "Câu hỏi nằm ngoài phạm vi tư vấn pháp luật.")
    set_cell(t, 4, 1, "Lý do")
    set_cell(t, 4, 2, "Yêu cầu không liên quan đến pháp luật Việt Nam.")
    set_cell(t, 5, 1, "Phản hồi")
    set_cell(t, 5, 2, "Hệ thống giải thích ngắn gọn lý do từ chối và hướng người dùng quay lại phạm vi phù hợp.")

    # Bảng 2.7.
    t = doc.tables[6]
    set_cell(t, 0, 1, "Thông tin")
    set_cell(t, 1, 1, "Câu hỏi hiện tại")
    set_cell(t, 2, 1, "Ngữ cảnh hội thoại nếu có")
    set_cell(t, 3, 1, "Kết quả phân tích")
    set_cell(t, 3, 2, "Ý định, chủ đề pháp lý, phạm vi tìm kiếm và truy vấn đã được tối ưu.")
    set_cell(t, 4, 1, "Thông tin chuyển tiếp")
    set_cell(t, 4, 2, "Phân tích được chuyển sang bước suy luận để định hướng tra cứu.")

    # Bảng 2.8.
    t = doc.tables[7]
    set_cell(t, 0, 1, "Thông tin")
    set_cell(t, 1, 1, "Ngữ cảnh xử lý")
    set_cell(t, 1, 2, "Câu hỏi, phân tích truy vấn và kết quả tra cứu trước đó.")
    set_cell(t, 2, 1, "Số vòng suy luận")
    set_cell(t, 3, 0, "Đầu ra tiếp tục")
    set_cell(t, 3, 1, "Yêu cầu tra cứu bổ sung")
    set_cell(t, 3, 2, "Hệ thống cần thêm bằng chứng trước khi trả lời.")
    set_cell(t, 4, 0, "Đầu ra tiếp tục")
    set_cell(t, 4, 1, "Số vòng suy luận")
    set_cell(t, 4, 2, "Tăng lên để kiểm soát giới hạn lặp.")
    set_cell(t, 5, 0, "Đầu ra kết thúc")
    set_cell(t, 5, 1, "Câu trả lời dự thảo")
    set_cell(t, 5, 2, "Câu trả lời được tổng hợp từ các căn cứ đã thu thập.")

    # Bảng 2.9.
    t = doc.tables[8]
    set_cell(t, 1, 1, "Câu trả lời dự thảo")
    set_cell(t, 2, 1, "Các căn cứ đã truy hồi")
    set_cell(t, 2, 2, "Nguồn thông tin được dùng để đối chiếu nội dung trả lời.")
    set_cell(t, 3, 0, "Đầu ra đạt")
    set_cell(t, 3, 1, "Câu trả lời giữ nguyên")
    set_cell(t, 3, 2, "Câu trả lời đã có đủ căn cứ để trả về người dùng.")
    set_cell(t, 4, 0, "Đầu ra cần sửa")
    set_cell(t, 4, 1, "Câu trả lời đã hiệu chỉnh")
    set_cell(t, 4, 2, "Loại bỏ hoặc sửa phần chưa có đủ căn cứ.")

    # Một số bảng sau vẫn giữ thuật ngữ cài đặt; chuyển về cách gọi trong báo cáo.
    set_cell(doc.tables[0], 5, 5, "Tra cứu văn bản và điều khoản pháp luật")

    t = doc.tables[3]
    set_cell(t, 2, 0, "MongoDB")
    set_cell(t, 2, 2, "Đọc thông qua lớp dịch vụ pháp luật của Main Service")
    set_cell(t, 2, 3, "Lưu toàn văn, metadata và hỗ trợ tìm kiếm, thống kê, hiển thị chi tiết nguồn")

    t = doc.tables[11]
    set_cell(t, 0, 0, "Tham số")
    set_cell(t, 1, 0, "Số kết quả truy hồi ban đầu")
    set_cell(t, 1, 2, "Số kết quả từ kho vector được chuyển sang bước xếp hạng lại")
    set_cell(t, 2, 0, "Số ứng viên xếp hạng lại")
    set_cell(t, 2, 2, "Số đoạn tối đa được đánh giá sâu ở giai đoạn hai")

    t = doc.tables[12]
    set_cell(t, 0, 0, "Loại sự kiện")
    set_cell(t, 0, 2, "Ý nghĩa")
    set_cell(t, 1, 0, "Đang xử lý")
    set_cell(t, 1, 1, "Danh sách bước xử lý")
    set_cell(t, 1, 2, "Cập nhật vùng hiển thị tiến trình theo các bước: chuẩn bị, tra cứu, tổng hợp và kiểm chứng")
    set_cell(t, 2, 0, "Hoàn tất")
    set_cell(t, 2, 1, "Câu trả lời và nguồn trích dẫn")
    set_cell(t, 2, 2, "Kết thúc pipeline và hiển thị câu trả lời cuối cùng")

    t = doc.tables[13]
    set_cell(t, 0, 0, "Sự kiện")
    set_cell(t, 0, 3, "Vai trò trên giao diện")
    set_cell(t, 2, 0, "Cập nhật tiến trình")
    set_cell(t, 2, 1, "Khi pipeline chuyển bước")
    set_cell(t, 2, 2, "Danh sách các bước xử lý")
    set_cell(t, 2, 3, "Cập nhật vùng tiến trình: kiểm tra, phân tích, tra cứu, tổng hợp và kiểm chứng")
    set_cell(t, 3, 0, "Hoàn tất")
    set_cell(t, 3, 1, "Sau khi câu trả lời đã được lưu")
    set_cell(t, 3, 2, "Câu trả lời, nguồn trích dẫn và thông tin phụ trợ")
    set_cell(t, 3, 3, "Hiển thị kết quả cuối cùng cho người dùng")

    t = doc.tables[14]
    set_cell(t, 1, 0, "Chia sẻ logic nghiệp vụ")
    set_cell(t, 1, 1, "Các lớp truy cập API, ánh xạ dữ liệu, kiểm tra hợp lệ và quản lý trạng thái dùng chung giữa Android/iOS")
    set_cell(t, 2, 0, "Trải nghiệm native")
    set_cell(t, 2, 1, "Giao diện Compose chạy trực tiếp trên nền tảng, không phụ thuộc WebView")

    t = doc.tables[15]
    set_cell(t, 4, 0, "Tư vấn có hướng dẫn")
    set_cell(t, 4, 1, "Câu hỏi ban đầu, câu hỏi làm rõ, lựa chọn của người dùng và luồng trả lời")
    set_cell(t, 4, 2, "Bắt đầu tư vấn, chọn đáp án, gửi yêu cầu trả lời")
    set_cell(t, 4, 3, "Hiển thị kết quả hoặc lỗi")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
