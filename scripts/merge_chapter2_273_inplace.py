from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX = Path(__file__).resolve().parents[1] / "02_Chuong_2_ban_bo_sung_phan_tich_rag.docx"

MERGED_PARAGRAPH = (
    "Trong cùng giai đoạn xử lý PDF, một số tác vụ không phụ thuộc trực tiếp vào nhau được tổ chức song song, "
    "chẳng hạn lưu file gốc để đối chiếu và trích xuất nội dung phục vụ chuẩn hóa. File gốc được giữ lại để "
    "quản trị viên có thể kiểm tra nguồn khi cần, trong khi phần văn bản trích xuất trở thành đầu vào cho bước "
    "nhận diện cấu trúc điều khoản. Cách tổ chức này giúp rút ngắn thời gian chờ nhưng không cần tách thành "
    "một nhánh nghiệp vụ riêng, vì mục tiêu chính của giai đoạn này vẫn là biến tài liệu PDF thành dữ liệu pháp "
    "luật có cấu trúc và có thể kiểm tra."
)


def insert_paragraph_after(paragraph, text, style="Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc, prefix):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Không tìm thấy đoạn bắt đầu bằng: {prefix}")


def main():
    doc = Document(DOCX)

    idx_272 = find_paragraph(doc, "2.7.2.")
    idx_273 = find_paragraph(doc, "2.7.3.")
    idx_274 = find_paragraph(doc, "2.7.4.")

    # Insert the useful operational idea from 2.7.3 into 2.7.2, before the
    # validation paragraph, so the narrative stays in one coherent subsection.
    validation_idx = None
    for i in range(idx_272 + 1, idx_273):
        if doc.paragraphs[i].text.strip().startswith("Kết quả chuẩn hoá không được lưu ngay"):
            validation_idx = i
            break
    if validation_idx is None:
        raise ValueError("Không tìm thấy đoạn kiểm tra hợp lệ trong mục 2.7.2")

    insert_paragraph_after(doc.paragraphs[validation_idx - 1], MERGED_PARAGRAPH)

    # Refresh indexes after insertion, then remove the separate 2.7.3 section.
    idx_273 = find_paragraph(doc, "2.7.3.")
    idx_274 = find_paragraph(doc, "2.7.4.")
    for paragraph in list(doc.paragraphs[idx_273:idx_274]):
        remove_paragraph(paragraph)

    # Renumber the next subsection to keep the chapter sequence continuous.
    idx_274 = find_paragraph(doc, "2.7.4.")
    doc.paragraphs[idx_274].text = "2.7.3. Cơ chế hoàn tác để giữ dữ liệu nhất quán"

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
