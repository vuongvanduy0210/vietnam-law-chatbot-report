from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUT = Path("Bao_Cao/kich_ban_luyen_noi_bao_cao_do_an.html")
OUT_DOCX = Path("Bao_Cao/kich_ban_luyen_noi_bao_cao_do_an.docx")


slides = [
    {
        "no": 1,
        "title": "Mở đầu và giới thiệu đề tài",
        "time": "25-30 giây",
        "talk": [
            "Em kính chào quý thầy cô trong hội đồng. Em là Vương Văn Duy, sinh viên thực hiện đồ án tốt nghiệp với đề tài: “Nghiên cứu phát triển trợ lý ảo pháp luật cho chuyển đổi số”. Đề tài được thực hiện dưới sự hướng dẫn của thầy ThS. Trần Đức Thịnh.",
            "Trong phần trình bày hôm nay, em sẽ tập trung vào lý do lựa chọn bài toán, cơ sở lý thuyết, phương pháp xây dựng hệ thống, các chức năng chính và kết quả thực nghiệm.",
        ],
        "note": "Nói chậm, rõ tên đề tài. Không cần giải thích kỹ ngay ở slide đầu.",
    },
    {
        "no": 2,
        "title": "Nội dung trình bày",
        "time": "25-30 giây",
        "talk": [
            "Bài báo cáo của em được chia thành bốn phần chính. Phần thứ nhất là phát biểu bài toán, trình bày lý do chọn đề tài và mục tiêu cần đạt được. Phần thứ hai là cơ sở lý thuyết, tập trung vào chatbot, LLMs, AI Agent và LangGraph.",
            "Phần thứ ba là nội dung trọng tâm: xây dựng hệ thống, trong đó em trình bày môi trường của agent, mô hình tổng quan, graph xử lý, các node, hệ thống RAG nội bộ và luồng cập nhật văn bản pháp luật. Phần cuối cùng là thực nghiệm và kết luận.",
        ],
    },
    {
        "no": 3,
        "title": "Chuyển sang phần 1 - Phát biểu bài toán",
        "time": "10-15 giây",
        "talk": [
            "Trước hết, em xin trình bày phần phát biểu bài toán để làm rõ vì sao đề tài này cần thiết và hệ thống hướng tới giải quyết vấn đề gì.",
        ],
    },
    {
        "no": 4,
        "title": "Lý do chọn đề tài",
        "time": "45-55 giây",
        "talk": [
            "Trong thực tế, nhu cầu tra cứu và tiếp cận thông tin pháp luật của người dân là rất lớn. Tuy nhiên, văn bản pháp luật thường có khối lượng lớn, cấu trúc dài, nhiều điều khoản và cách diễn đạt mang tính chuyên môn. Vì vậy, người dùng phổ thông thường gặp khó khăn khi tự tìm kiếm hoặc tự xác định quy định nào phù hợp với tình huống của mình.",
            "Một khó khăn khác là pháp luật thay đổi theo thời gian. Một văn bản có thể được sửa đổi, bổ sung hoặc thay thế bởi văn bản mới hơn. Nếu hệ thống chỉ dựa vào dữ liệu tĩnh hoặc chỉ dựa vào khả năng sinh câu trả lời của mô hình ngôn ngữ, câu trả lời có thể thiếu căn cứ hoặc không còn phù hợp với hiệu lực hiện tại.",
            "Từ thực tế đó, em lựa chọn xây dựng một trợ lý ảo pháp luật có khả năng tiếp nhận câu hỏi tự nhiên, truy hồi nguồn pháp luật nội bộ, đối chiếu nguồn cập nhật và đưa ra câu trả lời có căn cứ.",
        ],
        "transition": "Từ lý do này, slide tiếp theo trình bày mục tiêu cụ thể của hệ thống.",
    },
    {
        "no": 5,
        "title": "Mục tiêu đề tài",
        "time": "40-50 giây",
        "talk": [
            "Mục tiêu của đề tài là xây dựng một hệ thống chatbot pháp luật có thể hỗ trợ người dùng tra cứu và tư vấn pháp luật bằng ngôn ngữ tự nhiên. Hệ thống không chỉ sinh câu trả lời, mà cần đưa ra câu trả lời có căn cứ, có nguồn tham chiếu và phù hợp với ngữ cảnh câu hỏi.",
            "Về mặt dữ liệu, đề tài hướng tới tổ chức kho tri thức pháp luật để hệ thống có thể tìm kiếm, truy hồi và cập nhật dữ liệu. Về mặt sản phẩm, hệ thống gồm ứng dụng Mobile cho người dùng cuối và Web Admin cho quản trị viên cập nhật, quản lý nguồn dữ liệu pháp luật.",
            "Trọng tâm kỹ thuật của đồ án là kết hợp AI Agent, RAG, LangGraph và cơ chế kiểm chứng để giảm rủi ro trả lời sai hoặc bịa đặt.",
        ],
    },
    {
        "no": 6,
        "title": "Chuyển sang phần 2 - Cơ sở lý thuyết",
        "time": "10-15 giây",
        "talk": [
            "Sau khi xác định bài toán và mục tiêu, em xin trình bày các cơ sở lý thuyết được sử dụng để xây dựng hệ thống.",
        ],
    },
    {
        "no": 7,
        "title": "Khái niệm Chatbot",
        "time": "35-45 giây",
        "talk": [
            "Chatbot là hệ thống phần mềm mô phỏng hội thoại với người dùng thông qua văn bản hoặc giọng nói. Trong các hệ thống hiện đại, chatbot không chỉ phản hồi theo kịch bản có sẵn mà còn có thể kết hợp với mô hình ngôn ngữ lớn để phân tích câu hỏi và sinh phản hồi linh hoạt.",
            "Đối với bài toán pháp luật, chatbot cần đáp ứng yêu cầu cao hơn so với hội thoại thông thường. Câu trả lời phải đúng với quy định, có căn cứ nguồn và cần hạn chế tối đa hiện tượng mô hình tự suy diễn nội dung không có trong dữ liệu.",
        ],
    },
    {
        "no": 8,
        "title": "Chatbot dựa trên tập luật",
        "time": "35-45 giây",
        "talk": [
            "Cách tiếp cận đầu tiên là chatbot dựa trên tập luật. Hệ thống sẽ kiểm tra câu hỏi đầu vào với các kịch bản hoặc tập luật đã được chuẩn bị sẵn, sau đó trả lời theo mẫu tương ứng.",
            "Ưu điểm của cách này là dễ kiểm soát và phù hợp với các tình huống cố định. Tuy nhiên, nhược điểm là thiếu linh hoạt. Khi người dùng hỏi khác kịch bản, diễn đạt tự nhiên hơn hoặc đặt câu hỏi phức tạp hơn, hệ thống rất khó xử lý. Vì vậy, cách tiếp cận này chưa phù hợp với bài toán tư vấn pháp luật mở.",
        ],
    },
    {
        "no": 9,
        "title": "Chatbot dựa trên LLMs",
        "time": "45-55 giây",
        "talk": [
            "Cách tiếp cận thứ hai là sử dụng LLMs. LLMs có khả năng hiểu ngôn ngữ tự nhiên, nắm bắt ngữ cảnh và sinh câu trả lời linh hoạt hơn rất nhiều so với chatbot tập luật.",
            "Tuy nhiên, trong bài toán pháp luật, nếu chỉ dùng LLMs thuần túy thì có một rủi ro lớn là hiện tượng ảo giác. Mô hình có thể tạo ra nội dung nghe có vẻ hợp lý nhưng không có căn cứ trong văn bản pháp luật, hoặc không cập nhật với quy định mới.",
            "Vì vậy, đồ án không sử dụng LLMs như một thành phần trả lời độc lập, mà đặt LLMs trong kiến trúc có truy hồi dữ liệu và kiểm chứng.",
        ],
    },
    {
        "no": 10,
        "title": "Chatbot dựa trên AI Agent",
        "time": "50-60 giây",
        "talk": [
            "Cách tiếp cận thứ ba, cũng là hướng chính của đồ án, là xây dựng chatbot theo kiến trúc AI Agent. Thay vì chỉ sinh câu trả lời trực tiếp, agent có thể phân tích yêu cầu, lập kế hoạch xử lý, gọi công cụ truy hồi dữ liệu, đối chiếu nguồn và sau đó mới tổng hợp phản hồi.",
            "Cách tiếp cận này phù hợp với bài toán pháp luật vì câu trả lời thường cần nhiều bước: hiểu câu hỏi, tìm quy định liên quan, kiểm tra văn bản còn hiệu lực hay không, rồi mới diễn giải cho người dùng. Do đó, AI Agent giúp hệ thống chủ động hơn và kiểm soát được quá trình xử lý tốt hơn.",
        ],
    },
    {
        "no": 11,
        "title": "AI Agent",
        "time": "35-45 giây",
        "talk": [
            "AI Agent có thể hiểu là một tác tử thông minh có khả năng nhận yêu cầu, quan sát trạng thái hiện tại, lựa chọn hành động và sử dụng công cụ để đạt mục tiêu. Trong hệ thống của em, LLMs đóng vai trò bộ phận suy luận của agent.",
            "Điểm quan trọng là agent không chỉ trả lời dựa trên kiến thức đã học sẵn. Agent được cấp các công cụ như truy hồi cơ sở dữ liệu pháp luật nội bộ và tìm kiếm nguồn pháp luật cập nhật. Nhờ đó, câu trả lời được xây dựng dựa trên bằng chứng thay vì chỉ dựa trên suy đoán.",
        ],
    },
    {
        "no": 12,
        "title": "LangGraph",
        "time": "40-50 giây",
        "talk": [
            "Để tổ chức luồng xử lý của agent, đồ án sử dụng LangGraph. LangGraph cho phép mô hình hóa hệ thống thành graph gồm các node, edge và state.",
            "Node tương ứng với từng bước xử lý như kiểm tra đầu vào, phân tích câu hỏi, agent suy luận, gọi công cụ và kiểm chứng câu trả lời. Edge thể hiện hướng di chuyển giữa các node. State lưu thông tin xuyên suốt quá trình xử lý, ví dụ câu hỏi, lịch sử hội thoại, kết quả truy hồi và số vòng lặp xử lý.",
            "Việc dùng LangGraph giúp luồng agent rõ ràng hơn, dễ kiểm soát điều kiện chuyển node và tránh việc LLM hoạt động hoàn toàn tự do.",
        ],
    },
    {
        "no": 13,
        "title": "Chuyển sang phần 3 - Xây dựng hệ thống",
        "time": "10-15 giây",
        "talk": [
            "Tiếp theo là phần trọng tâm của bài báo cáo: cách xây dựng hệ thống trợ lý ảo pháp luật.",
        ],
    },
    {
        "no": 14,
        "title": "Xác định môi trường của agent",
        "time": "55-65 giây",
        "talk": [
            "Trước khi xây dựng agent, cần xác định rõ môi trường mà agent hoạt động và tập hành động mà agent được phép thực hiện. Với bài toán của em, môi trường gồm ba nhóm nguồn chính.",
            "Thứ nhất là cơ sở dữ liệu pháp luật nội bộ, tức kho dữ liệu đã được chuẩn hóa, chunking, vector hóa và lưu phục vụ truy hồi. Thứ hai là nguồn pháp luật cập nhật từ bên ngoài, ví dụ các website tra cứu văn bản pháp luật. Thứ ba là nguồn tri thức mới do quản trị viên đưa vào thông qua chức năng upload PDF trên Web Admin.",
            "Tập hành động của agent gồm trả lời câu hỏi, truy hồi quy định pháp luật, kiểm chứng và trích dẫn nguồn, đồng thời đối chiếu nguồn cập nhật khi cần. Cách xác định này giúp agent hoạt động trong phạm vi rõ ràng và phù hợp với bài toán pháp luật.",
        ],
    },
    {
        "no": 15,
        "title": "Mô hình tổng quan chatbot",
        "time": "55-65 giây",
        "talk": [
            "Slide này mô tả mô hình tổng quan của hệ thống chatbot. Người dùng đặt câu hỏi từ giao diện ứng dụng. Câu hỏi được đưa vào AI Agent để phân tích và điều phối công cụ.",
            "Agent có thể sử dụng công cụ truy hồi pháp luật nội bộ từ ChromaDB. Dữ liệu trong ChromaDB không phải là toàn bộ văn bản thô, mà là các vector chunks được tạo từ văn bản pháp luật đã chuẩn hóa. Ngoài ra, hệ thống có PostgreSQL để lưu các dữ liệu nghiệp vụ như người dùng, lịch sử hội thoại và phiên làm việc.",
            "Sau khi thu thập đủ thông tin, LLMs tổng hợp câu trả lời và hệ thống trả lại kết quả cho người dùng kèm căn cứ liên quan.",
        ],
    },
    {
        "no": 16,
        "title": "Luồng xử lý tin nhắn người dùng",
        "time": "60-70 giây",
        "talk": [
            "Ở mức chi tiết hơn, khi người dùng gửi một tin nhắn mới, hệ thống trước hết lấy lịch sử hội thoại từ PostgreSQL để bổ sung ngữ cảnh. Sau đó câu hỏi và context hội thoại được đưa vào luồng xử lý của agent.",
            "LLMs thực hiện bước phân tích câu hỏi, từ đó xác định cần truy hồi dữ liệu nào. Nếu cần, agent điều khiển các tool để thực hiện task, ví dụ truy hồi dữ liệu pháp luật nội bộ hoặc tìm kiếm nguồn cập nhật. Kết quả từ công cụ được đưa ngược lại cho agent để tiếp tục suy luận.",
            "Trước khi trả lời, câu trả lời được chuyển qua bước verifier để kiểm chứng lại. Đây là điểm quan trọng giúp giảm rủi ro trả lời bịa đặt hoặc trích dẫn không đúng nguồn.",
        ],
    },
    {
        "no": 17,
        "title": "Đồ thị xử lý của hệ thống",
        "time": "45-55 giây",
        "talk": [
            "Slide này biểu diễn graph xử lý của hệ thống. Luồng bắt đầu từ node chuẩn bị context, sau đó đi qua guardrail, query analysis, agent, tool execution và verifier.",
            "Điểm quan trọng của graph là mỗi bước có trách nhiệm riêng, không gom toàn bộ xử lý vào một prompt duy nhất. Graph cũng có các cạnh điều kiện: nếu câu hỏi không hợp lệ thì dừng sớm; nếu agent cần thêm bằng chứng thì quay lại gọi tool; nếu đã đủ điều kiện thì chuyển sang verifier.",
            "Nhờ cách tổ chức này, hệ thống có khả năng kiểm soát tốt hơn số vòng lặp, luồng gọi công cụ và điều kiện tạo câu trả lời cuối cùng.",
        ],
    },
    {
        "no": 18,
        "title": "Node Guardrail",
        "time": "45-55 giây",
        "talk": [
            "Node đầu tiên là Guardrail. Nhiệm vụ của node này là kiểm tra câu hỏi đầu vào, loại bỏ các yêu cầu không phù hợp hoặc nằm ngoài phạm vi tư vấn pháp luật.",
            "Trong hệ thống, prompt của Guardrail đóng vai trò như một bộ lọc kiểm duyệt nghiêm ngặt. Nếu câu hỏi hợp lệ, luồng xử lý tiếp tục sang bước phân tích truy vấn. Nếu câu hỏi không phù hợp, hệ thống có thể dừng sớm và phản hồi phù hợp cho người dùng.",
            "Việc tách riêng Guardrail giúp hệ thống tránh đưa các câu hỏi không liên quan vào pipeline RAG, từ đó giảm chi phí xử lý và giảm rủi ro sinh phản hồi sai phạm vi.",
        ],
    },
    {
        "no": 19,
        "title": "Node Query Analysis",
        "time": "50-60 giây",
        "talk": [
            "Sau Guardrail là node Query Analysis. Node này phân tích câu hỏi của người dùng để xác định vấn đề pháp lý chính, thuật ngữ pháp lý liên quan và tạo truy vấn tối ưu cho từng nguồn dữ liệu.",
            "Trong hệ thống, truy vấn dùng cho ChromaDB và truy vấn dùng cho nguồn web được tách riêng. Điều này quan trọng vì truy hồi vector nội bộ và tìm kiếm web có đặc điểm khác nhau. Truy vấn nội bộ cần bám vào thuật ngữ pháp lý, còn truy vấn web cần phù hợp với cách tìm kiếm văn bản cập nhật.",
            "Kết quả của node này không phải là câu trả lời cuối cùng, mà là đầu vào có cấu trúc để Agent sử dụng trong các bước tiếp theo.",
        ],
    },
    {
        "no": 20,
        "title": "Node Agent",
        "time": "55-65 giây",
        "talk": [
            "Node Agent là thành phần suy luận trung tâm. Agent nhận kết quả phân tích truy vấn, lịch sử hội thoại và các kết quả tool nếu đã có, sau đó quyết định bước tiếp theo.",
            "Agent có thể gọi tool truy hồi pháp luật nội bộ, gọi tool tìm kiếm nguồn cập nhật hoặc tổng hợp kết quả nếu đã có đủ bằng chứng tối thiểu. Trong graph của hệ thống, việc chuyển sang verifier không hoàn toàn do LLM tự quyết định, mà còn được kiểm soát bằng điều kiện graph: agent cần gọi đủ các nhóm công cụ cần thiết hoặc đạt giới hạn vòng lặp.",
            "Cách thiết kế này giúp cân bằng giữa khả năng suy luận linh hoạt của LLMs và sự kiểm soát cần thiết trong một hệ thống tư vấn pháp luật.",
        ],
    },
    {
        "no": 21,
        "title": "Node Verifier",
        "time": "45-55 giây",
        "talk": [
            "Node Verifier là bước kiểm chứng cuối cùng trước khi trả lời người dùng. Node này nhận câu trả lời do Agent tạo ra cùng toàn bộ kết quả công cụ đã sử dụng.",
            "Verifier kiểm tra xem câu trả lời có bịa đặt hay không, có bám vào bằng chứng truy hồi hay không, và có cần chỉnh sửa phần nào chưa đủ căn cứ hay không. Nếu phát hiện nội dung chưa phù hợp, node này có thể tự sửa đổi và cập nhật câu trả lời.",
            "Đây là lớp bảo vệ quan trọng của hệ thống, đặc biệt trong lĩnh vực pháp luật, nơi câu trả lời cần có căn cứ và không nên suy diễn tùy tiện.",
        ],
    },
    {
        "no": 22,
        "title": "Xây dựng dữ liệu pháp luật cho RAG",
        "time": "60-70 giây",
        "talk": [
            "Phần tiếp theo là xây dựng hệ thống truy vấn pháp luật nội bộ theo hướng RAG. Bước đầu tiên là xây dựng dữ liệu pháp luật.",
            "Dữ liệu ban đầu được chuẩn hóa theo đơn vị điều luật thay vì lưu toàn bộ văn bản thành một khối lớn. Mỗi điều luật được tổ chức với các trường chính như law_id, article_id, title, text, year, topics, keywords và summary.",
            "Việc chuẩn hóa này rất quan trọng vì văn bản pháp luật có cấu trúc đặc thù. Nếu lưu dữ liệu quá thô, hệ thống khó truy hồi đúng điều khoản. Nếu chỉ lưu văn bản không có metadata, hệ thống cũng khó lọc theo năm, chủ đề hoặc nguồn văn bản.",
            "Sau khi dữ liệu được làm sạch và gắn metadata, dữ liệu sẽ được đưa sang bước chunking và vector hóa để lưu vào ChromaDB.",
        ],
    },
    {
        "no": 23,
        "title": "Chunking, Embedding và lưu trữ ChromaDB",
        "time": "70-80 giây",
        "talk": [
            "Ở bước này, hệ thống áp dụng chiến lược chunking lai. Những điều luật ngắn được giữ nguyên để tránh làm mất ngữ cảnh. Chỉ những điều luật dài mới được chia nhỏ, với ngưỡng tối đa khoảng 1000 từ và overlap khoảng 150 từ.",
            "Sau khi chunking, mỗi chunk được vector hóa bằng embedding model tiếng Việt. Embedding được dùng để tìm kiếm tương đồng ngữ nghĩa giữa câu hỏi của người dùng và nội dung pháp luật.",
            "Trong ChromaDB, mỗi bản ghi gồm id, document, embedding và metadata. Trong đó embedding là vector số phục vụ tìm kiếm; document là đoạn văn bản pháp luật được trả về làm context; metadata vẫn là dữ liệu dạng text hoặc scalar để phục vụ lọc, hiển thị và trích dẫn.",
            "Cách lưu trữ này giúp hệ thống vừa tìm kiếm được theo ngữ nghĩa, vừa giữ lại thông tin cần thiết để tạo câu trả lời có nguồn.",
        ],
    },
    {
        "no": 24,
        "title": "Retrieve, Rerank và đóng gói Context",
        "time": "70-80 giây",
        "talk": [
            "Khi người dùng đặt câu hỏi, hệ thống không tìm kiếm trực tiếp bằng câu hỏi gốc, mà sử dụng internal_search_query đã được tối ưu từ node Query Analysis.",
            "Truy vấn này được encode thành vector, sau đó tìm kiếm trong ChromaDB để lấy ra các candidate chunks có độ tương đồng cao. Tuy nhiên, vector search chỉ là bước lọc đầu tiên, nên hệ thống tiếp tục dùng cross-encoder rerank để chấm lại mức độ phù hợp giữa truy vấn và từng chunk.",
            "Sau rerank, hệ thống chọn top documents phù hợp nhất và đóng gói thành formatted context để đưa vào Agent. Nhờ vậy, LLMs không phải tự nhớ luật, mà trả lời dựa trên các đoạn pháp luật đã được truy hồi.",
            "Trong hệ thống còn có bước phát hiện văn bản cũ mới theo metadata để cảnh báo khả năng có văn bản đã lỗi thời, sau đó Agent và Verifier tiếp tục kiểm tra lại bằng nguồn cập nhật.",
        ],
    },
    {
        "no": 25,
        "title": "Cập nhật văn bản pháp luật từ Web Admin",
        "time": "60-70 giây",
        "talk": [
            "Bên cạnh kho dữ liệu ban đầu, hệ thống có chức năng cập nhật văn bản pháp luật mới từ Web Admin. Quản trị viên có thể upload file PDF văn bản pháp luật mới.",
            "Hệ thống tạo task xử lý để theo dõi tiến trình. File PDF được trích xuất nội dung bằng OCR hoặc parser, sau đó LLMs hỗ trợ cấu trúc, làm sạch và gắn metadata cho dữ liệu.",
            "Dữ liệu sau xử lý được đưa vào MongoDB để phục vụ tra cứu văn bản và đưa vào ChromaDB để phục vụ RAG. Như vậy, luồng cập nhật dữ liệu không chỉ lưu văn bản để hiển thị, mà còn làm mới kho tri thức mà Agent sử dụng khi trả lời câu hỏi.",
        ],
    },
    {
        "no": 26,
        "title": "Mô hình tổng quan hệ thống",
        "time": "50-60 giây",
        "talk": [
            "Slide này mô tả kiến trúc tổng quan của toàn hệ thống. Phía người dùng có Mobile App để chat, tư vấn có hướng dẫn và tra cứu văn bản. Phía quản trị có Admin Web để cập nhật và quản lý dữ liệu pháp luật.",
            "Backend Main Service xử lý các nghiệp vụ chính như người dùng, xác thực, lịch sử hội thoại và quản lý dữ liệu. Chatbot backend hay RAG Service chịu trách nhiệm xử lý pipeline AI Agent và truy hồi dữ liệu pháp luật.",
            "Hệ thống sử dụng nhiều loại cơ sở dữ liệu theo vai trò khác nhau: PostgreSQL cho dữ liệu nghiệp vụ và lịch sử hội thoại, MongoDB cho dữ liệu văn bản pháp luật, ChromaDB cho vector chunks phục vụ RAG, và LLMs API cho suy luận ngôn ngữ.",
        ],
    },
    {
        "no": 27,
        "title": "Các chức năng của hệ thống",
        "time": "45-55 giây",
        "talk": [
            "Về hình thức sản phẩm, hệ thống gồm Mobile App dành cho người dùng cuối và Web Admin dành cho quản trị viên.",
            "Các chức năng chính của người dùng gồm chat tư vấn pháp luật, tư vấn có hướng dẫn, tra cứu văn bản pháp luật, quản lý hội thoại và xem nguồn tham chiếu. Phía quản trị viên có các chức năng cập nhật văn bản pháp luật mới từ PDF và quản lý kho tri thức phục vụ RAG.",
            "Trong đó, các chức năng trọng tâm của đồ án là luồng chat chính với LangGraph, hệ thống RAG nội bộ, luồng tư vấn có hướng dẫn và luồng cập nhật văn bản pháp luật mới.",
        ],
    },
    {
        "no": 28,
        "title": "Usecase tổng quát",
        "time": "35-45 giây",
        "talk": [
            "Usecase tổng quát cho thấy hai nhóm tác nhân chính là người dùng và quản trị viên. Người dùng tương tác với hệ thống thông qua các chức năng hỏi đáp pháp luật, tư vấn theo hướng dẫn, tra cứu văn bản và quản lý hội thoại.",
            "Quản trị viên thực hiện các chức năng quản lý dữ liệu, upload văn bản pháp luật mới và theo dõi quá trình xử lý dữ liệu. Các usecase này phản ánh phạm vi chức năng chính mà hệ thống đã triển khai.",
        ],
    },
    {
        "no": 29,
        "title": "Chuyển sang phần thực nghiệm và kết luận",
        "time": "10-15 giây",
        "talk": [
            "Sau phần xây dựng hệ thống, em xin chuyển sang phần thực nghiệm và kết luận để đánh giá kết quả đạt được.",
        ],
    },
    {
        "no": 30,
        "title": "Demo các chức năng",
        "time": "20-30 giây",
        "talk": [
            "Trong phần demo, em tập trung minh họa các chức năng chính của hệ thống, bao gồm chat tư vấn pháp luật, tư vấn có hướng dẫn, tra cứu văn bản và cập nhật văn bản pháp luật từ Web Admin.",
            "Mục tiêu của phần demo là cho thấy hệ thống không chỉ tồn tại ở mức mô hình, mà đã được triển khai thành sản phẩm có giao diện người dùng và luồng xử lý hoàn chỉnh.",
        ],
    },
    {
        "no": 31,
        "title": "Bộ kiểm thử và phương pháp đánh giá",
        "time": "55-65 giây",
        "talk": [
            "Để đánh giá hệ thống, em xây dựng bộ kiểm thử gồm 100 câu hỏi pháp luật, chia thành hai nhóm. Nhóm N1 gồm 50 câu hỏi factual có đáp án xác định, dùng để đánh giá độ chính xác thông tin, điều khoản và nguồn trích dẫn. Nhóm N2 gồm 50 câu hỏi tình huống mở, dùng để đánh giá chất lượng tư vấn và khả năng diễn giải.",
            "Phương pháp đánh giá tập trung vào các chỉ số: Accuracy@1, Citation Accuracy, Temporal Conflict, Answer Quality và Response Latency. Các chỉ số này phản ánh cả độ đúng, chất lượng câu trả lời và hiệu năng phản hồi của hệ thống.",
        ],
    },
    {
        "no": 32,
        "title": "Công thức và ý nghĩa các chỉ số",
        "time": "55-65 giây",
        "talk": [
            "Accuracy@1 đo tỷ lệ câu trả lời đầu tiên chứa đúng thông tin kỳ vọng. Citation Accuracy đo tỷ lệ câu trả lời trích dẫn đúng nguồn văn bản pháp luật chứa đáp án.",
            "Answer Quality được đánh giá theo thang điểm 1 đến 5 với nhóm câu hỏi mở, dựa trên độ đúng pháp lý, độ đầy đủ, khả năng áp dụng vào tình huống và cách trình bày câu trả lời.",
            "Response Latency đo thời gian từ lúc gửi câu hỏi đến khi hệ thống hoàn tất phản hồi. Trong đó P50 là độ trễ trung vị, nghĩa là 50% câu hỏi có thời gian phản hồi thấp hơn hoặc bằng giá trị này. P95 phản ánh nhóm câu hỏi chậm hơn, tức 95% câu hỏi có thời gian phản hồi thấp hơn hoặc bằng giá trị đó.",
        ],
    },
    {
        "no": 33,
        "title": "Kết quả thực nghiệm",
        "time": "50-60 giây",
        "talk": [
            "Kết quả thực nghiệm cho thấy hệ thống đạt Accuracy@1 là 90%, Citation Accuracy là 56,7%, Answer Quality đạt 4,13 trên 5. Về độ trễ, Latency P50 là 65,5 giây và P95 là 100,6 giây.",
            "Từ kết quả này có thể thấy hệ thống đạt độ chính xác tốt ở nhóm câu hỏi factual và chất lượng câu trả lời ở mức khá tốt đối với nhóm câu hỏi tình huống. Tuy nhiên, hạn chế lớn nhất nằm ở thời gian phản hồi.",
            "Nguyên nhân chính là pipeline Agentic RAG gồm nhiều bước: phân tích câu hỏi, gọi công cụ, truy hồi, rerank, tổng hợp và verifier. Vì vậy, trong hướng phát triển, tối ưu độ trễ là một nhiệm vụ quan trọng.",
        ],
    },
    {
        "no": 34,
        "title": "Kết luận",
        "time": "55-65 giây",
        "talk": [
            "Về kết quả đạt được, đồ án đã tìm hiểu và ứng dụng các khái niệm về chatbot, LLMs, RAG và AI Agent để xây dựng hệ thống trợ lý ảo pháp luật. Hệ thống đã triển khai được Mobile App, Web Admin, kho dữ liệu pháp luật nội bộ và pipeline Agentic RAG có bước kiểm chứng.",
            "Về hạn chế, hệ thống vẫn còn độ trễ phản hồi khá cao do pipeline gồm nhiều bước xử lý. Bộ kiểm thử cũng còn giới hạn và chưa thể bao phủ toàn bộ các lĩnh vực pháp luật Việt Nam.",
            "Trong hướng phát triển, em định hướng tối ưu thời gian phản hồi, mở rộng và cập nhật kho dữ liệu pháp luật, cải thiện truy hồi và trích dẫn nguồn, đồng thời mở rộng thêm các chức năng hỗ trợ người dùng và quản trị viên.",
        ],
    },
    {
        "no": 35,
        "title": "Kết thúc",
        "time": "15-20 giây",
        "talk": [
            "Phần trình bày của em đến đây là kết thúc. Em xin cảm ơn quý thầy cô và các bạn đã lắng nghe. Em rất mong nhận được góp ý và câu hỏi từ hội đồng để hoàn thiện đề tài tốt hơn.",
        ],
    },
]


def paragraph(text: str, cls: str = "") -> str:
    cls_attr = f' class="{cls}"' if cls else ""
    return f"<p{cls_attr}>{escape(text)}</p>"


def build_html() -> str:
    parts: list[str] = []
    parts.append(
        """<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
body { font-family: "Times New Roman", serif; font-size: 13pt; line-height: 1.35; color: #111; }
h1 { font-size: 20pt; text-align: center; margin: 0 0 6pt; }
h2 { font-size: 16pt; color: #0C1D4B; border-bottom: 1px solid #0C1D4B; padding-bottom: 4pt; margin-top: 18pt; }
h3 { font-size: 14pt; color: #0C1D4B; margin: 14pt 0 4pt; }
.meta { text-align: center; font-style: italic; margin-bottom: 18pt; }
.lead { background: #F2F5FA; border-left: 4px solid #0C1D4B; padding: 8pt 10pt; margin: 10pt 0 14pt; }
.time { font-weight: bold; color: #555; }
.label { font-weight: bold; color: #0C1D4B; margin-bottom: 2pt; }
.talk { margin: 3pt 0 7pt; text-align: justify; }
.note { font-style: italic; color: #555; margin: 4pt 0 8pt; }
.transition { color: #333; margin: 4pt 0 8pt; }
ul { margin-top: 3pt; }
li { margin-bottom: 3pt; }
</style>
</head>
<body>
"""
    )
    parts.append("<h1>KỊCH BẢN LUYỆN NÓI BÁO CÁO ĐỒ ÁN TỐT NGHIỆP</h1>")
    parts.append('<p class="meta">Đề tài: Nghiên cứu phát triển trợ lý ảo pháp luật cho chuyển đổi số</p>')
    parts.append(
        '<div class="lead"><p><b>Mục tiêu tài liệu:</b> hỗ trợ luyện trình bày theo từng slide trong khoảng 12-15 phút. Nội dung được viết theo văn nói, bám theo slide hiện tại và các nội dung chính trong báo cáo: chatbot, AI Agent, LangGraph, Agentic RAG, xây dựng kho dữ liệu pháp luật, cập nhật văn bản từ Web Admin, thực nghiệm và kết luận.</p></div>'
    )
    parts.append("<h2>Gợi ý nhịp trình bày chung</h2>")
    parts.append("<ul>")
    for item in [
        "Không đọc nguyên văn toàn bộ chữ trên slide; dùng slide làm điểm tựa để nói mạch logic.",
        "Các slide 14-25 là phần trọng tâm kỹ thuật, nên nói rõ hơn và chậm hơn.",
        "Khi nói về RAG, cần nhấn mạnh: dữ liệu pháp luật được chuẩn hóa, chunking, embedding, lưu ChromaDB, retrieve, rerank và đóng gói context cho Agent.",
        "Khi nói về Agent, cần nhấn mạnh hệ thống có graph kiểm soát luồng, không để LLM tự do quyết định hoàn toàn.",
        "Phần thực nghiệm nên trình bày trung thực: kết quả tốt ở độ chính xác và chất lượng câu trả lời, nhưng hạn chế còn nằm ở latency.",
    ]:
        parts.append(f"<li>{escape(item)}</li>")
    parts.append("</ul>")
    parts.append("<h2>Kịch bản chi tiết theo slide</h2>")
    for slide in slides:
        parts.append(f'<h3>Slide {slide["no"]}. {escape(slide["title"])}</h3>')
        parts.append(paragraph(f'Thời lượng gợi ý: {slide["time"]}', "time"))
        parts.append(paragraph("Lời thoại gợi ý:", "label"))
        for t in slide["talk"]:
            parts.append(paragraph(t, "talk"))
        if slide.get("transition"):
            parts.append(paragraph("Câu chuyển: " + slide["transition"], "transition"))
        if slide.get("note"):
            parts.append(paragraph("Lưu ý: " + slide["note"], "note"))
    parts.append("</body></html>")
    return "\n".join(parts)


def set_run_font(run, size: int = 13, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_body_paragraph(document: Document, text: str, *, italic: bool = False) -> None:
    paragraph_obj = document.add_paragraph()
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_obj.paragraph_format.first_line_indent = Cm(0.75)
    paragraph_obj.paragraph_format.space_after = Pt(6)
    paragraph_obj.paragraph_format.line_spacing = 1.25
    run = paragraph_obj.add_run(text)
    set_run_font(run, italic=italic)


def add_label(document: Document, text: str) -> None:
    paragraph_obj = document.add_paragraph()
    paragraph_obj.paragraph_format.space_before = Pt(2)
    paragraph_obj.paragraph_format.space_after = Pt(2)
    run = paragraph_obj.add_run(text)
    set_run_font(run, bold=True, color=RGBColor(12, 29, 75))


def add_slide_section(document: Document, slide: dict) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(f'Slide {slide["no"]}. {slide["title"]}')
    set_run_font(run, size=14, bold=True, color=RGBColor(12, 29, 75))

    time_para = document.add_paragraph()
    time_para.paragraph_format.space_after = Pt(4)
    time_run = time_para.add_run(f'Thời lượng gợi ý: {slide["time"]}')
    set_run_font(time_run, size=12, bold=True, color=RGBColor(80, 80, 80))

    add_label(document, "Lời thoại gợi ý:")
    for text in slide["talk"]:
        add_body_paragraph(document, text)

    if slide.get("transition"):
        add_label(document, "Câu chuyển:")
        add_body_paragraph(document, slide["transition"], italic=True)

    if slide.get("note"):
        add_label(document, "Lưu ý:")
        add_body_paragraph(document, slide["note"], italic=True)


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("KỊCH BẢN LUYỆN NÓI BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
    set_run_font(title_run, size=16, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Đề tài: Nghiên cứu phát triển trợ lý ảo pháp luật cho chuyển đổi số")
    set_run_font(subtitle_run, size=13, italic=True)

    add_label(document, "Mục tiêu tài liệu:")
    add_body_paragraph(
        document,
        "Tài liệu này hỗ trợ luyện trình bày theo từng slide trong khoảng 12-15 phút. Nội dung được viết theo văn nói, bám theo slide hiện tại và các nội dung chính trong báo cáo: chatbot, AI Agent, LangGraph, Agentic RAG, xây dựng kho dữ liệu pháp luật, cập nhật văn bản từ Web Admin, thực nghiệm và kết luận.",
    )

    add_label(document, "Gợi ý nhịp trình bày chung:")
    for item in [
        "Không đọc nguyên văn toàn bộ chữ trên slide; dùng slide làm điểm tựa để nói mạch logic.",
        "Các slide 14-25 là phần trọng tâm kỹ thuật, nên nói rõ hơn và chậm hơn.",
        "Khi nói về RAG, cần nhấn mạnh: dữ liệu pháp luật được chuẩn hóa, chunking, embedding, lưu ChromaDB, retrieve, rerank và đóng gói context cho Agent.",
        "Khi nói về Agent, cần nhấn mạnh hệ thống có graph kiểm soát luồng, không để LLM tự do quyết định hoàn toàn.",
        "Phần thực nghiệm nên trình bày trung thực: kết quả tốt ở độ chính xác và chất lượng câu trả lời, nhưng hạn chế còn nằm ở latency.",
    ]:
        paragraph_obj = document.add_paragraph(style=None)
        paragraph_obj.paragraph_format.left_indent = Cm(0.5)
        paragraph_obj.paragraph_format.space_after = Pt(3)
        run = paragraph_obj.add_run("- " + item)
        set_run_font(run)

    section_heading = document.add_paragraph()
    section_heading.paragraph_format.space_before = Pt(10)
    section_heading.paragraph_format.space_after = Pt(6)
    section_heading_run = section_heading.add_run("KỊCH BẢN CHI TIẾT THEO SLIDE")
    set_run_font(section_heading_run, size=14, bold=True, color=RGBColor(12, 29, 75))

    for slide in slides:
        add_slide_section(document, slide)

    document.save(OUT_DOCX)


def main() -> None:
    OUT.write_text(build_html(), encoding="utf-8")
    build_docx()
    print(OUT)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
