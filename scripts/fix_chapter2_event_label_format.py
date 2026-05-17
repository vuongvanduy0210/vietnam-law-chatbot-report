from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TARGETS = [
    ROOT / "Vương Văn Duy_Báo cáo.docx",
    ROOT / "Bao_Cao" / "02_Chuong_2_ban_dinh_dang_theo_mau.docx",
]


def set_times_12_black(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), "Times New Roman")


def normalize_event_label(path: Path) -> int:
    doc = Document(path)
    changed = 0
    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:4] != ["Event", "Thời điểm phát", "Nội dung chính", "Vai trò trên giao diện"]:
            continue

        for row in table.rows[1:]:
            if len(row.cells) < 2:
                continue
            cell = row.cells[1]
            if cell.text.strip() == "Sau khi lưu user message":
                paragraph = cell.paragraphs[0]
                for extra in cell.paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)
                for run in paragraph.runs:
                    run.text = ""
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = "Sau khi lưu tin nhắn người dùng"
                set_times_12_black(run)
                changed += 1
    if changed:
        doc.save(path)
    return changed


def main() -> None:
    for target in TARGETS:
        if target.exists():
            print(target, normalize_event_label(target))


if __name__ == "__main__":
    main()
