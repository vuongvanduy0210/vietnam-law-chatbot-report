from pathlib import Path

from docx import Document
from docx.shared import RGBColor


DOCX = Path(__file__).resolve().parents[1] / "02_Chuong_2_ban_dinh_dang_theo_mau.docx"
BLACK = RGBColor(0, 0, 0)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def force_black(doc):
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            try:
                style.font.color.rgb = BLACK
            except Exception:
                pass

    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            run.font.color.rgb = BLACK


def main():
    doc = Document(DOCX)

    # Only remove Markdown divider leftovers used as standalone separators.
    # Keep meaningful symbols such as arrows, bullets, formulas, and em dashes.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "---":
            remove_paragraph(paragraph)

    force_black(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
