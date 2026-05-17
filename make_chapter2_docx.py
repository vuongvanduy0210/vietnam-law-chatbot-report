#!/usr/bin/env python3
"""
Generate 02_Chuong_2.docx from 02_Chuong_2.md.

This reuses the Chapter 1 formatter but overrides image insertion so tall
Mermaid diagrams are constrained to the printable page area for draft review.
"""

import struct
from pathlib import Path

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import make_chapter1_docx as base

HERE = Path(__file__).parent
MD_FILE = HERE / "02_Chuong_2.md"
OUT_FILE = HERE / "02_Chuong_2.docx"
IMG_DIR = HERE / "sample_images"

MAX_IMG_WIDTH_CM = 16.5
MAX_IMG_HEIGHT_CM = 21.0


def _png_size(path: Path) -> tuple[int, int] | None:
    try:
        with path.open("rb") as f:
            sig = f.read(8)
            if sig != b"\x89PNG\r\n\x1a\n":
                return None
            length = struct.unpack(">I", f.read(4))[0]
            chunk_type = f.read(4)
            if chunk_type != b"IHDR" or length < 8:
                return None
            width, height = struct.unpack(">II", f.read(8))
            return width, height
    except Exception:
        return None


def chapter2_image_para(doc, filename):
    img_path = IMG_DIR / filename
    if not img_path.exists():
        base.normal_para(doc, f"[Hình: {filename} - không tìm thấy file]")
        return

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base._para_fmt(p, before=12, after=0, indent=0)

    size = _png_size(img_path)
    run = p.add_run()
    if not size:
        run.add_picture(str(img_path), width=Cm(MAX_IMG_WIDTH_CM))
        return

    px_w, px_h = size
    aspect = px_w / px_h if px_h else 1.0

    width_cm = MAX_IMG_WIDTH_CM
    height_cm = width_cm / aspect
    if height_cm > MAX_IMG_HEIGHT_CM:
        height_cm = MAX_IMG_HEIGHT_CM
        width_cm = height_cm * aspect

    run.add_picture(str(img_path), width=Cm(width_cm), height=Cm(height_cm))


def main():
    base.MD_FILE = MD_FILE
    base.OUT_FILE = OUT_FILE
    base.IMG_DIR = IMG_DIR
    base.image_para = chapter2_image_para

    lines = MD_FILE.read_text(encoding="utf-8").splitlines()
    doc = base.build_doc(lines)
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    main()
